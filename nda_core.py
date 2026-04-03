#!/usr/bin/env python3
"""
Smart Attorney NDA Redlining System v3
Rules-engine architecture with surgical edits.
"""

import os
import sys
import subprocess
import json
import secrets
import re
import shutil
import time
from datetime import datetime

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = "/home/cliff/redact/TestOutput"


class ConfigLoader:
    """Loads client config, prompts, and rule templates from external files."""

    def __init__(self, config_name="default_receiver"):
        self.config = self._load_json(f"config/{config_name}.json")
        self.templates = self._load_json("config/rule_templates.json")
        self.prompt = self._load_text("prompts/analysis_prompt.md")

    def _load_json(self, rel_path):
        path = os.path.join(BASE_DIR, rel_path)
        with open(path, 'r') as f:
            return json.load(f)

    def _load_text(self, rel_path):
        path = os.path.join(BASE_DIR, rel_path)
        with open(path, 'r') as f:
            return f.read()


class PIIRedactor:
    """Redacts and restores PII using hex-mapped placeholders."""

    LEGAL_SUFFIXES = {'llc', 'inc', 'corp', 'corporation', 'company', 'lp', 'llp', 'limited', 'ltd'}
    DESC_SUFFIXES = {'center', 'centre', 'clinic', 'lab', 'laboratory', 'foundation', 'trust',
                     'fund', 'enterprises', 'holdings', 'services', 'solutions', 'properties',
                     'realty', 'management', 'consulting', 'international', 'associates', 'group'}

    def __init__(self):
        self.whitelist = self._load_whitelist()

    def _load_whitelist(self):
        whitelist = set()
        wl_path = "/home/cliff/redact/redaction_whitelist.txt"
        if os.path.exists(wl_path):
            with open(wl_path, 'r') as f:
                for line in f:
                    word = line.strip().lower()
                    if word and not word.startswith('#'):
                        whitelist.add(word)
        return whitelist

    def _extract_entities(self, text):
        """Find business entity names from the preamble (before numbered clauses)."""
        # Only look in the preamble — before numbered clauses start
        preamble = text
        for m in re.finditer(r'^\s*\d+[\.\)]\s', text, re.MULTILINE):
            preamble = text[:m.start()]
            break

        all_suffixes = self.LEGAL_SUFFIXES | self.DESC_SUFFIXES
        pat = r'\b(?:' + '|'.join(re.escape(s) for s in sorted(all_suffixes, key=len, reverse=True)) + r')\.?\b'

        suffix_matches = list(re.finditer(pat, preamble, re.IGNORECASE))
        consumed = set()
        entities = []

        for i, m in enumerate(suffix_matches):
            if i in consumed:
                continue

            entity_end = m.end()

            if m.group().rstrip('.').lower() in self.DESC_SUFFIXES:
                after = preamble[m.end():]
                legal_follow = re.match(
                    r',\s*(' + '|'.join(re.escape(s) for s in self.LEGAL_SUFFIXES) + r')\b',
                    after, re.IGNORECASE)
                if legal_follow:
                    entity_end = m.end() + legal_follow.end()
                    for j in range(i + 1, len(suffix_matches)):
                        if suffix_matches[j].start() < entity_end:
                            consumed.add(j)

            before = preamble[:m.start()]
            best_start = 0
            for sep in re.finditer(r'(?:,\s*)|(?:(?:WHEREAS|and|between)\s+)', before, re.IGNORECASE):
                if sep.group().strip() == ',':
                    after_sep = preamble[sep.end():].lstrip()
                    first_word = after_sep.split()[0].rstrip('.,').lower() if after_sep.split() else ''
                    if first_word in self.LEGAL_SUFFIXES:
                        continue
                best_start = sep.end()

            entity = preamble[best_start:entity_end].strip().strip(',').strip()
            if entity and len(entity.split()) >= 2 and len(entity) <= 80:
                entities.append(entity)

        return entities

    def redact_text(self, text):
        """Returns (redacted_text, hex_mapping)."""
        # Company suffixes used in multiple patterns
        co_suffix = r'(?:LLC|L\.L\.C\.|Inc\.?|Corp\.?|Corporation|Company|Co\.|LP|L\.P\.|LLP|Associates|Group|Center|Centre|Clinic|Lab|Laboratory|Foundation|Trust|Fund|Enterprises|Holdings|Services|Solutions|Properties|Realty|Management|Consulting|International|Limited|Ltd\.?)'

        patterns = [
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL'),
            (r'\(\d{3}\)\s*\d{3}-\d{4}', 'PHONE'),
            (r'\d{3}-\d{3}-\d{4}', 'PHONE'),
            # Company FIRST (longest matches) — number-prefixed (e.g., "4339 Main Street, LLC")
            (r'\b\d+\s+[A-Za-z][A-Za-z\s&,.-]{1,40},?\s*' + co_suffix + r'\b', 'COMPANY'),
            # Company: standard names, allows comma before suffix, title or ALL CAPS
            (r'\b[A-Z][A-Za-z\s&,.-]{1,50},?\s*' + co_suffix + r'\b', 'COMPANY'),
            # Address (after company so "4339 Main Street, LLC" isn't split)
            (r'\b\d+\s+[A-Za-z]+\s+(?:Street|St\.?|Avenue|Ave\.?|Road|Rd\.?|Drive|Dr\.?|Lane|Ln\.?|Court|Ct\.?|Boulevard|Blvd\.?|Way|Place|Pl\.?|Circle|Cir\.?|Parkway|Pkwy\.?|Highway|Hwy\.?|STREET|ST|AVENUE|AVE|ROAD|RD|DRIVE|DR|LANE|LN|COURT|CT|BOULEVARD|BLVD)(?:\s*,?\s*(?:Suite|Ste\.?|Apt\.?|Unit|#|SUITE|STE|APT|UNIT)\s*\w+)?(?:\s*,\s*[A-Za-z]+(?:\s+[A-Za-z]+)*)?(?:\s*,\s*[A-Z]{2})?\s*\d{5}(?:-\d{4})?\b', 'ADDRESS'),
            # Street (after company)
            (r'\b\d+\s+[A-Za-z]+\s+(?:Street|St\.?|Avenue|Ave\.?|Road|Rd\.?|Drive|Dr\.?|Lane|Ln\.?|Court|Ct\.?|Boulevard|Blvd\.?|Way|Place|Pl\.?|Circle|Cir\.?|Parkway|Pkwy\.?|Highway|Hwy\.?|STREET|ST|AVENUE|AVE|ROAD|RD|DRIVE|DR|LANE|LN|COURT|CT|BOULEVARD|BLVD)\b', 'STREET'),
            (r'\b\d{5}(?:-\d{4})?\b', 'ZIP'),
            (r'\b\d{1,2}/\d{1,2}/\d{4}\b', 'DATE'),
            (r'\b\d{1,2}-\d{1,2}-\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', 'PERSON'),
        ]
        redacted = text
        mapping = {}
        found_pii = []  # collect (original_text, label) for second pass

        # Pre-pass: extract full entity names by tracing back from suffixes
        entities = self._extract_entities(text)
        for entity in entities:
            for match in list(re.finditer(re.escape(entity), redacted, re.IGNORECASE)):
                original = match.group()
                hex_id = secrets.token_hex(8)
                placeholder = f"[COMPANY:{hex_id}]"
                mapping[hex_id] = {'type': 'COMPANY', 'original': original, 'placeholder': placeholder}
                redacted = redacted.replace(original, placeholder, 1)
                found_pii.append((original, 'COMPANY'))

        # First pass: regex pattern matching
        for pattern, label in patterns:
            for match in list(re.finditer(pattern, redacted)):
                original = match.group()
                if label == 'PERSON' and all(w.lower() in self.whitelist for w in original.split()):
                    continue
                hex_id = secrets.token_hex(8)
                placeholder = f"[{label}:{hex_id}]"
                mapping[hex_id] = {'type': label, 'original': original, 'placeholder': placeholder}
                redacted = redacted.replace(original, placeholder, 1)
                found_pii.append((original, label))

        # Second pass: case-insensitive sweep for all found PII
        # Also generate sub-phrases from long matches (e.g., "Midwest Fertility Center"
        # from "Midwest Fertility Center and Ambulatory Surgery Center")
        # And individual non-whitelisted proper nouns from PERSON/COMPANY matches
        search_terms = []
        for original, label in found_pii:
            search_terms.append((original, label))
            if label == 'COMPANY':
                # Normalize whitespace (newlines→spaces) before splitting
                normalized = re.sub(r'\s+', ' ', original)
                for sep in [' and ', ' & ', ', ']:
                    if sep in normalized:
                        for part in normalized.split(sep):
                            part = part.strip().strip(',')
                            if len(part.split()) >= 2:
                                search_terms.append((part, label))
            # Individual proper nouns (catches "Downers" from "Downers Grove")
            if label in ('PERSON', 'COMPANY'):
                for word in original.split():
                    w = re.sub(r'[,.]', '', word)
                    if w and w[0].isupper() and w.lower() not in self.whitelist and len(w) > 2:
                        search_terms.append((w, label))

        for original, label in search_terms:
            for match in list(re.finditer(re.escape(original), redacted, re.IGNORECASE)):
                variant = match.group()
                hex_id = secrets.token_hex(8)
                placeholder = f"[{label}:{hex_id}]"
                mapping[hex_id] = {'type': label, 'original': variant, 'placeholder': placeholder}
                redacted = redacted.replace(variant, placeholder, 1)

        return redacted, mapping

    def redact_docx(self, docx_path, base_name):
        """Redact PII in a DOCX copy. Returns (redacted_path, mapping_path)."""
        try:
            from docx import Document
            redacted_path = f"/tmp/{base_name}_redacted.docx"
            mapping_path = f"/tmp/{base_name}_mapping.json"
            shutil.copy2(docx_path, redacted_path)

            text = self._extract_text(docx_path)
            if not text:
                return None, None

            _, mapping = self.redact_text(text)
            if not mapping:
                print("   ✓ No PII found to redact")
                with open(mapping_path, 'w') as f:
                    json.dump({}, f)
                return redacted_path, mapping_path

            doc = Document(redacted_path)
            for hex_id, data in mapping.items():
                for para in doc.paragraphs:
                    if data['original'] in para.text:
                        # Try run-level replacement first
                        found_in_run = False
                        for run in para.runs:
                            if data['original'] in run.text:
                                run.text = run.text.replace(data['original'], data['placeholder'])
                                found_in_run = True
                        # Fallback: PII spans multiple runs — merge and rewrite
                        if not found_in_run and para.runs:
                            full = para.text.replace(data['original'], data['placeholder'])
                            para.runs[0].text = full
                            for run in para.runs[1:]:
                                run.text = ""
            doc.save(redacted_path)
            with open(mapping_path, 'w') as f:
                json.dump(mapping, f, indent=2)
            print(f"   ✓ Redacted {len(mapping)} PII items")
            return redacted_path, mapping_path
        except Exception as e:
            print(f"   ❌ Redaction failed: {e}")
            return None, None

    def restore_docx(self, docx_path, mapping_path):
        """Restore PII in a DOCX from hex mapping."""
        try:
            from docx import Document
            with open(mapping_path, 'r') as f:
                mapping = json.load(f)
            if not mapping:
                return True
            doc = Document(docx_path)
            restored = 0
            for para in doc.paragraphs:
                for hex_id, data in mapping.items():
                    if data['placeholder'] in para.text:
                        found = False
                        for run in para.runs:
                            if data['placeholder'] in run.text:
                                run.text = run.text.replace(data['placeholder'], data['original'])
                                found = True
                                restored += 1
                        if not found and para.runs:
                            full = para.text.replace(data['placeholder'], data['original'])
                            para.runs[0].text = full
                            for run in para.runs[1:]:
                                run.text = ""
                            restored += 1
            doc.save(docx_path)
            print(f"   ✓ Restored {restored} PII items")
            return True
        except Exception as e:
            print(f"   ❌ PII restoration failed: {e}")
            return False

    def _extract_text(self, docx_path):
        try:
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'txt', '--outdir', '/tmp', docx_path],
                capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                txt_file = f"/tmp/{os.path.splitext(os.path.basename(docx_path))[0]}.txt"
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        return f.read()
            return None
        except Exception as e:
            print(f"   ❌ Text extraction error: {e}")
            return None
