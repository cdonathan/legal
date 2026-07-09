"""
Parse HTML from the browser editor and apply it to DOCX paragraphs
with proper run-level formatting (bold, italic, underline, strikethrough, font size).
"""

import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor
from typing import Optional
import logging

logger = logging.getLogger(__name__)

# Map browser font size values (1-7) to point sizes
FONT_SIZE_MAP = {
    '1': 8, '2': 10, '3': 11, '4': 14, '5': 18, '6': 24, '7': 36,
}


class RunSpec:
    """Describes a single text run with its formatting."""
    def __init__(self, text: str = "", bold: bool = False, italic: bool = False,
                 underline: bool = False, strike: bool = False, font_size: Optional[int] = None):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.strike = strike
        self.font_size = font_size


class HtmlToRunsParser(HTMLParser):
    """Parse HTML into a list of RunSpecs with formatting."""

    def __init__(self):
        super().__init__()
        self.runs: list[RunSpec] = []
        self.format_stack: list[dict] = [{}]  # Stack of active formatting
        self._current_text = ""

    def _current_format(self) -> dict:
        """Get merged formatting from the stack."""
        merged = {}
        for frame in self.format_stack:
            merged.update(frame)
        return merged

    def _flush_text(self):
        """Save accumulated text as a run with current formatting."""
        if self._current_text:
            fmt = self._current_format()
            self.runs.append(RunSpec(
                text=self._current_text,
                bold=fmt.get('bold', False),
                italic=fmt.get('italic', False),
                underline=fmt.get('underline', False),
                strike=fmt.get('strike', False),
                font_size=fmt.get('font_size', None),
            ))
            self._current_text = ""

    def handle_starttag(self, tag, attrs):
        self._flush_text()
        attrs_dict = dict(attrs)
        fmt = {}

        if tag in ('strong', 'b'):
            fmt['bold'] = True
        elif tag in ('em', 'i'):
            fmt['italic'] = True
        elif tag == 'u':
            fmt['underline'] = True
        elif tag in ('s', 'del', 'strike'):
            fmt['strike'] = True
        elif tag == 'span':
            # Check for inline styles
            style = attrs_dict.get('style', '')
            if 'font-weight' in style and ('bold' in style or '700' in style):
                fmt['bold'] = True
            if 'font-style' in style and 'italic' in style:
                fmt['italic'] = True
            if 'text-decoration' in style:
                if 'underline' in style:
                    fmt['underline'] = True
                if 'line-through' in style:
                    fmt['strike'] = True
            # Check class for our highlight spans (ignore formatting)
            cls = attrs_dict.get('class', '')
            if 'highlight' in cls:
                pass  # Don't add any formatting for highlight spans
        elif tag == 'font':
            size = attrs_dict.get('size', '')
            if size in FONT_SIZE_MAP:
                fmt['font_size'] = FONT_SIZE_MAP[size]
        elif tag == 'br':
            self._current_text += "\n"
            return  # br is self-closing, don't push to stack

        self.format_stack.append(fmt)

    def handle_endtag(self, tag):
        self._flush_text()
        if tag != 'br' and len(self.format_stack) > 1:
            self.format_stack.pop()

    def handle_data(self, data):
        self._current_text += data

    def handle_entityref(self, name):
        entities = {'amp': '&', 'lt': '<', 'gt': '>', 'quot': '"', 'nbsp': ' '}
        self._current_text += entities.get(name, f'&{name};')

    def handle_charref(self, name):
        if name.startswith('x'):
            self._current_text += chr(int(name[1:], 16))
        else:
            self._current_text += chr(int(name))

    def get_runs(self) -> list[RunSpec]:
        self._flush_text()
        return self.runs


def parse_html_to_runs(html: str) -> list[RunSpec]:
    """Parse an HTML string into a list of formatted text runs."""
    parser = HtmlToRunsParser()
    parser.feed(html)
    return parser.get_runs()


def apply_manual_edits_to_docx(docx_path: str, edits: list[dict], output_path: str):
    """
    Apply manual edits (with HTML formatting) to a DOCX file.

    Each edit has:
        - paragraph_index: which paragraph to modify (-1 for new)
        - insert_after: for new paragraphs, insert after this index
        - original: original plain text (for find/match)
        - new_text: new plain text
        - new_html: HTML with formatting from the browser editor
    """
    doc = Document(docx_path)

    for edit in edits:
        para_index = edit.get('paragraph_index', -1)
        new_html = edit.get('new_html', '')
        original = edit.get('original', '')

        if not new_html:
            continue

        # Parse HTML into runs
        runs = parse_html_to_runs(new_html)
        if not runs:
            continue

        if para_index >= 0 and para_index < len(doc.paragraphs):
            # Modify existing paragraph
            para = doc.paragraphs[para_index]
            _replace_paragraph_runs(para, runs)
        elif original:
            # Find paragraph by original text content
            for para in doc.paragraphs:
                if para.text.strip() == original.strip():
                    _replace_paragraph_runs(para, runs)
                    break
        else:
            # New paragraph — insert at end or after specified index
            insert_after = edit.get('insert_after', -1)
            if insert_after >= 0 and insert_after < len(doc.paragraphs):
                # Insert after the specified paragraph
                ref_para = doc.paragraphs[insert_after]
                new_para = _insert_paragraph_after(ref_para)
                _set_paragraph_runs(new_para, runs)
            else:
                # Append at end
                new_para = doc.add_paragraph()
                _set_paragraph_runs(new_para, runs)

    doc.save(output_path)
    logger.info(f"Applied {len(edits)} manual edits to {output_path}")


def _replace_paragraph_runs(para, runs: list[RunSpec]):
    """Replace all runs in a paragraph with new formatted runs."""
    # Clear existing runs
    for run in para.runs:
        run.text = ""
    # Remove all run elements
    for child in list(para._p):
        from docx.oxml.ns import qn
        if child.tag == qn('w:r'):
            para._p.remove(child)

    # Add new runs
    _set_paragraph_runs(para, runs)


def _set_paragraph_runs(para, runs: list[RunSpec]):
    """Add formatted runs to a paragraph."""
    for run_spec in runs:
        if not run_spec.text:
            continue
        # Handle newlines as separate runs
        parts = run_spec.text.split('\n')
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                if run_spec.bold:
                    run.bold = True
                if run_spec.italic:
                    run.italic = True
                if run_spec.underline:
                    run.underline = True
                if run_spec.strike:
                    run.font.strike = True
                if run_spec.font_size:
                    run.font.size = Pt(run_spec.font_size)
            # Add line break between parts (not after last)
            if i < len(parts) - 1:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                run = para.add_run()
                br = OxmlElement('w:br')
                run._r.append(br)


def _insert_paragraph_after(ref_para):
    """Insert a new paragraph after the reference paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)

    # Create a proper Paragraph object
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())
