"""
Convert DOCX to structured HTML for browser preview and editing.
Uses python-docx to read run-level formatting accurately.
Returns paragraph-indexed HTML for diff tracking on save.
"""

import os
from xml.sax.saxutils import escape
from docx import Document
from docx.oxml.ns import qn


def convert_docx_to_html(docx_path: str) -> dict:
    """
    Convert a DOCX file to HTML preserving formatting.
    Returns:
        {
            "html": full HTML string for rendering,
            "paragraphs": [{"index": 0, "text": "plain text", "style": "Heading 1"}, ...]
        }
    Paragraphs list is used for diff tracking — each paragraph has its plain text
    and index so edits can be mapped back to the original DOCX.
    """
    doc = Document(docx_path)
    html_parts = []
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        plain_text = para.text
        para_html = _para_to_html(para, i)
        if para_html:
            html_parts.append(para_html)
        paragraphs.append({
            "index": i,
            "text": plain_text,
            "style": para.style.name if para.style else "Normal"
        })

    # Handle tables
    for table in doc.tables:
        html_parts.append(_table_to_html(table))

    return {
        "html": "\n".join(html_parts),
        "paragraphs": paragraphs
    }


def _para_to_html(para, index: int) -> str:
    """Convert a paragraph to an HTML element with data-index for tracking."""
    runs_html = "".join(_run_to_html(r) for r in para.runs)
    if not runs_html.strip():
        return ""

    style_name = para.style.name.lower() if para.style else ""

    # Determine element type and CSS class
    if "heading 1" in style_name or style_name == "title page":
        tag = "h2"
        cls = "doc-heading1"
    elif "heading 2" in style_name:
        tag = "h3"
        cls = "doc-heading2"
    elif "heading 3" in style_name:
        tag = "h4"
        cls = "doc-heading3"
    elif "toc" in style_name:
        tag = "p"
        cls = "doc-toc"
    elif "center" in style_name:
        tag = "p"
        cls = "doc-center"
    elif "title" in style_name:
        tag = "h1"
        cls = "doc-title"
    else:
        tag = "p"
        cls = "doc-body"

    # Check indentation
    indent_style = ""
    if para.paragraph_format and para.paragraph_format.left_indent:
        indent_pt = para.paragraph_format.left_indent.pt
        if indent_pt and indent_pt > 0:
            indent_style = f" style=\"margin-left:{int(indent_pt * 0.75)}px\""

    return f'<{tag} class="{cls}" data-para-index="{index}"{indent_style}>{runs_html}</{tag}>'


def _run_to_html(run) -> str:
    """Convert a single run to HTML with inline formatting."""
    text = escape(run.text)
    if not text:
        return ""

    # Check XML directly for bold — avoids style inheritance issues
    rpr = run._r.find(qn('w:rPr'))
    is_bold = False
    is_italic = False
    is_underline = False
    is_strike = False

    if rpr is not None:
        b = rpr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            is_bold = (val is None or (val != '0' and val != 'false'))

        i = rpr.find(qn('w:i'))
        if i is not None:
            val = i.get(qn('w:val'))
            is_italic = (val is None or (val != '0' and val != 'false'))

        u = rpr.find(qn('w:u'))
        if u is not None:
            val = u.get(qn('w:val'))
            is_underline = (val is not None and val != 'none')

        strike = rpr.find(qn('w:strike'))
        if strike is not None:
            val = strike.get(qn('w:val'))
            is_strike = (val is None or (val != '0' and val != 'false'))

    if is_strike:
        text = f'<del>{text}</del>'
    if is_bold:
        text = f'<strong>{text}</strong>'
    if is_italic:
        text = f'<em>{text}</em>'
    if is_underline:
        text = f'<u>{text}</u>'

    return text


def _table_to_html(table) -> str:
    """Convert a table to HTML."""
    html = '<table class="doc-table">'
    for row in table.rows:
        html += "<tr>"
        for cell in row.cells:
            cell_text = escape(cell.text)
            html += f"<td>{cell_text}</td>"
        html += "</tr>"
    html += "</table>"
    return html
