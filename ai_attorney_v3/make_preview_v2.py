#!/usr/bin/env python3
"""
Generate preview HTML from DOCX using python-docx directly.
Reads actual run-level formatting (bold, italic, underline) without
misinterpreting paragraph style inheritance.
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from xml.sax.saxutils import escape

input_file = sys.argv[1] if len(sys.argv) > 1 else "/home/cliff/redact/ai_attorney/jobs/70fe68c1-3d61-4905-b9fe-e36ca6a2c23d/PSA_Form2 (1).docx"
output_file = sys.argv[2] if len(sys.argv) > 2 else "/home/cliff/redact/ai_attorney_v3/static/preview_example_psa.html"

doc = Document(input_file)

def run_to_html(run):
    """Convert a single run to HTML with inline formatting."""
    text = escape(run.text)
    if not text:
        return ""
    
    # Only apply bold/italic if explicitly set on the run (not inherited from style)
    is_bold = run.bold is True  # Explicitly True, not None (inherited)
    is_italic = run.italic is True
    is_underline = run.underline is True
    
    if is_bold:
        text = f"<strong>{text}</strong>"
    if is_italic:
        text = f"<em>{text}</em>"
    if is_underline:
        text = f"<u>{text}</u>"
    
    return text


def para_to_html(para):
    """Convert a paragraph to HTML."""
    # Determine paragraph type from style
    style_name = para.style.name.lower() if para.style else ""
    
    # Get paragraph text
    runs_html = "".join(run_to_html(r) for r in para.runs)
    if not runs_html.strip():
        return ""
    
    # Map styles to HTML elements
    if "heading 1" in style_name or "title" in style_name:
        return f"<h1>{runs_html}</h1>"
    elif "heading 2" in style_name or "subtitle" in style_name:
        return f"<h2>{runs_html}</h2>"
    elif "heading 3" in style_name:
        return f"<h3>{runs_html}</h3>"
    elif "toc" in style_name:
        return f'<p class="toc">{runs_html}</p>'
    elif "center" in style_name:
        return f'<p style="text-align:center">{runs_html}</p>'
    else:
        # Check indentation
        indent = ""
        if para.paragraph_format.left_indent:
            indent_pt = para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0
            if indent_pt > 0:
                indent = f' style="margin-left:{int(indent_pt * 0.75)}px"'
        return f"<p{indent}>{runs_html}</p>"


# Build HTML
body_parts = []
for para in doc.paragraphs:
    html = para_to_html(para)
    if html:
        body_parts.append(html)

# Also handle tables
for table in doc.tables:
    table_html = "<table>"
    for row in table.rows:
        table_html += "<tr>"
        for cell in row.cells:
            cell_text = escape(cell.text)
            table_html += f"<td>{cell_text}</td>"
        table_html += "</tr>"
    table_html += "</table>"
    body_parts.append(table_html)

body = "\n".join(body_parts)

html_doc = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Document Preview — python-docx Rendering</title>
<style>
  * {{ box-sizing: border-box; }}
  body {{
    font-family: 'Times New Roman', Georgia, serif;
    max-width: 860px;
    margin: 0 auto;
    padding: 40px 60px;
    line-height: 1.6;
    color: #000;
    background: white;
    font-size: 11pt;
  }}
  h1 {{ font-size: 14pt; text-align: center; font-weight: bold; margin: 20px 0 10px; text-transform: uppercase; }}
  h2 {{ font-size: 12pt; font-weight: bold; margin: 18px 0 8px; }}
  h3 {{ font-size: 11pt; font-weight: bold; margin: 14px 0 6px; }}
  p {{ margin: 4px 0; text-align: justify; }}
  p.toc {{ margin: 2px 0; color: #444; font-size: 10pt; }}
  strong {{ font-weight: bold; }}
  em {{ font-style: italic; }}
  u {{ text-decoration: underline; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 10pt; }}
  td, th {{ border: 1px solid #999; padding: 5px 8px; vertical-align: top; }}
  .note {{
    background: #fffbeb;
    border: 1px solid #f59e0b;
    border-radius: 6px;
    padding: 12px 16px;
    margin-bottom: 24px;
    font-family: -apple-system, sans-serif;
    font-size: 0.85rem;
    color: #78350f;
  }}
</style>
</head>
<body>
<div class="note">
  <strong>Option A Preview (python-docx rendering)</strong> — Bold/italic applied only when explicitly set on text runs, not inherited from paragraph styles.
</div>

{body}

</body>
</html>"""

with open(output_file, "w", encoding="utf-8") as f:
    f.write(html_doc)

print(f"Preview written to: {output_file}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"HTML body size: {len(body):,} chars")
