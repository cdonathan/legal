"""
Document Processor — file conversion, text extraction, DOCX modification, output generation.
Handles redlined and clean document output with PII reconstruction.
"""

import os
import re
import subprocess
import shutil
import logging
from typing import Optional

from models import ProposedChange
from text_utils import fix_ligatures

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles file I/O, DOCX modification, and output generation."""

    def extract_text(self, file_path: str) -> str:
        """
        Extract plaintext from supported file formats.
        Supports: .docx, .pdf, .txt, .mhtml/.mht
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return fix_ligatures(f.read())

        if ext in (".mhtml", ".mht"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            text = re.sub(r"<[^>]+>", " ", content)
            return fix_ligatures(re.sub(r"\s+", " ", text).strip())

        if ext == ".pdf":
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return fix_ligatures("\n".join(parts))

        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return fix_ligatures(text)

        raise ValueError(f"Unsupported file type: {ext}")

    def apply_changes(
        self,
        docx_path: str,
        changes: list[ProposedChange],
        output_path: str,
        redline: bool = False
    ):
        """
        Apply accepted changes to a DOCX file.
        - redline=True: strikethrough (red) old text + underline (green) new text
        - redline=False: clean replacement, no markup
        """
        from docx import Document
        from docx.shared import RGBColor

        shutil.copy2(docx_path, output_path)
        doc = Document(output_path)

        # Fix ligatures in all paragraphs first
        for para in doc.paragraphs:
            orig = para.text
            fixed = fix_ligatures(orig)
            if orig != fixed:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = fixed
                else:
                    para.add_run(fixed)

        # Sort changes by document position (descending) to avoid offset shifts
        sorted_changes = sorted(
            [c for c in changes if c.type == "replace" and c.find and c.replace],
            key=lambda c: c.document_position,
            reverse=True
        )

        # Detect and remove overlapping changes (higher confidence wins)
        non_overlapping = self._remove_overlaps(sorted_changes)

        # Apply changes
        for change in non_overlapping:
            find_text = fix_ligatures(change.find).strip()
            replace_text = change.replace

            if not find_text:
                continue

            matched = False
            for para in doc.paragraphs:
                para_text = fix_ligatures(para.text)

                # Exact match
                if find_text in para_text:
                    if redline:
                        self._redline_replace(para, find_text, replace_text)
                    else:
                        self._clean_replace(para, find_text, replace_text)
                    matched = True
                    break

            # Normalized match fallback
            if not matched:
                find_normalized = re.sub(r'\s+', ' ', find_text).strip()
                find_normalized = re.sub(r'\s+([.,;:!?])', r'\1', find_normalized)

                for para in doc.paragraphs:
                    para_text = fix_ligatures(para.text)
                    para_normalized = re.sub(r'\s+', ' ', para_text).strip()
                    para_normalized = re.sub(r'\s+([.,;:!?])', r'\1', para_normalized)

                    if find_normalized in para_normalized:
                        actual_find = self._find_actual_span(para_text, find_normalized)
                        if actual_find:
                            if redline:
                                self._redline_replace(para, actual_find, replace_text)
                            else:
                                self._clean_replace(para, actual_find, replace_text)
                            matched = True
                            break

            if not matched:
                logger.warning(f"Change {change.id}: Could not locate '{find_text[:50]}...' in document")

        doc.save(output_path)

    def create_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """Generate PDF from DOCX using LibreOffice headless. Returns True on success."""
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", os.path.dirname(pdf_path), docx_path],
                capture_output=True, text=True, timeout=60,
            )
            # LibreOffice names the output based on input filename
            generated = docx_path.replace(".docx", ".pdf")
            if os.path.exists(generated) and generated != pdf_path:
                os.rename(generated, pdf_path)
            return os.path.exists(pdf_path)
        except Exception as e:
            logger.warning(f"PDF generation failed: {e}")
            return False

    def _remove_overlaps(self, changes: list[ProposedChange]) -> list[ProposedChange]:
        """
        Remove overlapping changes. When two changes target overlapping text,
        keep the higher-confidence one.
        """
        confidence_rank = {"exact": 4, "fuzzy": 3, "full_clause": 2, "manual": 1}
        result = []
        used_ranges = []  # List of (start, end) tuples

        # Sort by confidence (highest first), then by position
        sorted_by_conf = sorted(
            changes,
            key=lambda c: (-confidence_rank.get(c.confidence, 0), c.document_position)
        )

        for change in sorted_by_conf:
            start = change.document_position
            end = start + len(change.find)
            overlaps = False

            for used_start, used_end in used_ranges:
                if start < used_end and end > used_start:
                    overlaps = True
                    break

            if not overlaps:
                result.append(change)
                used_ranges.append((start, end))

        # Return in document order (descending for safe replacement)
        return sorted(result, key=lambda c: c.document_position, reverse=True)

    def _clean_replace(self, para, find_text: str, replace_text: str):
        """Replace text preserving paragraph formatting."""
        full_text = fix_ligatures(para.text)
        new_full = full_text.replace(find_text, replace_text, 1)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_full
        else:
            para.add_run(new_full)

    def _redline_replace(self, para, find_text: str, replace_text: str):
        """Replace with redline formatting: strikethrough red old, underline green new."""
        from docx.shared import RGBColor

        full_text = fix_ligatures(para.text)
        idx = full_text.find(find_text)
        if idx < 0:
            return

        before = full_text[:idx]
        after = full_text[idx + len(find_text):]

        # Clear all runs
        for run in para.runs:
            run.text = ""

        # Rebuild with formatting
        if para.runs:
            para.runs[0].text = before
        else:
            para.add_run(before)

        # Strikethrough old text (red)
        strike_run = para.add_run(find_text)
        strike_run.font.strike = True
        strike_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        # Underline new text (green)
        ins_run = para.add_run(replace_text)
        ins_run.font.color.rgb = RGBColor(0x16, 0x63, 0x34)
        ins_run.font.underline = True

        # Rest of paragraph
        para.add_run(after)

    def _find_actual_span(self, para_text: str, normalized_find: str) -> Optional[str]:
        """Find the actual text span in a paragraph matching a normalized search."""
        para_norm = re.sub(r'\s+', ' ', para_text).strip()
        para_norm = re.sub(r'\s+([.,;:!?])', r'\1', para_norm)
        idx = para_norm.find(normalized_find)
        if idx < 0:
            return None

        # Build mapping from normalized to original positions
        norm_to_orig = []
        in_whitespace = False

        orig_start = 0
        while orig_start < len(para_text) and para_text[orig_start] in ' \t\n\r':
            orig_start += 1

        for orig_i in range(orig_start, len(para_text)):
            ch = para_text[orig_i]
            if ch in ' \t\n\r':
                if not in_whitespace:
                    norm_to_orig.append(orig_i)
                    in_whitespace = True
            else:
                norm_to_orig.append(orig_i)
                in_whitespace = False

        if idx + len(normalized_find) <= len(norm_to_orig):
            start = norm_to_orig[idx]
            end_idx = idx + len(normalized_find)
            end = norm_to_orig[end_idx] if end_idx < len(norm_to_orig) else len(para_text)
            return para_text[start:end]

        return None
