#!/usr/bin/env python3
"""
Enhanced LibreOffice track changes test with all change types
"""

import os
from docx import Document
from docx.shared import RGBColor

def create_comprehensive_track_changes_test():
    """Create test document with all types of track changes"""
    
    input_file = "/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Conf_Agr_Sample1-pre-redline.docx"
    output_file = "/home/cliff/redact/redline_project/LibreOffice_ALL_Changes_Test.docx"
    
    print("🔄 Creating comprehensive track changes test...")
    
    doc = Document(input_file)
    
    # Find and modify existing paragraphs
    for i, para in enumerate(doc.paragraphs):
        if i == 1 and "Confidentiality Agreement" in para.text:
            # REPLACEMENT: Replace entire sentence
            para.clear()
            
            # Original text (strikethrough, red)
            del_run = para.add_run("This Confidentiality Agreement")
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Replacement text (underline, green)
            ins_run = para.add_run("This CONFIDENTIALITY AND NON-DISCLOSURE AGREEMENT")
            ins_run.underline = True
            ins_run.font.color.rgb = RGBColor(0, 128, 0)
            
            para.add_run(' (this "Agreement"), dated as of [DATE] ("Effective Date"), is between [COMPANY A] and [COMPANY B].')
            
        elif i == 3 and "connection with" in para.text:
            # WORD INSERTIONS AND DELETIONS within paragraph
            para.clear()
            
            para.add_run("In connection with ")
            
            # Delete word
            del_run = para.add_run("engaging in ")
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Insert word
            ins_run = para.add_run("participating in ")
            ins_run.underline = True
            ins_run.font.color.rgb = RGBColor(0, 128, 0)
            
            para.add_run("the ")
            
            # Delete phrase
            del_run = para.add_run("participation, financing, and/or investment activities")
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Insert replacement phrase
            ins_run = para.add_run("evaluation, due diligence, and potential acquisition activities")
            ins_run.underline = True
            ins_run.font.color.rgb = RGBColor(0, 128, 0)
            
            para.add_run(" concerning the premises located at [PROPERTY ADDRESS] (the \"Purpose\").")
    
    # ADD NEW PARAGRAPHS (insertions)
    new_para1 = doc.add_paragraph()
    label_run = new_para1.add_run("[INSERTED CLAUSE] ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(255, 0, 0)
    
    clause_run = new_para1.add_run("PERMITTED RECIPIENTS: The Receiving Party may disclose Confidential Information to its directors, officers, employees, attorneys, accountants, financial advisors, consultants, lenders, investors, and other professional advisors who have a need to know such information for the Purpose and who are bound by confidentiality obligations.")
    clause_run.underline = True
    clause_run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Add another insertion
    new_para2 = doc.add_paragraph()
    label_run2 = new_para2.add_run("[INSERTED CLAUSE] ")
    label_run2.bold = True
    label_run2.font.color.rgb = RGBColor(255, 0, 0)
    
    clause_run2 = new_para2.add_run("RETURN OF MATERIALS: Upon termination of this Agreement or upon written request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and any copies, notes, analyses, or derivatives thereof.")
    clause_run2.underline = True
    clause_run2.font.color.rgb = RGBColor(0, 128, 0)
    
    # Add a DELETED paragraph
    del_para = doc.add_paragraph()
    del_label = del_para.add_run("[DELETED PARAGRAPH] ")
    del_label.bold = True
    del_label.font.color.rgb = RGBColor(255, 0, 0)
    
    del_text = del_para.add_run("This paragraph was removed because it contained outdated legal language that is no longer standard in institutional NDAs.")
    del_text.font.strike = True
    del_text.font.color.rgb = RGBColor(255, 0, 0)
    
    # Add MOVED text
    move_para = doc.add_paragraph()
    move_label = move_para.add_run("[MOVED TEXT] ")
    move_label.bold = True
    move_label.font.color.rgb = RGBColor(0, 0, 255)  # Blue
    
    move_text = move_para.add_run("This clause was moved from Section 5 to Section 3 for better organization: The obligations of confidentiality shall survive termination of this Agreement for a period of three (3) years.")
    move_text.font.color.rgb = RGBColor(0, 0, 255)
    move_text.underline = True
    
    # Add FORMATTING changes
    format_para = doc.add_paragraph()
    format_label = format_para.add_run("[FORMATTING CHANGE] ")
    format_label.bold = True
    format_label.font.color.rgb = RGBColor(128, 0, 128)  # Purple
    
    # Show old formatting (strikethrough)
    old_format = format_para.add_run("confidential information")
    old_format.font.strike = True
    old_format.font.color.rgb = RGBColor(255, 0, 0)
    
    format_para.add_run(" → ")
    
    # Show new formatting (underlined)
    new_format = format_para.add_run("Confidential Information")
    new_format.underline = True
    new_format.font.color.rgb = RGBColor(0, 128, 0)
    new_format.bold = True
    
    format_para.add_run(" (capitalized defined term)")
    
    # Add summary of all changes
    doc.add_page_break()
    
    summary_title = doc.add_paragraph()
    title_run = summary_title.add_run("TRACK CHANGES SUMMARY - ALL CHANGE TYPES")
    title_run.bold = True
    title_run.font.size = 16
    
    changes_list = [
        "1. REPLACEMENTS: Entire sentences replaced (red strikethrough → green underline)",
        "2. WORD DELETIONS: Individual words removed (red strikethrough)",
        "3. WORD INSERTIONS: New words added within paragraphs (green underline)",
        "4. PHRASE REPLACEMENTS: Multi-word phrases changed (red → green)",
        "5. PARAGRAPH INSERTIONS: New clauses added (green underline with labels)",
        "6. PARAGRAPH DELETIONS: Entire paragraphs removed (red strikethrough)",
        "7. MOVED TEXT: Content relocated (blue underline)",
        "8. FORMATTING CHANGES: Capitalization/styling (purple labels)"
    ]
    
    for change in changes_list:
        change_para = doc.add_paragraph()
        change_para.add_run(change)
    
    # Save document
    doc.save(output_file)
    
    print(f"✅ Comprehensive test created: {os.path.basename(output_file)}")
    print("📝 Contains ALL track change types:")
    print("   🔴 Red strikethrough = Deletions")
    print("   🟢 Green underline = Insertions") 
    print("   🔵 Blue underline = Moved text")
    print("   🟣 Purple labels = Formatting changes")
    print("🔍 Open in Microsoft Word to test compatibility")
    
    return output_file

if __name__ == "__main__":
    create_comprehensive_track_changes_test()
