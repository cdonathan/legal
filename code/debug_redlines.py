#!/usr/bin/env python3
"""
Debug the redlining process - check AI output and document changes
"""

from docx import Document
import json
import os
import openai

def check_ai_response():
    """Test what AI is actually returning"""
    try:
        with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
            api_key = f.read().strip()
        client = openai.OpenAI(api_key=api_key)
    except:
        print("No OpenAI client available")
        return
    
    sample_text = """This CONFIDENTIALITY AND NONDISCLOSURE AGREEMENT is made between Company A and Company B for the purpose of evaluating a potential transaction."""
    
    prompt = f"""Redline this NDA. Return ONLY JSON with structured changes.

NDA TEXT:
{sample_text}

Return JSON format:
{{
  "changes": [
    {{
      "type": "insert",
      "after_paragraph": 0,
      "text": "The Receiving Party may disclose Confidential Information to its attorneys and advisors.",
      "reason": "Add permitted recipients clause"
    }}
  ]
}}

Add missing: permitted recipients, return of materials."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.1
        )
        
        print("AI Response:")
        print(response.choices[0].message.content)
        
        # Try to extract JSON
        content = response.choices[0].message.content
        json_start = content.find('{')
        json_end = content.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            json_str = content[json_start:json_end]
            changes = json.loads(json_str)
            print(f"\nParsed {len(changes.get('changes', []))} changes:")
            for change in changes.get('changes', []):
                print(f"  - {change['type']}: {change['reason']}")
        
    except Exception as e:
        print(f"AI Error: {e}")

def check_document_changes():
    """Check what changes are actually in the document"""
    doc_path = "/home/cliff/redact/redline_project/REDLINE_NDA_Sample5_pre_redline_redlined.docx"
    doc = Document(doc_path)
    
    print(f"\nDocument Analysis: {os.path.basename(doc_path)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    redline_count = 0
    for i, para in enumerate(doc.paragraphs):
        has_formatting = False
        for run in para.runs:
            if run.underline or (run.font.color.rgb and str(run.font.color.rgb) != 'None'):
                has_formatting = True
                redline_count += 1
                print(f"Paragraph {i}: {para.text[:80]}...")
                break
    
    print(f"Paragraphs with redline formatting: {redline_count}")
    
    if redline_count == 0:
        print("❌ No redline formatting found!")
    else:
        print(f"✅ Found {redline_count} redlined paragraphs")

if __name__ == "__main__":
    print("=== AI Response Check ===")
    check_ai_response()
    
    print("\n=== Document Changes Check ===")
    check_document_changes()
