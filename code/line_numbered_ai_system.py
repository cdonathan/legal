#!/usr/bin/env python3
"""
Line-Numbered AI Analysis System
Adds line numbers to document, AI provides precise location instructions
"""

import openai
import json
import os
import re
from docx import Document
from docx.shared import RGBColor

class LineNumberedAISystem:
    def __init__(self):
        self.personal_info = {}
        self.line_numbered_text = ""
        self.original_paragraphs = []
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def create_line_numbered_document(self, doc_path):
        """Create line-numbered version of document for AI analysis"""
        doc = Document(doc_path)
        self.original_paragraphs = [p.text for p in doc.paragraphs]
        
        line_numbered_lines = []
        line_number = 1
        
        for para_text in self.original_paragraphs:
            if para_text.strip():  # Only number non-empty paragraphs
                # Redact personal info
                redacted_para = self.redact_personal_info(para_text)
                line_numbered_lines.append(f"LINE {line_number:03d}: {redacted_para}")
                line_number += 1
            else:
                line_numbered_lines.append("")  # Keep empty lines
        
        self.line_numbered_text = '\n'.join(line_numbered_lines)
        
        # Save line-numbered version for reference
        with open('/home/cliff/redact/redline_project/Line_Numbered_Sample_2.txt', 'w') as f:
            f.write("# Line-Numbered Document for AI Analysis\n\n")
            f.write(self.line_numbered_text)
        
        print(f"   ✓ Created line-numbered document with {line_number-1} lines")
        return self.line_numbered_text
    
    def get_precise_ai_recommendations(self, line_numbered_text):
        """AI analyzes line-numbered document and provides precise location instructions"""
        if not self.openai_client:
            return []
        
        # Load the problems analysis
        with open('/home/cliff/redact/redline_project/Sample_2_Problems_Analysis.md', 'r') as f:
            problems_analysis = f.read()
        
        # Load Golden NDA for reference
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        prompt = f"""You are an attorney analyzing a line-numbered NDA. Provide PRECISE location instructions for fixes.

GOLDEN NDA STANDARD:
{golden_nda}

PROBLEMS IDENTIFIED:
{problems_analysis}

LINE-NUMBERED NDA:
{line_numbered_text}

TASK: Provide precise change instructions using line numbers. Be SPECIFIC about locations.

Return ONLY JSON array:
[
  {{
    "problem_addressed": "Definition of Confidential Information",
    "change_type": "replace_text_on_line",
    "line_number": 5,
    "find_text": "exact text to find on that line",
    "replace_with": "exact replacement text",
    "reason": "why this fixes the problem",
    "priority": "HIGH"
  }},
  {{
    "problem_addressed": "Missing Purpose Clause",
    "change_type": "insert_after_line",
    "line_number": 3,
    "insert_text": "exact text to insert",
    "reason": "why this addresses the problem",
    "priority": "HIGH"
  }},
  {{
    "problem_addressed": "Add Governing Law",
    "change_type": "insert_at_end",
    "insert_text": "exact text to add at document end",
    "reason": "why this is needed",
    "priority": "MEDIUM"
  }}
]

CRITICAL RULES:
- Use EXACT line numbers from the line-numbered document
- For replacements, specify the EXACT text to find on that specific line
- For insertions, specify the EXACT line number to insert after
- Be precise - the code will follow your instructions exactly"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            
            # Save AI recommendations with line numbers
            with open('/home/cliff/redact/redline_project/Precise_AI_Recommendations_Sample_2.md', 'w') as f:
                f.write("# Precise AI Recommendations with Line Numbers - Sample 2\n\n")
                f.write("**AI Analysis Output:**\n\n")
                f.write("```\n")
                f.write(content)
                f.write("\n```\n\n")
            
            # Extract JSON
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = content[json_start:json_end]
                recommendations = json.loads(json_str)
                
                # Add detailed recommendations to file
                with open('/home/cliff/redact/redline_project/Precise_AI_Recommendations_Sample_2.md', 'a') as f:
                    f.write(f"**Parsed Recommendations:** {len(recommendations)}\n\n")
                    
                    for i, rec in enumerate(recommendations, 1):
                        f.write(f"## Recommendation {i}: {rec.get('problem_addressed', 'Unknown')}\n\n")
                        f.write(f"**Priority:** {rec.get('priority', 'UNKNOWN')}\n\n")
                        f.write(f"**Change Type:** {rec.get('change_type', 'unknown')}\n\n")
                        f.write(f"**Line Number:** {rec.get('line_number', 'N/A')}\n\n")
                        f.write(f"**Find Text:**\n```\n{rec.get('find_text', 'N/A')}\n```\n\n")
                        f.write(f"**Replace/Insert:**\n```\n{rec.get('replace_with', rec.get('insert_text', 'N/A'))}\n```\n\n")
                        f.write(f"**Reason:** {rec.get('reason', 'N/A')}\n\n")
                        f.write("---\n\n")
                
                return recommendations
            else:
                return []
                
        except Exception as e:
            print(f"AI recommendation error: {e}")
            return []
    
    def implement_precise_recommendations(self, original_doc_path, recommendations, output_path_redlined, output_path_clean):
        """Code implements AI's precise line-numbered recommendations"""
        doc = Document(original_doc_path)
        
        print(f"📋 Implementing {len(recommendations)} precise AI recommendations...")
        
        # Create redlined version
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("PRECISE AI RECOMMENDATIONS IMPLEMENTED")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = Problems | 🟢 Green underline = AI fixes with line numbers")
        
        redlined_doc.add_paragraph("=" * 80)
        
        # Process each paragraph with line-based changes
        implemented_count = 0
        
        for line_num, para in enumerate(doc.paragraphs, 1):
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply line-specific replacements
            for rec in recommendations:
                if (rec.get('change_type') == 'replace_text_on_line' and 
                    rec.get('line_number') == line_num and 
                    rec.get('find_text') in para_text):
                    
                    parts = para_text.split(rec['find_text'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Show problem being fixed
                    del_run = new_para.add_run(rec['find_text'])
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Show AI fix
                    ins_run = new_para.add_run(rec['replace_with'])
                    ins_run.underline = True
                    ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    implemented_count += 1
                    break
            
            if para_text:
                new_para.add_run(para_text)
            
            # Add line-specific insertions
            for rec in recommendations:
                if (rec.get('change_type') == 'insert_after_line' and 
                    rec.get('line_number') == line_num):
                    
                    insert_para = redlined_doc.add_paragraph()
                    
                    label_run = insert_para.add_run(f"[AI LINE {line_num} INSERT: {rec['problem_addressed']}] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(0, 100, 0)
                    
                    insert_run = insert_para.add_run(rec['insert_text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    implemented_count += 1
        
        # Add end insertions
        for rec in recommendations:
            if rec.get('change_type') == 'insert_at_end':
                insert_para = redlined_doc.add_paragraph()
                
                label_run = insert_para.add_run(f"[AI END INSERT: {rec['problem_addressed']}] ")
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(0, 100, 0)
                
                insert_run = insert_para.add_run(rec['insert_text'])
                insert_run.underline = True
                insert_run.font.color.rgb = RGBColor(0, 128, 0)
                
                implemented_count += 1
        
        # Create clean version (same logic but without markup)
        clean_doc = Document()
        
        for line_num, para in enumerate(doc.paragraphs, 1):
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for rec in recommendations:
                if (rec.get('change_type') == 'replace_text_on_line' and 
                    rec.get('line_number') == line_num and 
                    rec.get('find_text') in para_text):
                    para_text = para_text.replace(rec['find_text'], rec['replace_with'])
            
            new_para.add_run(para_text)
            
            # Add insertions
            for rec in recommendations:
                if (rec.get('change_type') == 'insert_after_line' and 
                    rec.get('line_number') == line_num):
                    insert_para = clean_doc.add_paragraph()
                    insert_para.add_run(rec['insert_text'])
        
        # Add end insertions to clean version
        for rec in recommendations:
            if rec.get('change_type') == 'insert_at_end':
                insert_para = clean_doc.add_paragraph()
                insert_para.add_run(rec['insert_text'])
        
        # Save both versions
        redlined_doc.save(output_path_redlined)
        clean_doc.save(output_path_clean)
        
        print(f"   ✓ Successfully implemented {implemented_count} precise AI recommendations")
        return implemented_count
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Complete line-numbered AI analysis + precise code implementation"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Step 1: Create line-numbered document
        print("[█░░░░] Step 1: Creating line-numbered document...")
        line_numbered_text = self.create_line_numbered_document(input_path)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: AI Analysis with line numbers
        print("[██░░░] Step 2: Getting precise AI recommendations...")
        recommendations = self.get_precise_ai_recommendations(line_numbered_text)
        print(f"   ✓ AI provided {len(recommendations)} precise recommendations")
        
        # Step 3: Precise Code Implementation
        print("[███░░] Step 3: Implementing precise recommendations...")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_precise_redlined.docx"
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_precise_clean-version.docx"
        
        implemented = self.implement_precise_recommendations(input_path, recommendations, redlined_path, clean_path)
        
        # Step 4: Restore personal info
        print("[████░] Step 4: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("[█████] Step 5: Complete!")
        print("✅ PRECISE LINE-NUMBERED AI SYSTEM COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"📋 Line-numbered doc: Line_Numbered_Sample_2.txt")
        print(f"📋 AI Recommendations: Precise_AI_Recommendations_Sample_2.md")
        print(f"🎯 AI recommended {len(recommendations)} changes, code implemented {implemented}")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 line_numbered_ai_system.py input.docx")
        sys.exit(1)
    
    processor = LineNumberedAISystem()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
