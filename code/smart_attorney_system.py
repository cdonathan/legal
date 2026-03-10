#!/usr/bin/env python3
"""
Smart Attorney Pattern-Based NDA Redlining System
Uses enhanced patterns library instead of golden NDA templates
"""

import os
import sys
import subprocess
import json
import secrets
import re
import re
import openai
import time

class SmartAttorneySystem:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
        self.patterns_prompt = self._load_patterns_prompt()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def _load_patterns_prompt(self):
        """Load the smart attorney prompt"""
        try:
            with open('/home/cliff/redact/redline_project/components/smart_attorney_prompt.md', 'r') as f:
                return f.read()
        except:
            return None
    
    def convert_to_docx(self, input_path):
        """Convert any document type to DOCX format first"""
        try:
            file_ext = os.path.splitext(input_path)[1].lower()
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            
            # If already DOCX, return as-is
            if file_ext == '.docx':
                return input_path
            
            # Output path for converted DOCX
            output_dir = "/tmp"
            docx_path = os.path.join(output_dir, f"{base_name}_converted.docx")
            
            # For PDFs, use specialized PDF conversion
            if file_ext == '.pdf':
                return self.convert_pdf_to_docx(input_path, docx_path)
            
            # For other formats, try LibreOffice conversion to DOCX
            cmd = [
                'libreoffice', 
                '--headless', 
                '--convert-to', 'docx',
                '--outdir', output_dir,
                input_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0 and os.path.exists(docx_path):
                print(f"   ✓ Converted {file_ext} to DOCX")
                return docx_path
            
            print(f"   ❌ Could not convert {file_ext} to DOCX")
            return None
            
        except Exception as e:
            print(f"   ❌ Conversion failed: {e}")
            return None

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX using multiple methods"""
        
        # Method 1: Try pdf2docx (most accurate)
        try:
            from pdf2docx import Converter
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            
            if os.path.exists(docx_path):
                print("   ✓ Converted PDF to DOCX using pdf2docx")
                return docx_path
        except ImportError:
            print("   ⚠️ pdf2docx not available, trying alternative methods...")
        except Exception as e:
            print(f"   ⚠️ pdf2docx failed: {e}")
        
        # Method 2: Try PyMuPDF (fitz) + python-docx
        try:
            import fitz  # PyMuPDF
            from docx import Document
            
            doc = fitz.open(pdf_path)
            word_doc = Document()
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    word_doc.add_paragraph(text)
                    if page_num < len(doc) - 1:  # Add page break except for last page
                        word_doc.add_page_break()
            
            word_doc.save(docx_path)
            doc.close()
            
            if os.path.exists(docx_path):
                print("   ✓ Converted PDF to DOCX using PyMuPDF")
                return docx_path
                
        except ImportError:
            print("   ⚠️ PyMuPDF not available, trying next method...")
        except Exception as e:
            print(f"   ⚠️ PyMuPDF failed: {e}")
        
        # Method 3: Try pdfplumber + python-docx
        try:
            import pdfplumber
            from docx import Document
            
            word_doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        word_doc.add_paragraph(text)
                        word_doc.add_page_break()
            
            word_doc.save(docx_path)
            
            if os.path.exists(docx_path):
                print("   ✓ Converted PDF to DOCX using pdfplumber")
                return docx_path
                
        except ImportError:
            print("   ⚠️ pdfplumber not available, trying basic method...")
        except Exception as e:
            print(f"   ⚠️ pdfplumber failed: {e}")
        
        # Method 4: Basic PyPDF2 fallback
        try:
            import PyPDF2
            from docx import Document
            
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                word_doc = Document()
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        word_doc.add_paragraph(text)
                        word_doc.add_page_break()
                
                word_doc.save(docx_path)
            
            if os.path.exists(docx_path):
                print("   ✓ Converted PDF to DOCX using PyPDF2")
                return docx_path
                
        except ImportError:
            print("   ❌ No PDF conversion libraries available")
        except Exception as e:
            print(f"   ❌ PyPDF2 failed: {e}")
        
        print("   ❌ All PDF conversion methods failed")
        return None

    def convert_with_libreoffice(self, input_path):
        """Convert document using LibreOffice"""
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            input_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
        
        return None
    
    def redact_personal_info(self, text):
        """Redact personal information and add line numbers"""
        redacted_text = text
        counter = len(self.personal_info) + 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'broker', 'seller']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                if placeholder not in self.personal_info:
                    self.personal_info[placeholder] = match.group()
                    redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                    counter += 1
        
        # Add line numbers after redaction
        lines = redacted_text.split('\n')
        numbered_lines = []
        for i, line in enumerate(lines, 1):
            numbered_lines.append(f"LINE {i:03d}: {line}")
        
        return '\n'.join(numbered_lines)
    
    def smart_attorney_analysis(self, redacted_text, base_name):
        """Smart attorney pattern analysis with timing"""
        if not self.openai_client or not self.patterns_prompt:
            return None, []
        
        # Extract the implementation prompt from the patterns prompt
        implementation_section = self.patterns_prompt.split("## 🎯 IMPLEMENTATION PROMPT FOR SPECIFIC NDA")[1]
        
        prompt = f"""You are an experienced real estate attorney who redlines NDAs to protect purchaser interests. You follow systematic patterns based on years of practice.

PATTERN DETECTION FRAMEWORK:

UNIVERSAL PATTERNS (Check Every NDA):
- Pattern 1: Confidential Information Exclusions - TRIGGER: No exclusions for publicly available, already possessed, or independently developed information
- Pattern 2: Expanded Disclosure Recipients - TRIGGER: Limited disclosure list (only employees, legal counsel)  
- Pattern 3: Effective Date Definition - TRIGGER: No "Effective Date" defined

CONDITIONAL PATTERNS (Apply When Triggered):
- Pattern 4: Return/Destruction Flexibility - TRIGGER: Rigid "must return" without destroy option
- Pattern 5: Sophistication Assessment - TRIGGER: Simple vs sophisticated NDA analysis
- Pattern 6: Commercial Reasonableness - TRIGGER: "Take all steps" or absolute obligations
- Pattern 7: Legal Compliance Exceptions - TRIGGER: No court order/legal requirement exceptions
- Pattern 8: Fee Protection - TRIGGER: "Attorney's fees" without reasonableness qualifier
- Pattern 9: Injunctive Relief Balance - TRIGGER: Missing injunctive relief or requires bond
- Pattern 10: Term Limitation - TRIGGER: Perpetual or indefinite obligations
- Pattern 11: Defined Term Consistency - TRIGGER: Mixed capitalization of defined terms
- Pattern 12: Execution Flexibility - TRIGGER: Original signatures only
- Pattern 13: Disclosure Recipients Expansion - TRIGGER: Limited recipient list (only employees, officers, attorneys)
- Pattern 14: Business Purpose Expansion - TRIGGER: Generic "participation, financing" without "purchasing"
- Pattern 15: Defined Terms Addition - TRIGGER: Undefined key terms (property references without definition)
- Pattern 16: Confidentiality Requirement Softening - TRIGGER: "bound by written confidentiality agreements"

TASK: Analyze this LINE-NUMBERED NDA and provide EXACT line numbers and text for changes. 

SYSTEMATIC EVALUATION REQUIRED: You MUST evaluate every single pattern 1-16 and state FOUND or NOT FOUND for each.

EVALUATION PROCESS:
1. Check Pattern 1: Confidential Information Exclusions - FOUND or NOT FOUND
2. Check Pattern 2: Expanded Disclosure Recipients - FOUND or NOT FOUND  
3. Check Pattern 3: Effective Date Definition - FOUND or NOT FOUND
4. Check Pattern 4: Return/Destruction Flexibility - FOUND or NOT FOUND
5. Check Pattern 5: Sophistication Assessment - FOUND or NOT FOUND
6. Check Pattern 6: Commercial Reasonableness - FOUND or NOT FOUND
7. Check Pattern 7: Legal Compliance Exceptions - FOUND or NOT FOUND
8. Check Pattern 8: Fee Protection - FOUND or NOT FOUND
9. Check Pattern 9: Injunctive Relief Balance - FOUND or NOT FOUND
10. Check Pattern 10: Term Limitation - FOUND or NOT FOUND
11. Check Pattern 11: Defined Term Consistency - FOUND or NOT FOUND
12. Check Pattern 12: Execution Flexibility - FOUND or NOT FOUND
13. Check Pattern 13: Disclosure Recipients Expansion - FOUND or NOT FOUND
14. Check Pattern 14: Business Purpose Expansion - FOUND or NOT FOUND
15. Check Pattern 15: Defined Terms Addition - FOUND or NOT FOUND
16. Check Pattern 16: Confidentiality Requirement Softening - FOUND or NOT FOUND

For each FOUND pattern, provide implementation instructions. For NOT FOUND patterns, you may skip implementation.

For each triggered pattern, specify:
1. LINE NUMBER where change occurs
2. EXACT current text on that line (copy exactly)
3. EXACT replacement text for that line
4. OR specify "insert_after" with line number and text to insert

Return your analysis in this format:

SYSTEMATIC EVALUATION:
Pattern 1: FOUND/NOT FOUND - [brief reason]
Pattern 2: FOUND/NOT FOUND - [brief reason]
Pattern 3: FOUND/NOT FOUND - [brief reason]
Pattern 4: FOUND/NOT FOUND - [brief reason]
Pattern 5: FOUND/NOT FOUND - [brief reason]
Pattern 6: FOUND/NOT FOUND - [brief reason]
Pattern 7: FOUND/NOT FOUND - [brief reason]
Pattern 8: FOUND/NOT FOUND - [brief reason]
Pattern 9: FOUND/NOT FOUND - [brief reason]
Pattern 10: FOUND/NOT FOUND - [brief reason]
Pattern 11: FOUND/NOT FOUND - [brief reason]
Pattern 12: FOUND/NOT FOUND - [brief reason]
Pattern 13: FOUND/NOT FOUND - [brief reason]
Pattern 14: FOUND/NOT FOUND - [brief reason]
Pattern 15: FOUND/NOT FOUND - [brief reason]
Pattern 16: FOUND/NOT FOUND - [brief reason]

IMPLEMENTATION INSTRUCTIONS:
[JSON array for FOUND patterns only]
[
  {{
    "pattern_name": "Pattern 8: Fee Protection",
    "line_number": 45,
    "action": "replace",
    "current_text": "including attorney's fees, arising out of any breach",
    "new_text": "including reasonable attorney's fees, arising out of any breach",
    "attorney_reasoning": "Purchaser should not be liable for unreasonable attorney's fees",
    "purchaser_benefit": "Limits financial exposure for legal costs"
  }},
  {{
    "pattern_name": "Pattern 1: Confidential Information Exclusions",
    "line_number": 12,
    "action": "insert_after",
    "insert_text": "However, Confidential Information does not include: (i) information already in possession; (ii) information publicly available; (iii) information independently developed; (iv) information received from third parties without confidentiality obligations.",
    "attorney_reasoning": "Purchaser needs protection from overly broad confidentiality scope",
    "purchaser_benefit": "Limits confidentiality obligations to truly confidential information"
  }}
]

LINE-NUMBERED NDA:
{redacted_text}

CRITICAL: Use exact LINE numbers and exact text from the document above. Copy/paste exactly from the LINE XXX: content."""

        try:
            # Time the OpenAI API call
            import time
            import re
            start_time = time.time()
            print(f"   🔄 Starting OpenAI API call at {time.strftime('%H:%M:%S')}")
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.1
            )
            
            end_time = time.time()
            api_duration = end_time - start_time
            print(f"   ✅ OpenAI API call completed at {time.strftime('%H:%M:%S')}")
            print(f"   ⏱️ API call duration: {api_duration:.2f} seconds")
            
            analysis_content = response.choices[0].message.content
            
            # Save analysis
            analysis_file = f"/home/cliff/redact/redline_project/{base_name}_Smart_Attorney_Analysis.md"
            with open(analysis_file, 'w') as f:
                f.write(f"# Smart Attorney Pattern Analysis: {base_name}\n\n")
                f.write("**Pattern-Based Attorney Analysis (No Templates)**\n\n")
                f.write(f"**OpenAI API Duration: {api_duration:.2f} seconds**\n\n")
                f.write("```json\n")
                f.write(analysis_content)
                f.write("\n```\n\n")
            
            # Parse JSON
            # Handle duplicate ```json markers and fix line number formatting
            analysis_content = analysis_content.replace('```json\n```json', '```json')
            # Fix leading zeros in line numbers that break JSON parsing
            analysis_content = re.sub(r'"line_number":\s*0*(\d+)', r'"line_number": \1', analysis_content)
            # Clean control characters (tabs, etc.) that break JSON parsing
            analysis_content = re.sub(r'[\t\r\n\f\v]', ' ', analysis_content)
            
            json_start = analysis_content.find('[')
            json_end = analysis_content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = analysis_content[json_start:json_end]
                instructions = json.loads(json_str)
                
                print(f"   ✓ Smart attorney analysis: {os.path.basename(analysis_file)}")
                print(f"   ✓ Triggered {len(instructions)} attorney patterns")
                return analysis_file, instructions
            else:
                print("   ❌ Could not parse attorney analysis JSON")
                return analysis_file, []
                
        except Exception as e:
            print(f"   ❌ Smart attorney analysis error: {e}")
            return None, []
    
    def find_text_flexibly(self, doc, target_text):
        """Find text using flexible matching to handle redacted text and variations"""
        import re
        
        if not target_text:
            return None
        
        # Clean target text by removing redaction placeholders
        clean_target = re.sub(r'\[NAME_\d+\]', '', target_text)
        clean_target = re.sub(r'\[COMPANY_\d+\]', '', target_text)
        clean_target = re.sub(r'\[ADDRESS_\d+\]', '', target_text)
        clean_target = re.sub(r'\[DATE_\d+\]', '', target_text)
        clean_target = clean_target.strip()
        
        # Try exact match first
        search = doc.createSearchDescriptor()
        search.setSearchString(clean_target)
        found = doc.findFirst(search)
        if found:
            return found
        
        # Extract key phrases (remove common words, keep meaningful terms)
        words = clean_target.split()
        key_phrases = []
        
        # Try progressively shorter phrases
        for length in range(min(len(words), 6), 2, -1):
            for i in range(len(words) - length + 1):
                phrase = ' '.join(words[i:i+length])
                if len(phrase) > 10:  # Only meaningful phrases
                    key_phrases.append(phrase)
        
        # Try each key phrase
        for phrase in key_phrases:
            search.setSearchString(phrase)
            found = doc.findFirst(search)
            if found:
                return found
        
        # Try individual significant words (longer than 4 chars)
        significant_words = [w for w in words if len(w) > 4 and w.lower() not in 
                           ['line', 'that', 'this', 'with', 'from', 'they', 'have', 'will', 'been']]
        
        for word in significant_words:
            search.setSearchString(word)
            found = doc.findFirst(search)
            if found:
                return found
        
        return None

    def create_smart_redlined_document(self, instructions, original_path, base_name):
        """Create redlined document using line-based precise changes"""
        import uno
        from com.sun.star.beans import PropertyValue
        
        print("   🔄 Starting LibreOffice for line-based attorney redlining...")
        os.system("pkill -f libreoffice")
        time.sleep(2)
        
        os.system("libreoffice --headless --invisible --accept='socket,host=localhost,port=2002;urp;StarOffice.ServiceManager' &")
        time.sleep(5)
        
        try:
            # Connect to LibreOffice
            local_context = uno.getComponentContext()
            resolver = local_context.ServiceManager.createInstanceWithContext(
                "com.sun.star.bridge.UnoUrlResolver", local_context)
            
            context = resolver.resolve("uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")
            desktop = context.ServiceManager.createInstanceWithContext(
                "com.sun.star.frame.Desktop", context)
            
            def create_property(name, value):
                prop = PropertyValue()
                prop.Name = name
                prop.Value = value
                return prop
            
            # Load document with auto-detection of file type
            original_url = uno.systemPathToFileUrl(os.path.abspath(original_path))
            load_props = (
                create_property("Hidden", True),
                create_property("UpdateDocMode", 0)  # Don't update links/fields that might cause formatting issues
            )
            doc = desktop.loadComponentFromURL(original_url, "_blank", 0, load_props)
            
            # Fix specific page break: "protective order" + line break + "or other remedy"
            try:
                text_obj = doc.getText()
                full_text = text_obj.getString()
                
                # Look for the specific pattern and fix it
                import re
                # Pattern: "protective order" followed by whitespace/newlines followed by "or other remedy"
                pattern = r'(protective order)\s*[\r\n]+\s*(or other remedy)'
                
                if re.search(pattern, full_text):
                    fixed_text = re.sub(pattern, r'\1 \2', full_text)
                    
                    # Replace the entire document text
                    cursor = text_obj.createTextCursor()
                    cursor.gotoStart(False)
                    cursor.gotoEnd(True)
                    cursor.setString(fixed_text)
                    print("   ✓ Fixed 'protective order' -> 'or other remedy' page break")
                else:
                    print("   ✓ No 'protective order' page break found")
                    
            except Exception as e:
                print(f"   ⚠️ Page break fix failed: {e}")
                pass
            
            # Enable track changes BEFORE making changes
            doc.RecordChanges = True
            print(f"   ✓ Track changes enabled: {doc.RecordChanges}")
            
            # Check initial redlines count
            try:
                initial_redlines = doc.getRedlines().getCount()
                print(f"   📊 Initial redlines: {initial_redlines}")
            except:
                initial_redlines = 0
            
            implemented_count = 0
            
            # Apply line-based changes from AI instructions
            for inst in instructions:
                try:
                    pattern_name = inst.get('pattern_name', '')
                    action = inst.get('action', '')
                    line_number = inst.get('line_number', 0)
                    
                    print(f"   🔄 Applying: {pattern_name} (Line {line_number})")
                    
                    if action == 'replace':
                        current_text = inst.get('current_text', '')
                        new_text = inst.get('new_text', '')
                        
                        if current_text and new_text:
                            # Use flexible text matching
                            found = self.find_text_flexibly(doc, current_text)
                            
                            if found:
                                text = doc.getText()
                                cursor = text.createTextCursor()
                                cursor.gotoRange(found, False)
                                cursor.gotoRange(found.getEnd(), True)
                                # Clear selection completely, then insert new text
                                selected_text = cursor.getString()
                                cursor.setString("")  # Clear selection
                                text.insertString(cursor, new_text, False)
                                implemented_count += 1
                                print(f"   ✓ Line {line_number}: Replaced '{selected_text[:30]}...' with track changes")
                            else:
                                print(f"   ❌ Line {line_number}: Text not found with flexible matching - '{current_text[:50]}...'")
                        else:
                            print(f"   ❌ Line {line_number}: Missing current_text or new_text")
                    
                    elif action == 'insert_after':
                        insert_text = inst.get('insert_text', '')
                        
                        if insert_text:
                            # For insert_after, we need to find a reference point
                            # We'll use the pattern name to determine where to insert
                            if 'exclusion' in pattern_name.lower():
                                # Try to find common confidential info patterns
                                search_terms = [
                                    "financial information.",
                                    "Confidential Information",
                                    "confidential information"
                                ]
                                
                                inserted = False
                                for search_term in search_terms:
                                    search = doc.createSearchDescriptor()
                                    search.setSearchString(search_term)
                                    found = doc.findFirst(search)
                                    
                                    if found:
                                        text = doc.getText()
                                        cursor = text.createTextCursor()
                                        cursor.gotoRange(found.getEnd(), False)
                                        text.insertString(cursor, f" {insert_text}", False)
                                        implemented_count += 1
                                        print(f"   ✓ Line {line_number}: Inserted exclusions after '{search_term}'")
                                        inserted = True
                                        break
                                
                                if not inserted:
                                    print(f"   ❌ Line {line_number}: Could not find insertion point for exclusions")
                            
                            else:
                                # Handle new patterns for insert_after
                                if 'recipients expansion' in pattern_name.lower():
                                    # Find disclosure recipients to expand
                                    search_terms = [
                                        "employees, officers, directors, shareholders, attorneys, accountants",
                                        "legal counsel, real estate broker, real estate agent"
                                    ]
                                    
                                    inserted = False
                                    for search_term in search_terms:
                                        found = self.find_text_flexibly(doc, search_term)
                                        if found:
                                            text = doc.getText()
                                            cursor = text.createTextCursor()
                                            cursor.gotoRange(found, False)
                                            cursor.gotoRange(found.getEnd(), True)
                                            
                                            if "legal counsel" in search_term:
                                                expanded = "legal counsel, real estate broker, real estate agent, institutional lenders, investors, members, managers, officers, and financial advisors"
                                            else:
                                                expanded = "affiliates, members, managers, partners, directors, investors, employees, officers, directors, shareholders, attorneys, accountants, and financial advisors"
                                            
                                            text.insertString(cursor, expanded, False)
                                            implemented_count += 1
                                            print(f"   ✓ Line {line_number}: Expanded disclosure recipients")
                                            inserted = True
                                            break
                                    
                                    if not inserted:
                                        print(f"   ❌ Line {line_number}: Could not expand disclosure recipients")
                                
                                elif 'business purpose' in pattern_name.lower():
                                    # Add purchasing to business purpose
                                    found = self.find_text_flexibly(doc, "participation, financing")
                                    if found:
                                        text = doc.getText()
                                        cursor = text.createTextCursor()
                                        cursor.gotoRange(found, False)
                                        cursor.gotoRange(found.getEnd(), True)
                                        text.insertString(cursor, "participation, financing, purchasing", False)
                                        implemented_count += 1
                                        print(f"   ✓ Line {line_number}: Added purchasing to business purpose")
                                    else:
                                        print(f"   ❌ Line {line_number}: Could not find business purpose")
                                
                                elif 'defined terms' in pattern_name.lower():
                                    # Add Property definition
                                    found = self.find_text_flexibly(doc, "premises located at")
                                    if not found:
                                        found = self.find_text_flexibly(doc, "property located at")
                                    
                                    if found:
                                        text = doc.getText()
                                        cursor = text.createTextCursor()
                                        cursor.gotoRange(found.getEnd(), False)
                                        # Don't move cursor arbitrarily - stay at found position
                                        text.insertString(cursor, ' ("Property")', False)
                                        implemented_count += 1
                                        print(f"   ✓ Line {line_number}: Added Property definition")
                                    else:
                                        print(f"   ❌ Line {line_number}: Could not find property reference")
                                
                                elif 'confidentiality requirement' in pattern_name.lower():
                                    # Soften confidentiality requirements
                                    found = self.find_text_flexibly(doc, "bound by written confidentiality agreements")
                                    if found:
                                        text = doc.getText()
                                        cursor = text.createTextCursor()
                                        cursor.gotoRange(found, False)
                                        cursor.gotoRange(found.getEnd(), True)
                                        text.insertString(cursor, "informed about the confidentiality nature of the Confidential Information", False)
                                        implemented_count += 1
                                        print(f"   ✓ Line {line_number}: Softened confidentiality requirements")
                                    else:
                                        print(f"   ❌ Line {line_number}: Could not find confidentiality requirements")
                                
                                else:
                                    # Generic insertion at end of document
                                    text = doc.getText()
                                    cursor = text.createTextCursor()
                                    cursor.gotoEnd(False)
                                    # Insert with minimal spacing to avoid page breaks
                                    text.insertString(cursor, f" {insert_text}", False)
                                    implemented_count += 1
                                    print(f"   ✓ Line {line_number}: Inserted at end of document")
                        else:
                            print(f"   ❌ Line {line_number}: Missing insert_text")
                    
                    else:
                        print(f"   ❌ Line {line_number}: Unknown action '{action}'")
                
                except Exception as e:
                    print(f"   ❌ Failed to apply {pattern_name}: {e}")
            
            # Check final redlines count
            try:
                final_redlines = doc.getRedlines().getCount()
                print(f"   📊 Final redlines: {final_redlines} (added {final_redlines - initial_redlines})")
            except:
                print("   📊 Could not count final redlines")
            
            # Save redlined document (with track changes)
            redlined_path = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Smart_Attorney_Redlined.docx"
            redlined_url = uno.systemPathToFileUrl(os.path.abspath(redlined_path))
            
            save_props = (
                create_property("FilterName", "MS Word 2007 XML"),
                create_property("Overwrite", True)
            )
            doc.storeAsURL(redlined_url, save_props)
            
            # Create clean version by opening original again and applying changes without track changes
            load_props = (
                create_property("Hidden", True),
                create_property("UpdateDocMode", 0)
            )
            clean_doc = desktop.loadComponentFromURL(original_url, "_blank", 0, load_props)
            clean_doc.RecordChanges = False  # Disable track changes for clean version
            
            # Apply same changes to clean document without track changes
            implemented_clean = 0
            for inst in instructions:
                try:
                    action = inst.get('action', '')
                    if action == 'replace':
                        current_text = inst.get('current_text', '')
                        new_text = inst.get('new_text', '')
                        if current_text and new_text:
                            found = self.find_text_flexibly(clean_doc, current_text)
                            if found:
                                text = clean_doc.getText()
                                cursor = text.createTextCursor()
                                cursor.gotoRange(found, False)
                                cursor.gotoRange(found.getEnd(), True)
                                cursor.setString("")
                                text.insertString(cursor, new_text, False)
                                implemented_clean += 1
                    elif action == 'insert_after':
                        insert_text = inst.get('insert_text', '')
                        if insert_text:
                            # Apply insertion logic for clean version
                            if 'exclusions' in inst.get('pattern_name', '').lower():
                                search_terms = ['Confidential Information', 'financial information.']
                                for search_term in search_terms:
                                    found = self.find_text_flexibly(clean_doc, search_term)
                                    if found:
                                        text = clean_doc.getText()
                                        cursor = text.createTextCursor()
                                        cursor.gotoRange(found.getEnd(), False)
                                        text.insertString(cursor, f" {insert_text}", False)
                                        implemented_clean += 1
                                        break
                            else:
                                # Generic insertion at end
                                text = clean_doc.getText()
                                cursor = text.createTextCursor()
                                cursor.gotoEnd(False)
                                text.insertString(cursor, f" {insert_text}", False)
                                implemented_clean += 1
                except:
                    continue
            
            # Save clean document
            clean_path = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Smart_Attorney_Clean.docx"
            clean_url = uno.systemPathToFileUrl(os.path.abspath(clean_path))
            clean_doc.storeAsURL(clean_url, save_props)
            clean_doc.close(True)
            doc.close(True)
            
            print(f"   ✅ Created redlined document: {os.path.basename(redlined_path)}")
            print(f"   ✅ Created clean document: {os.path.basename(clean_path)} ({implemented_clean} changes applied)")
            print(f"   ✅ Applied {implemented_count}/{len(instructions)} attorney patterns with precise line targeting")
            
            return redlined_path, clean_path
            
        except Exception as e:
            print(f"   ❌ LibreOffice error: {e}")
            return None
        
        finally:
            try:
                desktop.terminate()
            except:
                pass
            os.system("pkill -f libreoffice")
    
    def restore_personal_info_libreoffice(self, doc_path):
        """Restore personal information using LibreOffice"""
        if not doc_path or not os.path.exists(doc_path):
            return
        
        import uno
        from com.sun.star.beans import PropertyValue
        
        os.system("libreoffice --headless --invisible --accept='socket,host=localhost,port=2003;urp;StarOffice.ServiceManager' &")
        time.sleep(3)
        
        try:
            local_context = uno.getComponentContext()
            resolver = local_context.ServiceManager.createInstanceWithContext(
                "com.sun.star.bridge.UnoUrlResolver", local_context)
            
            context = resolver.resolve("uno:socket,host=localhost,port=2003;urp;StarOffice.ComponentContext")
            desktop = context.ServiceManager.createInstanceWithContext(
                "com.sun.star.frame.Desktop", context)
            
            def create_property(name, value):
                prop = PropertyValue()
                prop.Name = name
                prop.Value = value
                return prop
            
            doc_url = uno.systemPathToFileUrl(os.path.abspath(doc_path))
            doc = desktop.loadComponentFromURL(doc_url, "_blank", 0, (create_property("Hidden", True),))
            
            # Restore personal info
            for placeholder, original in self.personal_info.items():
                replace_desc = doc.createReplaceDescriptor()
                replace_desc.setSearchString(placeholder)
                replace_desc.setReplaceString(original)
                doc.replaceAll(replace_desc)
            
            save_props = (
                create_property("FilterName", "MS Word 2007 XML"),
                create_property("Overwrite", True)
            )
            doc.storeAsURL(doc_url, save_props)
            doc.close(True)
            
        except Exception as e:
            print(f"   ❌ Personal info restore error: {e}")
        
        finally:
            try:
                desktop.terminate()
            except:
                pass
            os.system("pkill -f libreoffice")
    
    def redact_docx_pii(self, docx_path, base_name):
        """Redact PII from DOCX using hex mapping system"""
        try:
            # Convert DOCX to text for redaction
            text_content = self.convert_with_libreoffice(docx_path)
            if not text_content:
                return None, None
            
            # Apply hex mapping redaction
            redacted_text, hex_mapping = self.apply_hex_redaction(text_content)
            redacted_count = len(hex_mapping)
            
            # Create new DOCX with redacted content
            from docx import Document
            new_doc = Document()
            
            # Split redacted text into lines and add as paragraphs
            for line in redacted_text.split('\n'):
                if line.strip():
                    new_doc.add_paragraph(line.strip())
            
            # Save redacted DOCX and mapping file
            redacted_path = f"/tmp/{base_name}_redacted.docx"
            mapping_path = f"/tmp/{base_name}_mapping.json"
            
            new_doc.save(redacted_path)
            
            with open(mapping_path, 'w') as f:
                json.dump(hex_mapping, f, indent=2)
            
            print(f"   ✓ Redacted {redacted_count} PII items with hex mapping")
            return redacted_path, mapping_path
            
        except Exception as e:
            print(f"   ❌ Redaction failed: {e}")
            return None, None
    
    def apply_hex_redaction(self, text):
        """Apply hex mapping redaction patterns"""
        patterns = [
            # Email addresses (match first to avoid conflicts)
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL'),
            
            # Phone numbers
            (r'\(\d{3}\)\s*\d{3}-\d{4}', 'PHONE'),
            (r'\d{3}-\d{3}-\d{4}', 'PHONE'),
            
            # Full addresses (number + street name) - match before street alone
            (r'\b\d+\s+[A-Z][a-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Court|Ct|Boulevard|Blvd)(?:\s*,\s*[A-Z][a-z]+)?(?:\s*,\s*[A-Z]{2})?\s*\d{5}?\b', 'ADDRESS'),
            
            # Street names (just the street part) - match after full addresses
            (r'\b[A-Z][a-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Court|Ct|Boulevard|Blvd)\b', 'STREET'),
            
            # ZIP codes
            (r'\b\d{5}(?:-\d{4})?\b', 'ZIP'),
            
            # Specific company names (with business suffixes)
            (r'\b[A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|Corporation|Company|Co\.)\b', 'COMPANY'),
            
            # Dates in signature lines
            (r'\b\d{1,2}/\d{1,2}/\d{4}\b', 'DATE'),
            (r'\b\d{1,2}-\d{1,2}-\d{4}\b', 'DATE'),
            
            # Personal names (avoid common legal terms) - match last to avoid conflicts
            (r'\b(?!Real|Estate|This|Agreement|Property|Information|Party|Buyer|Seller|Company)[A-Z][a-z]+ [A-Z][a-z]+\b', 'PERSON'),
        ]
        
        redacted_text = text
        hex_mapping = {}
        
        for pattern, label in patterns:
            matches = list(re.finditer(pattern, redacted_text))
            for match in matches:
                original = match.group()
                hex_id = secrets.token_hex(8)
                placeholder = f"[{label}:{hex_id}]"
                
                # Store mapping
                hex_mapping[hex_id] = {
                    'type': label,
                    'original': original,
                    'placeholder': placeholder
                }
                
                redacted_text = redacted_text.replace(original, placeholder, 1)
        
        return redacted_text, hex_mapping

    def restore_pii_in_docx(self, redacted_docx_path, mapping_path, output_path):
        """Restore PII in DOCX using hex mapping"""
        try:
            # Load mapping
            with open(mapping_path, 'r') as f:
                hex_mapping = json.load(f)
            
            # Convert redacted DOCX to text
            redacted_text = self.convert_with_libreoffice(redacted_docx_path)
            if not redacted_text:
                return False
            
            # Restore PII
            restored_text = redacted_text
            for hex_id, data in hex_mapping.items():
                restored_text = restored_text.replace(data['placeholder'], data['original'])
            
            # Create restored DOCX
            from docx import Document
            doc = Document()
            
            for line in restored_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"   ❌ PII restoration failed: {e}")
            return False

    def process_smart_attorney_system(self, input_path):
        """Complete smart attorney pattern-based system"""
        print(f"🔄 Smart Attorney Pattern System: {os.path.basename(input_path)}")
        
        # Step 0: Convert to DOCX if needed
        print("[░░░░] Step 0: Converting to DOCX format...")
        docx_path = self.convert_to_docx(input_path)
        if not docx_path:
            print("❌ Could not convert to DOCX format")
            return
        
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        
        # Step 1: Redact PII from DOCX before processing
        print("[█░░░] Step 1: Redacting personal information...")
        redacted_docx_path, mapping_path = self.redact_docx_pii(docx_path, base_name)
        if not redacted_docx_path:
            print("❌ Document redaction failed")
            return
        
        # Step 2: Convert redacted document to text for analysis
        print("[██░░] Step 2: Converting redacted document...")
        text_content = self.convert_with_libreoffice(redacted_docx_path)
        if not text_content:
            print("❌ Document conversion failed")
            return
        
        # Step 3: Smart attorney pattern analysis
        print("[██░░] Step 2: Smart attorney pattern analysis...")
        analysis_file, instructions = self.smart_attorney_analysis(text_content, base_name)
        
        if not instructions:
            print("❌ No attorney patterns triggered")
            return
        
        # Step 4: Create redlined document
        print("[███░] Step 3: Creating smart attorney redlined document...")
        redlined_path, clean_path = self.create_smart_redlined_document(instructions, redacted_docx_path, base_name)
        
        # Step 5: Create final documents with PII restoration
        print("[████] Step 4: Creating final documents with PII restoration...")
        
        if redlined_path and clean_path and mapping_path:
            # Copy original to final directory
            import shutil
            original_final = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Original.docx"
            shutil.copy2(docx_path, original_final)
            
            # Copy redacted to final directory  
            redacted_final = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Redacted.docx"
            shutil.copy2(redacted_docx_path, redacted_final)
            
            # Copy mapping to final directory
            mapping_final = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Mapping.json"
            shutil.copy2(mapping_path, mapping_final)
            
            # Create redlined document by applying redlines to ORIGINAL document (not restoring PII)
            print("   🔄 Applying redlines to original document for proper track changes...")
            redlined_original_path = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Smart_Attorney_Redlined_Original.docx"
            clean_original_path = f"/home/cliff/redact/redline_project/libreTest/{base_name}_Smart_Attorney_Clean_Original.docx"
            
            # Apply redlines to original document (with PII intact)
            redlined_orig, clean_orig = self.create_smart_redlined_document(instructions, docx_path, f"{base_name}_Original")
            
            if redlined_orig and clean_orig:
                # Move to final locations
                shutil.move(redlined_orig, redlined_original_path)
                shutil.move(clean_orig, clean_original_path)
                print(f"   ✅ Created redlined with original PII: {os.path.basename(redlined_original_path)}")
                print(f"   ✅ Created clean with original PII: {os.path.basename(clean_original_path)}")
            else:
                print("   ❌ Failed to create original redlined documents")
            
            print(f"\n📁 ALL 7 DOCUMENTS SAVED TO /home/cliff/redact/redline_project/libreTest/:")
            print(f"   1. Original: {original_final}")
            print(f"   2. Redacted: {redacted_final}")
            print(f"   3. Mapping: {mapping_final}")
            print(f"   4. Redlined (redacted): {redlined_path}")
            print(f"   5. Clean (redacted): {clean_path}")
            print(f"   6. Redlined (original): {redlined_original_path}")
            print(f"   7. Clean (original): {clean_original_path}")
        
        else:
            print("❌ Missing required files for PII restoration")
        
        print("✅ SMART ATTORNEY PATTERN SYSTEM COMPLETE!")
        print(f"📋 Pattern Analysis: {os.path.basename(analysis_file) if analysis_file else 'FAILED'}")
        print(f"📝 Smart Redlined: {os.path.basename(redlined_path) if redlined_path else 'FAILED'}")
        print(f"📄 Smart Clean: {os.path.basename(clean_path) if clean_path else 'FAILED'}")
        print(f"🎯 Attorney patterns applied: {len(instructions)}")
        
        # Show applied patterns
        if instructions:
            print("\n🧠 ATTORNEY PATTERNS APPLIED:")
            for i, inst in enumerate(instructions, 1):
                print(f"{i}. {inst.get('pattern_name', 'Unknown')}")
                print(f"   Reasoning: {inst.get('attorney_reasoning', 'Unknown')}")
                print(f"   Benefit: {inst.get('purchaser_benefit', 'Unknown')}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 smart_attorney_system.py input.docx")
        sys.exit(1)
    
    system = SmartAttorneySystem()
    system.process_smart_attorney_system(sys.argv[1])

if __name__ == "__main__":
    main()
