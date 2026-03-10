#!/usr/bin/env python3
"""
Complete NDA Redlining System
Creates: AI Problem Analysis + AI Recommendations + Redlined Doc + Clean Doc
"""

import os
import sys
import subprocess
import json
import re
import openai
from docx import Document
from docx.shared import RGBColor

class CompleteNDARedliningSystem:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = len(self.personal_info) + 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'broker', 'seller']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                if placeholder not in self.personal_info:
                    self.personal_info[placeholder] = match.group()
                    redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                    counter += 1
        
        return redacted_text
    
    def convert_with_libreoffice(self, input_path):
        """Convert document using LibreOffice"""
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            input_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
        
        return None
    
    def create_ai_problem_analysis(self, redacted_text, base_name):
        """Step 1: Create AI problem analysis"""
        if not self.openai_client:
            return None
        
        # Load Golden NDA
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        prompt = f"""You are an experienced real estate attorney analyzing an NDA against the Golden NDA standard.

GOLDEN NDA STANDARD:
{golden_nda}

NDA TO ANALYZE:
{redacted_text}

TASK: Identify what is WRONG with this NDA compared to the Golden NDA standard.

Create a detailed analysis identifying all problems:

# What's Wrong with This NDA

## Executive Summary
[Brief overview of major problems]

## Critical Legal Deficiencies

### 1. [Problem Category]
**Issue:** [What's wrong]
**Golden Standard:** [How Golden NDA handles this correctly]
**Risk:** [What legal/business risk this creates]
**Fix Required:** [What needs to be changed]

[Continue for all major problems]

## Missing Protections
[Critical clauses completely absent]

## Risk Assessment
**High Risk Issues:** [List most critical problems]
**Medium Risk Issues:** [List moderate problems]

## Conclusion
[Overall assessment of how far this NDA falls short of the gold standard]

Be thorough and identify every significant problem."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.1
            )
            
            analysis = response.choices[0].message.content
            
            # Save analysis
            analysis_file = f"/home/cliff/redact/redline_project/{base_name}_AI_Problem_Analysis.md"
            with open(analysis_file, 'w') as f:
                f.write(analysis)
            
            print(f"   ✓ Created AI problem analysis: {os.path.basename(analysis_file)}")
            return analysis_file
            
        except Exception as e:
            print(f"   ❌ AI problem analysis error: {e}")
            return None
    
    def create_ai_recommendations(self, redacted_text, base_name):
        """Step 2: Create AI recommendations with line numbers"""
        if not self.openai_client:
            return None
        
        # Add line numbers to text
        lines = redacted_text.split('\n')
        numbered_lines = []
        line_number = 1
        
        for line in lines:
            if line.strip():
                numbered_lines.append(f"LINE {line_number:03d}: {line.strip()}")
                line_number += 1
        
        line_numbered_text = '\n'.join(numbered_lines)
        
        # Load Golden NDA and problem analysis
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        prompt = f"""You are an attorney providing PRECISE redlining recommendations.

GOLDEN NDA STANDARD:
{golden_nda}

LINE-NUMBERED NDA:
{line_numbered_text}

TASK: Provide precise change instructions using line numbers. Address ALL major problems.

CRITICAL: Focus on adding missing definitions like "Confidential Information" in early paragraphs.

Return ONLY JSON array:
[
  {{
    "problem_addressed": "Definition of Confidential Information",
    "change_type": "replace_text_on_line",
    "line_number": 3,
    "find_text": "Informational Materials",
    "replace_with": "Confidential Information means any non-public financial, operational, legal, strategic, or business information relating to the Owner or the Property that is disclosed to the Potential Purchaser, whether orally, electronically, visually, or in written form. Confidential Information includes, without limitation: financial statements, rent rolls, leases, operating statements, tenant information, legal documentation, property information, transaction materials",
    "reason": "Adds comprehensive definition of confidential information",
    "priority": "HIGH"
  }},
  {{
    "problem_addressed": "Missing Purpose Clause",
    "change_type": "insert_after_line",
    "line_number": 2,
    "insert_text": "The purpose of this Agreement is to evaluate a potential transaction involving the Property.",
    "reason": "Explicitly states the purpose of disclosure",
    "priority": "HIGH"
  }}
]

Address 7-10 major problems with precise line numbers."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            
            # Save recommendations
            recommendations_file = f"/home/cliff/redact/redline_project/{base_name}_AI_Recommendations.md"
            with open(recommendations_file, 'w') as f:
                f.write(f"# AI Recommendations for {base_name}\n\n")
                f.write("**AI Analysis Output:**\n\n")
                f.write("```json\n")
                f.write(content)
                f.write("\n```\n\n")
            
            # Parse JSON
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = content[json_start:json_end]
                recommendations = json.loads(json_str)
                
                # Add detailed recommendations
                with open(recommendations_file, 'a') as f:
                    f.write(f"**Parsed Recommendations:** {len(recommendations)}\n\n")
                    
                    for i, rec in enumerate(recommendations, 1):
                        f.write(f"## Recommendation {i}: {rec.get('problem_addressed', 'Unknown')}\n\n")
                        f.write(f"**Priority:** {rec.get('priority', 'UNKNOWN')}\n\n")
                        f.write(f"**Change Type:** {rec.get('change_type', 'unknown')}\n\n")
                        f.write(f"**Line Number:** {rec.get('line_number', 'N/A')}\n\n")
                        f.write(f"**Find Text:**\n```\n{rec.get('find_text', 'N/A')}\n```\n\n")
                        f.write(f"**Replace/Insert:**\n```\n{rec.get('replace_with', rec.get('insert_text', 'N/A'))}\n```\n\n")
                        f.write(f"**Reason:** {rec.get('reason', 'N/A')}\n\n")
                        f.write("---\n\n")
                
                print(f"   ✓ Created AI recommendations: {os.path.basename(recommendations_file)}")
                return recommendations_file, recommendations
            else:
                print("   ❌ Could not parse AI recommendations JSON")
                return recommendations_file, []
                
        except Exception as e:
            print(f"   ❌ AI recommendations error: {e}")
            return None, []
    
    def implement_surgical_changes(self, original_path, recommendations, base_name):
        """Step 3 & 4: Implement changes and create redlined + clean docs"""
        # Convert to paragraphs
        text_content = self.convert_with_libreoffice(original_path)
        if not text_content:
            return None, None
        
        # Split into paragraphs
        paragraphs = []
        lines = text_content.split('\n')
        current_para = []
        para_number = 1
        
        for line in lines:
            if line.strip():
                current_para.append(line.strip())
            else:
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append({
                        'paragraph_number': para_number,
                        'original_text': para_text,
                        'current_text': para_text,
                        'changes': []
                    })
                    para_number += 1
                    current_para = []
        
        if current_para:
            para_text = ' '.join(current_para)
            paragraphs.append({
                'paragraph_number': para_number,
                'original_text': para_text,
                'current_text': para_text,
                'changes': []
            })
        
        # Implement changes
        implemented_count = 0
        
        for rec in recommendations:
            try:
                if rec.get('change_type') == 'replace_text_on_line':
                    line_num = rec.get('line_number')
                    find_text = rec.get('find_text', '')
                    replace_text = rec.get('replace_with', '')
                    
                    # Find best matching paragraph
                    for para in paragraphs:
                        if find_text[:30] in para['current_text'] or find_text in para['current_text']:
                            # Make surgical change
                            if find_text in para['current_text']:
                                new_text = para['current_text'].replace(find_text, replace_text, 1)
                            else:
                                # Partial match - replace key phrase
                                key_words = find_text.split()[:5]
                                key_phrase = ' '.join(key_words)
                                if key_phrase in para['current_text']:
                                    new_text = para['current_text'].replace(key_phrase, replace_text, 1)
                                else:
                                    continue
                            
                            para['current_text'] = new_text
                            para['changes'].append({
                                'type': 'replace',
                                'find': find_text if find_text in para['original_text'] else key_phrase,
                                'replace': replace_text,
                                'reason': rec.get('reason', '')
                            })
                            implemented_count += 1
                            break
                
                elif rec.get('change_type') == 'insert_after_line':
                    line_num = rec.get('line_number')
                    insert_text = rec.get('insert_text', '')
                    
                    if line_num <= len(paragraphs):
                        new_para = {
                            'paragraph_number': None,
                            'original_text': '',
                            'current_text': insert_text,
                            'changes': [{
                                'type': 'insert',
                                'text': insert_text,
                                'reason': rec.get('reason', '')
                            }]
                        }
                        paragraphs.insert(line_num, new_para)
                        implemented_count += 1
                
                elif rec.get('change_type') == 'insert_at_end':
                    insert_text = rec.get('insert_text', '')
                    new_para = {
                        'paragraph_number': None,
                        'original_text': '',
                        'current_text': insert_text,
                        'changes': [{
                            'type': 'insert',
                            'text': insert_text,
                            'reason': rec.get('reason', '')
                        }]
                    }
                    paragraphs.append(new_para)
                    implemented_count += 1
                    
            except Exception as e:
                print(f"   ❌ Failed to implement {rec.get('problem_addressed', 'unknown')}: {e}")
        
        # Create redlined document
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_Redlined.docx"
        doc = Document()
        
        # Header
        header = doc.add_paragraph()
        header_run = header.add_run("AI REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Deleted | 🟢 Green underline = Added")
        doc.add_paragraph("=" * 80)
        
        # Process paragraphs
        for para_data in paragraphs:
            doc_para = doc.add_paragraph()
            
            if para_data['changes']:
                for change in para_data['changes']:
                    if change['type'] == 'replace':
                        original_text = para_data['original_text']
                        find_text = change['find']
                        replace_text = change['replace']
                        
                        if find_text in original_text:
                            parts = original_text.split(find_text)
                            
                            if parts[0]:
                                doc_para.add_run(parts[0])
                            
                            # Deletion
                            del_run = doc_para.add_run(find_text)
                            del_run.font.strike = True
                            del_run.font.color.rgb = RGBColor(255, 0, 0)
                            
                            # Addition
                            ins_run = doc_para.add_run(replace_text)
                            ins_run.underline = True
                            ins_run.font.color.rgb = RGBColor(0, 128, 0)
                            
                            if len(parts) > 1 and parts[1]:
                                doc_para.add_run(parts[1])
                        else:
                            doc_para.add_run(para_data['current_text'])
                    
                    elif change['type'] == 'insert':
                        label_run = doc_para.add_run("[AI ADDITION] ")
                        label_run.bold = True
                        label_run.font.color.rgb = RGBColor(0, 100, 0)
                        
                        ins_run = doc_para.add_run(change['text'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                if para_data['current_text']:
                    doc_para.add_run(para_data['current_text'])
        
        doc.save(redlined_path)
        
        # Create clean document
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_Clean.docx"
        clean_doc = Document()
        
        for para_data in paragraphs:
            if para_data['current_text']:
                para = clean_doc.add_paragraph()
                para.add_run(para_data['current_text'])
        
        clean_doc.save(clean_path)
        
        print(f"   ✓ Created redlined document: {os.path.basename(redlined_path)}")
        print(f"   ✓ Created clean document: {os.path.basename(clean_path)}")
        print(f"   ✓ Implemented {implemented_count}/{len(recommendations)} changes")
        
        return redlined_path, clean_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information in document"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
    
    def process_complete_nda_redlining(self, input_path):
        """Complete end-to-end NDA redlining system"""
        print(f"🔄 Complete NDA Redlining: {os.path.basename(input_path)}")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        # Convert document
        print("[█░░░░] Step 1: Converting document with LibreOffice...")
        text_content = self.convert_with_libreoffice(input_path)
        if not text_content:
            print("❌ LibreOffice conversion failed")
            return
        
        # Redact personal info
        redacted_text = self.redact_personal_info(text_content)
        print(f"   ✓ Redacted {len(self.personal_info)} personal items")
        
        # Step 1: AI Problem Analysis
        print("[██░░░] Step 2: Creating AI problem analysis...")
        analysis_file = self.create_ai_problem_analysis(redacted_text, base_name)
        
        # Step 2: AI Recommendations
        print("[███░░] Step 3: Creating AI recommendations...")
        recommendations_file, recommendations = self.create_ai_recommendations(redacted_text, base_name)
        
        # Step 3 & 4: Implementation
        print("[████░] Step 4: Implementing changes and creating documents...")
        redlined_path, clean_path = self.implement_surgical_changes(input_path, recommendations, base_name)
        
        # Step 5: Restore personal info
        print("[█████] Step 5: Restoring personal information...")
        if redlined_path:
            self.restore_personal_info(redlined_path)
        if clean_path:
            self.restore_personal_info(clean_path)
        
        print("✅ COMPLETE NDA REDLINING SYSTEM FINISHED!")
        print(f"📋 AI Problem Analysis: {os.path.basename(analysis_file) if analysis_file else 'FAILED'}")
        print(f"📋 AI Recommendations: {os.path.basename(recommendations_file) if recommendations_file else 'FAILED'}")
        print(f"📝 Redlined Document: {os.path.basename(redlined_path) if redlined_path else 'FAILED'}")
        print(f"📄 Clean Document: {os.path.basename(clean_path) if clean_path else 'FAILED'}")
        print(f"🎯 Total AI recommendations: {len(recommendations)}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 complete_nda_redlining_system.py input.docx")
        sys.exit(1)
    
    system = CompleteNDARedliningSystem()
    system.process_complete_nda_redlining(sys.argv[1])

if __name__ == "__main__":
    main()
