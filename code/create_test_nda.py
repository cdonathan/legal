#!/usr/bin/env python3

from docx import Document

def create_clean_test_nda():
    """Create a clean NDA with realistic PII for testing"""
    
    doc = Document()
    
    # Title
    doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
    
    # Main content with PII
    content = """
This Non-Disclosure Agreement ("Agreement") is entered into on March 9, 2026, between John Smith, an individual residing at 123 Oak Street, Springfield, IL 62701 (email: john.smith@email.com, phone: (555) 123-4567), hereinafter referred to as the "Disclosing Party," and ABC Corporation, a Delaware corporation with its principal place of business at 456 Pine Avenue, Chicago, IL 60601 (email: legal@abccorp.com, phone: (555) 987-6543), hereinafter referred to as the "Receiving Party."

WHEREAS, the Disclosing Party possesses certain confidential and proprietary information relating to real estate investments; and

WHEREAS, the Receiving Party desires to evaluate such information for potential business opportunities;

NOW, THEREFORE, in consideration of the mutual covenants contained herein, the parties agree as follows:

1. CONFIDENTIAL INFORMATION
For purposes of this Agreement, "Confidential Information" means all non-public, proprietary information disclosed by John Smith to ABC Corporation.

2. OBLIGATIONS OF RECEIVING PARTY
ABC Corporation agrees to maintain the confidentiality of all Confidential Information and not to disclose such information to any third parties without the prior written consent of John Smith.

3. RETURN OF MATERIALS
Upon termination of this Agreement, ABC Corporation shall return all materials containing Confidential Information to John Smith at 123 Oak Street, Springfield, IL 62701.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

DISCLOSING PARTY:                    RECEIVING PARTY:

_________________________           _________________________
John Smith                          Sarah Johnson, CEO
Date: 3/9/2026                      ABC Corporation
                                    Date: 3/9/2026

Contact Information:
John Smith: john.smith@email.com, (555) 123-4567
Sarah Johnson: sarah.johnson@abccorp.com, (555) 987-6543
"""
    
    # Add content as paragraphs
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    # Save the document
    output_path = "/tmp/clean_test_nda.docx"
    doc.save(output_path)
    print(f"✅ Created clean test NDA: {output_path}")
    return output_path

if __name__ == "__main__":
    create_clean_test_nda()
