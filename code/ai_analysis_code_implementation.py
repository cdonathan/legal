#!/usr/bin/env python3
"""
AI Analysis + Code Implementation System
AI analyzes and recommends, code implements the changes
"""

import openai
import json
import os
import re
from docx import Document
from docx.shared import RGBColor

class AIAnalysisCodeImplementation:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_ai_recommendations(self, redacted_text):
        """AI analyzes and provides change recommendations"""
        if not self.openai_client:
            return []
        
        # Load the problems analysis
        with open('/home/cliff/redact/redline_project/Sample_2_Problems_Analysis.md', 'r') as f:
            problems_analysis = f.read()
        
        # Load Golden NDA for reference
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        prompt = f"""You are an attorney analyzing Sample 2 NDA. Provide specific change recommendations to fix identified problems.

GOLDEN NDA STANDARD:
{golden_nda}

PROBLEMS IDENTIFIED:
{problems_analysis}

SAMPLE 2 NDA:
{redacted_text}

TASK: Analyze the document and provide specific change instructions. DO NOT implement changes - only recommend them.

Return ONLY JSON array of recommendations:
[
  {{
    "problem_addressed": "Definition of Confidential Information",
    "change_type": "replace_text",
    "find_text": "exact text to find in document",
    "replace_with": "exact replacement text",
    "reason": "why this fixes the identified problem",
    "priority": "HIGH/MEDIUM/LOW"
  }},
  {{
    "problem_addressed": "Missing Purpose Clause",
    "change_type": "insert_after_paragraph",
    "paragraph_number": 1,
    "insert_text": "exact text to insert",
    "reason": "why this addresses the missing clause",
    "priority": "HIGH/MEDIUM/LOW"
  }}
]

Focus on HIGH priority problems first. Provide exact text matches for replacements."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            
            # Save AI recommendations
            with open('/home/cliff/redact/redline_project/AI_Recommendations_Sample_2.md', 'w') as f:
                f.write("# AI Change Recommendations for Sample 2\n\n")
                f.write("**AI Analysis Output:**\n\n")
                f.write("```\n")
                f.write(content)
                f.write("\n```\n\n")
            
            # Extract JSON
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = content[json_start:json_end]
                recommendations = json.loads(json_str)
                
                # Add detailed recommendations to file
                with open('/home/cliff/redact/redline_project/AI_Recommendations_Sample_2.md', 'a') as f:
                    f.write(f"**Parsed Recommendations:** {len(recommendations)}\n\n")
                    
                    for i, rec in enumerate(recommendations, 1):
                        f.write(f"## Recommendation {i}: {rec.get('problem_addressed', 'Unknown')}\n\n")
                        f.write(f"**Priority:** {rec.get('priority', 'UNKNOWN')}\n\n")
                        f.write(f"**Change Type:** {rec.get('change_type', 'unknown')}\n\n")
                        f.write(f"**Find Text:**\n```\n{rec.get('find_text', 'N/A')}\n```\n\n")
                        f.write(f"**Replace With:**\n```\n{rec.get('replace_with', rec.get('insert_text', 'N/A'))}\n```\n\n")
                        f.write(f"**Reason:** {rec.get('reason', 'N/A')}\n\n")
                        f.write("---\n\n")
                
                return recommendations
            else:
                return []
                
        except Exception as e:
            print(f"AI recommendation error: {e}")
            return []
    
    def implement_recommendations(self, original_doc_path, recommendations, output_path_redlined, output_path_clean):
        """Code implements the AI recommendations"""
        doc = Document(original_doc_path)
        
        print(f"📋 Implementing {len(recommendations)} AI recommendations...")
        
        # Create redlined version
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("AI RECOMMENDATIONS IMPLEMENTED")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = AI identified problems | 🟢 Green underline = AI recommended fixes")
        
        # Show AI recommendations summary
        summary = redlined_doc.add_paragraph()
        summary_run = summary.add_run(f"AI Recommendations: {len(recommendations)} changes to fix identified problems")
        summary_run.italic = True
        
        redlined_doc.add_paragraph("=" * 80)
        
        # Process each paragraph
        paragraph_index = 0
        implemented_count = 0
        
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply text replacements
            for rec in recommendations:
                if rec.get('change_type') == 'replace_text' and rec.get('find_text') in para_text:
                    parts = para_text.split(rec['find_text'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Show AI identified problem
                    del_run = new_para.add_run(rec['find_text'])
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Show AI recommended fix
                    ins_run = new_para.add_run(rec['replace_with'])
                    ins_run.underline = True
                    ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    implemented_count += 1
                    break
            
            if para_text:
                new_para.add_run(para_text)
            
            # Add insertions
            for rec in recommendations:
                if (rec.get('change_type') == 'insert_after_paragraph' and 
                    rec.get('paragraph_number') == paragraph_index):
                    
                    insert_para = redlined_doc.add_paragraph()
                    
                    label_run = insert_para.add_run(f"[AI RECOMMENDATION: {rec['problem_addressed']}] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(0, 100, 0)
                    
                    insert_run = insert_para.add_run(rec['insert_text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    reason_para = redlined_doc.add_paragraph()
                    reason_run = reason_para.add_run(f"AI REASON: {rec['reason']}")
                    reason_run.italic = True
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)
                    
                    implemented_count += 1
            
            paragraph_index += 1
        
        # Create clean version
        clean_doc = Document()
        paragraph_index = 0
        
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for rec in recommendations:
                if rec.get('change_type') == 'replace_text' and rec.get('find_text') in para_text:
                    para_text = para_text.replace(rec['find_text'], rec['replace_with'])
            
            new_para.add_run(para_text)
            
            # Add insertions
            for rec in recommendations:
                if (rec.get('change_type') == 'insert_after_paragraph' and 
                    rec.get('paragraph_number') == paragraph_index):
                    insert_para = clean_doc.add_paragraph()
                    insert_para.add_run(rec['insert_text'])
            
            paragraph_index += 1
        
        # Save both versions
        redlined_doc.save(output_path_redlined)
        clean_doc.save(output_path_clean)
        
        print(f"   ✓ Successfully implemented {implemented_count} AI recommendations")
        return implemented_count
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Complete AI analysis + code implementation workflow"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: AI Analysis
        print("[██░░░] Step 2: Getting AI recommendations...")
        recommendations = self.get_ai_recommendations(redacted_text)
        print(f"   ✓ AI provided {len(recommendations)} recommendations")
        
        # Step 3: Code Implementation
        print("[███░░] Step 3: Implementing AI recommendations...")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_ai_recommended_redlined.docx"
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_ai_recommended_clean-version.docx"
        
        implemented = self.implement_recommendations(input_path, recommendations, redlined_path, clean_path)
        
        # Step 4: Restore personal info
        print("[████░] Step 4: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("[█████] Step 5: Complete!")
        print("✅ AI ANALYSIS + CODE IMPLEMENTATION COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"📋 AI Recommendations: AI_Recommendations_Sample_2.md")
        print(f"🎯 AI recommended {len(recommendations)} changes, code implemented {implemented}")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 ai_analysis_code_implementation.py input.docx")
        sys.exit(1)
    
    processor = AIAnalysisCodeImplementation()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
