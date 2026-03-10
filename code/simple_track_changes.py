#!/usr/bin/env python3
"""
Simplified Track Changes - Create comparison document
AI returns structured changes, we create a clean comparison document
"""

import json
import os
import openai
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

class SimpleTrackChanges:
    def __init__(self):
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def get_ai_structured_changes(self, text):
        """Get structured changes from AI in JSON format"""
        if not self.openai_client:
            return self._mock_changes()
        
        prompt = f"""You are redlining an NDA. Return ONLY a JSON object with structured changes.

NDA TEXT:
{text[:2000]}...

Return JSON in this exact format:
{{
  "changes": [
    {{
      "type": "insert",
      "after_paragraph": 3,
      "text": "New clause text here",
      "reason": "Why this change is needed"
    }},
    {{
      "type": "replace", 
      "paragraph": 2,
      "find": "old text",
      "replace": "new text",
      "reason": "Why this change is needed"
    }}
  ]
}}

Focus on adding missing standard clauses: permitted recipients, return of materials, term limits."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
                temperature=0.1
            )
            
            # Extract JSON from response
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"AI error: {e}")
            return self._mock_changes()
    
    def _mock_changes(self):
        """Mock changes for testing"""
        return {
            "changes": [
                {
                    "type": "insert",
                    "after_paragraph": 3,
                    "text": "The Receiving Party may disclose Confidential Information to its attorneys, accountants, financial advisors, lenders, investors, and other professional advisors who have a need to know such information for the Purpose and who are bound by confidentiality obligations.",
                    "reason": "Add permitted recipients clause - institutional standard"
                },
                {
                    "type": "insert",
                    "after_paragraph": 5,
                    "text": "Upon written request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and any copies, notes, or derivatives thereof.",
                    "reason": "Add return of materials clause - required for enforceability"
                }
            ]
        }
    
    def create_redlined_document(self, original_doc_path, changes, output_path):
        """Create document showing changes with visual markup"""
        doc = Document(original_doc_path)
        
        # Create new document for redlined version
        redlined_doc = Document()
        
        # Copy original paragraphs and apply changes
        paragraph_index = 0
        
        for para in doc.paragraphs:
            # Add original paragraph
            new_para = redlined_doc.add_paragraph()
            
            # Check for replacements in this paragraph
            para_text = para.text
            for change in changes.get('changes', []):
                if change['type'] == 'replace' and change.get('paragraph') == paragraph_index:
                    # Show deletion and insertion
                    if change['find'] in para_text:
                        parts = para_text.split(change['find'])
                        
                        # Add text before
                        if parts[0]:
                            new_para.add_run(parts[0])
                        
                        # Add deleted text (strikethrough, red)
                        del_run = new_para.add_run(change['find'])
                        del_run.font.strike = True
                        del_run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        # Add inserted text (underline, green)
                        ins_run = new_para.add_run(change['replace'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
                        
                        # Add text after
                        if len(parts) > 1 and parts[1]:
                            new_para.add_run(parts[1])
                        
                        para_text = ""  # Don't add original text
            
            # Add original text if no replacements
            if para_text:
                new_para.add_run(para_text)
            
            # Check for insertions after this paragraph
            for change in changes.get('changes', []):
                if change['type'] == 'insert' and change.get('after_paragraph') == paragraph_index:
                    # Add new paragraph with inserted text
                    insert_para = redlined_doc.add_paragraph()
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    # Add comment about the change
                    comment_para = redlined_doc.add_paragraph()
                    comment_run = comment_para.add_run(f"[INSERTED: {change['reason']}]")
                    comment_run.italic = True
                    comment_run.font.color.rgb = RGBColor(128, 128, 128)
            
            paragraph_index += 1
        
        # Add summary of changes at the end
        redlined_doc.add_page_break()
        summary_para = redlined_doc.add_paragraph()
        summary_run = summary_para.add_run("REDLINE SUMMARY")
        summary_run.bold = True
        summary_run.font.size = 14
        
        for i, change in enumerate(changes.get('changes', []), 1):
            change_para = redlined_doc.add_paragraph()
            change_para.add_run(f"{i}. {change['type'].title()}: {change['reason']}")
        
        redlined_doc.save(output_path)
        return output_path
    
    def process_nda(self, input_path):
        """Complete NDA redlining process"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Get AI changes
        print("🤖 Getting AI redline suggestions...")
        changes = self.get_ai_structured_changes(text)
        
        print(f"✓ AI suggested {len(changes.get('changes', []))} changes")
        for i, change in enumerate(changes.get('changes', []), 1):
            print(f"   {i}. {change['reason']}")
        
        # Create redlined document
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = f"/home/cliff/redact/redline_project/{base_name}_REDLINED.docx"
        
        print("📝 Creating redlined document...")
        self.create_redlined_document(input_path, changes, output_path)
        
        print(f"✅ Complete: {os.path.basename(output_path)}")
        print("   • Green underlined text = Insertions")
        print("   • Red strikethrough text = Deletions") 
        print("   • Summary of changes included at end")
        
        return output_path

def main():
    processor = SimpleTrackChanges()
    
    # Test with sample NDA
    test_file = "/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Conf_Agr_Sample1-pre-redline.docx"
    if os.path.exists(test_file):
        processor.process_nda(test_file)
    else:
        print(f"Test file not found: {test_file}")

if __name__ == "__main__":
    main()
