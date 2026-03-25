#!/usr/bin/env python3
"""
Smart Attorney Pattern-Based NDA Redlining System v2
Uses assessment-first prompt with adaptive terminology matching.
"""

import os
import sys
import subprocess
import json
import secrets
import re
import shutil
import time
from datetime import datetime
import uno
from com.sun.star.beans import PropertyValue

OUTPUT_DIR = "/home/cliff/redact/OneDrive_1_3-23-2026/TestOutput"

class SmartAttorneySystem:
    def __init__(self):
        self.openai_client = None
        self.patterns_prompt = None
        self.setup_openai()
        self.load_patterns_prompt()
        os.makedirs(OUTPUT_DIR, exist_ok=True)

    def setup_openai(self):
        try:
            import openai
            api_key_file = "/home/cliff/redact/openai_api_key.txt"
            if os.path.exists(api_key_file):
                with open(api_key_file, 'r') as f:
                    api_key = f.read().strip()
                self.openai_client = openai.OpenAI(api_key=api_key)
                print("✅ OpenAI client initialized")
            else:
                print("❌ OpenAI API key file not found")
        except ImportError:
            print("❌ OpenAI library not installed")
        except Exception as e:
            print(f"❌ OpenAI setup failed: {e}")

    def load_patterns_prompt(self):
        try:
            p1 = "/home/cliff/redact/redline_project/components/prompt_call1_analysis.md"
            p2 = "/home/cliff/redact/redline_project/components/prompt_call2_changes.md"
            if os.path.exists(p1) and os.path.exists(p2):
                with open(p1, 'r') as f:
                    self.prompt_call1 = f.read()
                with open(p2, 'r') as f:
                    self.prompt_call2 = f.read()
                self.patterns_prompt = True  # flag that prompts loaded
                print("✅ Two-call prompts loaded")
            else:
                print("❌ Prompt files not found")
                self.patterns_prompt = None
        except Exception as e:
            print(f"❌ Failed to load prompts: {e}")

    def convert_to_docx(self, input_path):
        if input_path.endswith('.docx'):
            return input_path
        try:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = f"/tmp/{base_name}_converted.docx"
            cmd = ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', '/tmp', input_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and os.path.exists(output_path):
                return output_path
            else:
                print(f"   ❌ Conversion failed: {result.stderr}")
                return None
        except Exception as e:
            print(f"   ❌ Conversion error: {e}")
            return None

    def convert_with_libreoffice(self, docx_path):
        try:
            cmd = ['libreoffice', '--headless', '--convert-to', 'txt', '--outdir', '/tmp', docx_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                txt_file = f"/tmp/{base_name}.txt"
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        return f.read()
            return None
        except Exception as e:
            print(f"   ❌ Text conversion error: {e}")
            return None

    def redact_docx_pii(self, docx_path, base_name):
        """Redact PII directly in a copy of the DOCX, preserving formatting."""
        try:
            redacted_path = f"/tmp/{base_name}_redacted.docx"
            mapping_path = f"/tmp/{base_name}_mapping.json"
            shutil.copy2(docx_path, redacted_path)

            text_content = self.convert_with_libreoffice(docx_path)
            if not text_content:
                return None, None

            _, hex_mapping = self.apply_hex_redaction(text_content)
            if not hex_mapping:
                print(f"   ✓ No PII found to redact")
                with open(mapping_path, 'w') as f:
                    json.dump({}, f)
                return redacted_path, mapping_path

            # Apply redactions in the DOCX preserving formatting
            from docx import Document
            doc = Document(redacted_path)
            for hex_id, data in hex_mapping.items():
                original = data['original']
                placeholder = data['placeholder']
                for para in doc.paragraphs:
                    if original in para.text:
                        for run in para.runs:
                            if original in run.text:
                                run.text = run.text.replace(original, placeholder)

            doc.save(redacted_path)
            with open(mapping_path, 'w') as f:
                json.dump(hex_mapping, f, indent=2)

            print(f"   ✓ Redacted {len(hex_mapping)} PII items with hex mapping")
            return redacted_path, mapping_path
        except Exception as e:
            print(f"   ❌ Redaction failed: {e}")
            return None, None

    def load_whitelist(self):
        """Load redaction whitelist — words that should NOT be redacted."""
        whitelist = set()
        wl_path = "/home/cliff/redact/redaction_whitelist.txt"
        if os.path.exists(wl_path):
            with open(wl_path, 'r') as f:
                for line in f:
                    word = line.strip().lower()
                    if word and not word.startswith('#'):
                        whitelist.add(word)
        return whitelist

    def apply_hex_redaction(self, text):
        whitelist = self.load_whitelist()
        patterns = [
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL'),
            (r'\(\d{3}\)\s*\d{3}-\d{4}', 'PHONE'),
            (r'\d{3}-\d{3}-\d{4}', 'PHONE'),
            (r'\b\d+\s+[A-Z][a-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Court|Ct|Boulevard|Blvd)(?:\s*,\s*[A-Z][a-z]+)?(?:\s*,\s*[A-Z]{2})?\s*\d{5}?\b', 'ADDRESS'),
            (r'\b[A-Z][a-z]+\s+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Court|Ct|Boulevard|Blvd)\b', 'STREET'),
            (r'\b\d{5}(?:-\d{4})?\b', 'ZIP'),
            (r'\b[A-Z][A-Za-z\s&]+(?:LLC|Inc|Corp|Corporation|Company|Co\.)\b', 'COMPANY'),
            (r'\b\d{1,2}/\d{1,2}/\d{4}\b', 'DATE'),
            (r'\b\d{1,2}-\d{1,2}-\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', 'PERSON'),
        ]
        redacted_text = text
        hex_mapping = {}
        for pattern, label in patterns:
            matches = list(re.finditer(pattern, redacted_text))
            for match in matches:
                original = match.group()

                # For PERSON matches, skip if any word is in the whitelist
                if label == 'PERSON':
                    words = original.split()
                    if any(w.lower() in whitelist for w in words):
                        continue

                hex_id = secrets.token_hex(8)
                placeholder = f"[{label}:{hex_id}]"
                hex_mapping[hex_id] = {
                    'type': label,
                    'original': original,
                    'placeholder': placeholder
                }
                redacted_text = redacted_text.replace(original, placeholder, 1)
        return redacted_text, hex_mapping

    def restore_pii_in_docx(self, docx_path, mapping_path):
        """Restore PII in a DOCX by replacing hex placeholders with originals."""
        try:
            with open(mapping_path, 'r') as f:
                hex_mapping = json.load(f)
            if not hex_mapping:
                return True

            from docx import Document
            doc = Document(docx_path)
            restored = 0
            for para in doc.paragraphs:
                for hex_id, data in hex_mapping.items():
                    if data['placeholder'] in para.text:
                        for run in para.runs:
                            if data['placeholder'] in run.text:
                                run.text = run.text.replace(data['placeholder'], data['original'])
                                restored += 1
            doc.save(docx_path)
            print(f"   ✓ Restored {restored} PII items in {os.path.basename(docx_path)}")
            return True
        except Exception as e:
            print(f"   ❌ PII restoration failed: {e}")
            return False

    def smart_attorney_analysis(self, redacted_text, base_name):
        if not self.openai_client or not self.patterns_prompt:
            print("   ❌ OpenAI client or prompts not available")
            return None, None

        try:
            lines = redacted_text.split('\n')
            numbered_lines = []
            for i, line in enumerate(lines, 1):
                if line.strip():
                    numbered_lines.append(f"LINE {i:03d}: {line}")
            numbered_text = '\n'.join(numbered_lines)
            today = datetime.now().strftime('%B %d, %Y')

            # === CALL 1: Analysis & Scoring ===
            print(f"   🔄 Call 1: Analysis & Scoring...")
            start = datetime.now()

            call1_prompt = f"""{self.prompt_call1}

LINE-NUMBERED NDA:
{numbered_text}

Respond with valid JSON only — no markdown fencing, no commentary."""

            resp1 = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert real estate attorney. Read the NDA carefully and quote specific language for each category. Respond with valid JSON only."},
                    {"role": "user", "content": call1_prompt}
                ],
                temperature=0.1,
                max_tokens=5000
            )

            analysis_text = resp1.choices[0].message.content.strip()
            dur1 = (datetime.now() - start).total_seconds()
            print(f"   ✅ Call 1 done in {dur1:.1f}s")

            # Save raw Call 1
            with open(os.path.join(OUTPUT_DIR, f"{base_name}_call1_raw.txt"), 'w') as f:
                f.write(analysis_text)

            # Parse Call 1
            clean1 = analysis_text
            if clean1.startswith('```'):
                clean1 = re.sub(r'^```(?:json)?\s*', '', clean1)
                clean1 = re.sub(r'\s*```$', '', clean1)
            clean1 = re.sub(r':\s*0+(\d+)', r': \1', clean1)
            clean1 = clean1.replace('\t', '\\t').replace('\r', '\\r')
            analysis = json.loads(clean1)

            # Print scores
            total = analysis.get('total_score', 0)
            print(f"   📊 Total score: {total}")
            for k, v in analysis.get('analysis', {}).items():
                score = v.get('score', v) if isinstance(v, dict) else v
                if isinstance(v, dict) and v.get('score', 0) > 0:
                    print(f"      {k}: {v['score']} — {v.get('reason', '')}")

            # === THRESHOLD ENFORCEMENT (Python, not AI) ===
            scores = {}
            for k, v in analysis.get('analysis', {}).items():
                if isinstance(v, dict):
                    scores[k] = v.get('score', 0)
                else:
                    scores[k] = v

            items_8_plus = [k for k, s in scores.items() if s >= 8]
            items_5_7 = [k for k, s in scores.items() if 5 <= s <= 7]
            items_1_4 = [k for k, s in scores.items() if 1 <= s <= 4]

            applied_items = list(items_8_plus)  # Always change 8-10

            # 5-7: change if total >= 15, or 3+ items in 5-7, or any item 8+
            if items_5_7 and (total >= 15 or len(items_5_7) >= 3 or len(items_8_plus) > 0):
                applied_items.extend(items_5_7)

            # 1-4: change only if total >= 20 or 2+ items 8+
            if items_1_4 and (total >= 20 or len(items_8_plus) >= 2):
                applied_items.extend(items_1_4)

            threshold_result = {
                'items_8_plus': items_8_plus,
                'items_5_7': items_5_7,
                'items_1_4': items_1_4,
                'applied_items': applied_items
            }

            print(f"   📋 Threshold: {len(applied_items)} items qualify for changes")
            for item in applied_items:
                print(f"      ✓ {item} (score {scores.get(item, '?')})")

            if not applied_items:
                print("   ✅ No items clear threshold — document is adequate")
                instructions = {'patterns': [], 'document_assessment': {
                    'confidential_info_term': analysis.get('terminology', {}).get('confidential_info_term', 'N/A'),
                    'recipient_term': analysis.get('terminology', {}).get('recipient_term', 'N/A'),
                    'discloser_term': analysis.get('terminology', {}).get('discloser_term', 'N/A'),
                    'representatives_term': analysis.get('terminology', {}).get('representatives_term', 'N/A'),
                    'sophistication': analysis.get('sophistication', 'N/A'),
                    'total_score': total,
                    'changes_needed': 'none'
                }}
                analysis_file = os.path.join(OUTPUT_DIR, f"{base_name}_Smart_Attorney_Analysis.md")
                with open(analysis_file, 'w') as f:
                    f.write(f"# Smart Attorney Analysis: {base_name}\n\nNo items cleared threshold.\n")
                return analysis_file, instructions

            # Build applied items detail for Call 2
            applied_detail = []
            for item in applied_items:
                entry = analysis.get('analysis', {}).get(item, {})
                applied_detail.append(f"- {item}: score {scores.get(item, '?')} — {entry.get('reason', 'N/A')}")
            applied_items_text = '\n'.join(applied_detail)

            # === CALL 2: Generate Changes ===
            print(f"   🔄 Call 2: Generating changes...")
            start = datetime.now()

            call2_prompt = f"""{self.prompt_call2}

CURRENT DATE: {today}

ANALYSIS FROM CALL 1:
{analysis_text}

ITEMS TO CHANGE (determined by threshold rules — generate a pattern for EACH of these and ONLY these):
{applied_items_text}

LINE-NUMBERED NDA:
{numbered_text}

Respond with valid JSON only — no markdown fencing, no commentary."""

            resp2 = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert real estate attorney. Using the analysis provided, generate precise redline changes. Respond with valid JSON only."},
                    {"role": "user", "content": call2_prompt}
                ],
                temperature=0.1,
                max_tokens=10000
            )

            changes_text = resp2.choices[0].message.content.strip()
            dur2 = (datetime.now() - start).total_seconds()
            print(f"   ✅ Call 2 done in {dur2:.1f}s")

            # Save raw Call 2
            with open(os.path.join(OUTPUT_DIR, f"{base_name}_call2_raw.txt"), 'w') as f:
                f.write(changes_text)

            # Parse Call 2
            clean2 = changes_text
            if clean2.startswith('```'):
                clean2 = re.sub(r'^```(?:json)?\s*', '', clean2)
                clean2 = re.sub(r'\s*```$', '', clean2)
            clean2 = re.sub(r':\s*0+(\d+)', r': \1', clean2)
            clean2 = clean2.replace('\t', '\\t').replace('\r', '\\r')
            instructions = json.loads(clean2)

            if 'patterns' not in instructions:
                instructions['patterns'] = []

            # Merge assessment info for downstream
            instructions['document_assessment'] = {
                'confidential_info_term': analysis.get('terminology', {}).get('confidential_info_term', 'N/A'),
                'recipient_term': analysis.get('terminology', {}).get('recipient_term', 'N/A'),
                'discloser_term': analysis.get('terminology', {}).get('discloser_term', 'N/A'),
                'representatives_term': analysis.get('terminology', {}).get('representatives_term', 'N/A'),
                'sophistication': analysis.get('sophistication', 'N/A'),
                'total_score': total,
                'changes_needed': 'none' if not instructions['patterns'] else 'moderate'
            }

            # Save analysis report
            analysis_file = os.path.join(OUTPUT_DIR, f"{base_name}_Smart_Attorney_Analysis.md")
            with open(analysis_file, 'w') as f:
                f.write(f"# Smart Attorney Analysis: {base_name}\n\n")
                f.write(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(f"## Terminology\n")
                for k, v in analysis.get('terminology', {}).items():
                    f.write(f"- **{k}:** {v}\n")
                f.write(f"\n## Category Scores (Total: {total})\n\n")
                for k, v in analysis.get('analysis', {}).items():
                    if isinstance(v, dict):
                        f.write(f"### {k}: Score {v.get('score', 'N/A')}\n")
                        f.write(f"- **Reason:** {v.get('reason', 'N/A')}\n")
                        ql = v.get('quoted_language', '')
                        if ql and ql != 'NOT FOUND':
                            f.write(f"- **Quoted:** {ql[:200]}...\n")
                        elif ql == 'NOT FOUND':
                            f.write(f"- **Quoted:** NOT FOUND\n")
                        f.write(f"\n")
                ta = instructions.get('threshold_analysis', {})
                f.write(f"## Threshold Analysis\n")
                f.write(f"- **Applied:** {ta.get('applied_items', [])}\n\n")
                f.write(f"## Changes ({len(instructions['patterns'])})\n\n")
                for i, p in enumerate(instructions['patterns'], 1):
                    f.write(f"### {i}. {p.get('title', 'Unknown')} (score: {p.get('score', '?')})\n")
                    f.write(f"- **Line:** {p.get('line_number', 'N/A')}\n")
                    f.write(f"- **Reasoning:** {p.get('reasoning', 'N/A')}\n\n")

            print(f"   ✓ {len(instructions['patterns'])} patterns to apply")
            return analysis_file, instructions

        except json.JSONDecodeError as e:
            print(f"   ❌ JSON parse error: {e}")
            err_file = os.path.join(OUTPUT_DIR, f"{base_name}_parse_error.txt")
            with open(err_file, 'w') as f:
                f.write(f"JSON Error: {e}\n\n")
                f.write(locals().get('analysis_text', '') + '\n---\n' + locals().get('changes_text', ''))
            return None, None
        except Exception as e:
            print(f"   ❌ Analysis error: {e}")
            return None, None

    def apply_pattern_to_document(self, document, pattern):
        try:
            original_text = pattern.get('original_text', '')
            replacement_text = pattern.get('replacement_text', '')
            line_num = pattern.get('line_number', 0)

            if not original_text or not replacement_text:
                print(f"   ❌ Line {line_num}: Missing original or replacement text")
                return False

            search = document.createSearchDescriptor()
            search.setPropertyValue("SearchRegularExpression", False)
            search.setPropertyValue("SearchCaseSensitive", False)

            # Normalize tabs/whitespace in original_text for matching
            normalized = re.sub(r'[\t]+', ' ', original_text).strip()

            for search_str in [original_text, normalized]:
                search.setSearchString(search_str)
                found = document.findFirst(search)
                if found:
                    found.setString(replacement_text)
                    print(f"   ✓ Line {line_num}: Applied '{pattern.get('title', '')[:50]}'")
                    return True

            # Try regex: replace whitespace runs with flexible whitespace match
            regex_pattern = re.sub(r'\s+', '\\\\s+', re.escape(normalized))
            search.setPropertyValue("SearchRegularExpression", True)
            search.setSearchString(regex_pattern)
            found = document.findFirst(search)
            if found:
                search.setPropertyValue("SearchRegularExpression", False)
                found.setString(replacement_text)
                print(f"   ✓ Line {line_num}: Applied via regex '{pattern.get('title', '')[:50]}'")
                return True
            search.setPropertyValue("SearchRegularExpression", False)

            # Fallback: first 8 words of normalized text
            words = normalized.split()[:8]
            search_phrase = ' '.join(words)
            search.setSearchString(search_phrase)
            found = document.findFirst(search)
            if found:
                cursor = found.getText().createTextCursorByRange(found)
                cursor.goRight(len(original_text), True)
                cursor.setString(replacement_text)
                print(f"   ✓ Line {line_num}: Applied via partial match '{search_phrase[:40]}...'")
                return True

            print(f"   ❌ Line {line_num}: Text not found — '{original_text[:60]}...'")
            return False
        except Exception as e:
            print(f"   ❌ Pattern error: {e}")
            return False

    def create_redlined_document(self, instructions, original_path, base_name, mapping_path):
        if not instructions or not instructions.get('patterns'):
            print("   ⚠️ No changes recommended — document is adequate")
            return {}

        try:
            print("   🔄 Starting LibreOffice for redlining...")
            os.system("pkill -f 'soffice'")
            time.sleep(2)

            local_context = uno.getComponentContext()
            resolver = local_context.ServiceManager.createInstanceWithContext(
                "com.sun.star.bridge.UnoUrlResolver", local_context)

            try:
                context = resolver.resolve("uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")
            except:
                subprocess.Popen([
                    'libreoffice', '--headless', '--invisible', '--nocrashreport',
                    '--nodefault', '--nolockcheck', '--nologo', '--norestore',
                    '--accept=socket,host=localhost,port=2002;urp;'
                ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                time.sleep(3)
                context = resolver.resolve("uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")

            desktop = context.ServiceManager.createInstanceWithContext(
                "com.sun.star.frame.Desktop", context)

            file_url = uno.systemPathToFileUrl(os.path.abspath(original_path))
            properties = (PropertyValue("Hidden", 0, True, 0),)
            document = desktop.loadComponentFromURL(file_url, "_blank", 0, properties)

            # Enable track changes
            document.setPropertyValue("RecordChanges", True)
            print(f"   ✓ Track changes enabled")

            # Apply patterns
            applied = 0
            for pattern in instructions['patterns']:
                title = pattern.get('title', 'Unknown')
                print(f"   🔄 Applying: {title}")
                if self.apply_pattern_to_document(document, pattern):
                    applied += 1

            print(f"   📊 Applied {applied}/{len(instructions['patterns'])} patterns")

            # Post-processing: deterministic rules applied directly
            search = document.createSearchDescriptor()
            search.setPropertyValue("SearchRegularExpression", False)
            search.setPropertyValue("SearchCaseSensitive", False)
            search.setPropertyValue("SearchWords", False)

            # Rule: "attorney's fees" → "reasonable attorney's fees"
            for phrase in ["attorney's fees", "attorneys' fees", "attorneys fees"]:
                search.setSearchString(phrase)
                found = document.findFirst(search)
                while found:
                    if "reasonable" not in found.getString().lower():
                        ctx = found.getText()
                        cursor = ctx.createTextCursorByRange(found)
                        cursor.goLeft(12, True)
                        preceding = cursor.getString().lower()
                        if "reasonable" not in preceding:
                            found.setString(f"reasonable {found.getString()}")
                            print(f"   ✓ Post-processing: added 'reasonable' before '{phrase}'")
                    found = document.findNext(found.getEnd(), search)

            save_props = (PropertyValue("FilterName", 0, "MS Word 2007 XML", 0),)
            outputs = {}

            # 1. Redacted + Redlined (PII removed, track changes visible)
            p = os.path.join(OUTPUT_DIR, f"{base_name}_Redacted_Redlined.docx")
            document.storeAsURL(uno.systemPathToFileUrl(os.path.abspath(p)), save_props)
            outputs['redacted_redlined'] = p
            print(f"   ✅ Redacted+Redlined: {os.path.basename(p)}")

            # 2. Reconstructed + Redlined (PII restored, track changes visible)
            p2 = os.path.join(OUTPUT_DIR, f"{base_name}_Reconstructed_Redlined.docx")
            shutil.copy2(p, p2)
            self.restore_pii_in_docx(p2, mapping_path)
            outputs['reconstructed_redlined'] = p2
            print(f"   ✅ Reconstructed+Redlined: {os.path.basename(p2)}")

            # Accept all changes for clean versions
            try:
                redlines = document.getRedlines()
                for i in range(redlines.getCount()):
                    redlines.getByIndex(0).accept()
            except Exception as e:
                print(f"   ⚠️ Error accepting changes: {e}")

            # 3. Redacted + Clean (PII removed, changes accepted)
            p3 = os.path.join(OUTPUT_DIR, f"{base_name}_Redacted_Clean.docx")
            document.storeAsURL(uno.systemPathToFileUrl(os.path.abspath(p3)), save_props)
            outputs['redacted_clean'] = p3
            print(f"   ✅ Redacted+Clean: {os.path.basename(p3)}")

            # 4. Reconstructed + Clean (PII restored, changes accepted)
            p4 = os.path.join(OUTPUT_DIR, f"{base_name}_Reconstructed_Clean.docx")
            shutil.copy2(p3, p4)
            self.restore_pii_in_docx(p4, mapping_path)
            outputs['reconstructed_clean'] = p4
            print(f"   ✅ Reconstructed+Clean: {os.path.basename(p4)}")

            document.close(True)
            return outputs

        except Exception as e:
            print(f"   ❌ LibreOffice redlining failed: {e}")
            return {}
        finally:
            time.sleep(1)
            os.system("pkill -f 'soffice'")

    def process(self, input_path):
        print(f"\n{'='*60}")
        print(f"🔄 Smart Attorney v2: {os.path.basename(input_path)}")
        print(f"{'='*60}")

        # Step 1: File conversion
        print("\n[Step 1] Converting to DOCX...")
        docx_path = self.convert_to_docx(input_path)
        if not docx_path:
            print("❌ Could not convert to DOCX")
            return

        base_name = os.path.splitext(os.path.basename(docx_path))[0]

        # Step 2: Redaction with mapping
        print("\n[Step 2] Redacting PII...")
        redacted_path, mapping_path = self.redact_docx_pii(docx_path, base_name)
        if not redacted_path:
            print("❌ PII redaction failed")
            return

        # Extract text from redacted doc for AI
        print("\n[Step 3] AI attorney assessment...")
        text_content = self.convert_with_libreoffice(redacted_path)
        if not text_content:
            print("❌ Text extraction failed")
            return

        analysis_file, instructions = self.smart_attorney_analysis(text_content, base_name)
        if not instructions:
            print("❌ AI analysis failed")
            return

        # Check if AI says no changes needed
        assessment = instructions.get('document_assessment', {})
        if assessment.get('changes_needed') == 'none' or not instructions.get('patterns'):
            print("\n✅ Document is adequate — no changes recommended")
            shutil.copy2(docx_path, os.path.join(OUTPUT_DIR, f"{base_name}_Original.docx"))
            print(f"\n📁 Output: {OUTPUT_DIR}")
            print(f"   1. Original: {base_name}_Original.docx")
            print(f"   2. Analysis: {os.path.basename(analysis_file)}")
            return

        # Step 4: Redlining + Output creation
        print("\n[Step 4] Applying redlines and creating outputs...")
        outputs = self.create_redlined_document(instructions, redacted_path, base_name, mapping_path)

        # Step 5: Save original + mapping
        print("\n[Step 5] Saving final outputs...")
        shutil.copy2(docx_path, os.path.join(OUTPUT_DIR, f"{base_name}_Original.docx"))
        shutil.copy2(mapping_path, os.path.join(OUTPUT_DIR, f"{base_name}_Mapping.json"))

        print(f"\n{'='*60}")
        print(f"✅ COMPLETE — Output: {OUTPUT_DIR}")
        print(f"{'='*60}")
        print(f"   1. Original:               {base_name}_Original.docx")
        print(f"   2. Analysis:               {os.path.basename(analysis_file)}")
        print(f"   3. Redacted+Redlined:      {os.path.basename(outputs.get('redacted_redlined', 'FAILED'))}")
        print(f"   4. Redacted+Clean:         {os.path.basename(outputs.get('redacted_clean', 'FAILED'))}")
        print(f"   5. Reconstructed+Redlined: {os.path.basename(outputs.get('reconstructed_redlined', 'FAILED'))}")
        print(f"   6. Reconstructed+Clean:    {os.path.basename(outputs.get('reconstructed_clean', 'FAILED'))}")
        print(f"   7. Mapping:                {base_name}_Mapping.json")
        print(f"\n🎯 Patterns applied: {len(instructions.get('patterns', []))}")


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 smart_attorney_system_backup.py <input_file>")
        print("Supported formats: .docx, .doc, .pdf, .txt, .mhtml")
        sys.exit(1)

    input_file = sys.argv[1]
    if not os.path.exists(input_file):
        print(f"Error: File not found — {input_file}")
        sys.exit(1)

    system = SmartAttorneySystem()
    system.process(input_file)


if __name__ == "__main__":
    main()
