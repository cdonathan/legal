#!/usr/bin/env python3
"""
Institutional Standards Redliner - Uses gap analysis to make targeted changes
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class InstitutionalRedliner:
    def __init__(self):
        self.personal_info = {}
        self.gap_analysis = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def analyze_institutional_gaps(self, redacted_text):
        """Analyze what institutional requirements are missing"""
        if not self.openai_client:
            return {"missing_requirements": [2, 3, 4, 5, 6, 7, 9]}
        
        prompt = f"""Analyze this NDA against institutional standards and identify what's missing.

NDA TEXT:
{redacted_text}

INSTITUTIONAL REQUIREMENTS:
1. Clear definition of Confidential Information
2. Standard legal exceptions (public, third party, independent, legally required)
3. Advisor disclosure rights (attorneys, accountants, lenders, advisors)
4. Seller liability limitation ("no representation or warranty")
5. Injunctive relief clause ("irreparable harm")
6. No transaction obligation clause
7. Return/destruction of materials clause
8. Governing law and jurisdiction
9. Standard contract boilerplate
10. Broker protection

Return ONLY JSON:
{{
  "missing_requirements": [2, 3, 4, 5, 6, 7, 9],
  "priority_gaps": [2, 3, 7]
}}

List requirement numbers that are missing or inadequate."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Gap analysis error: {e}")
            return {"missing_requirements": [2, 3, 4, 5, 6, 7, 9], "priority_gaps": [2, 3, 7]}
    
    def get_targeted_changes(self, redacted_text, missing_requirements):
        """Get specific changes to address institutional gaps"""
        if not self.openai_client:
            return []
        
        req_descriptions = {
            2: "Standard legal exceptions (public, third party, independent, legally required)",
            3: "Advisor disclosure rights (attorneys, accountants, lenders, advisors)",
            4: "Seller liability limitation (no representation or warranty)",
            5: "Injunctive relief clause (irreparable harm)",
            6: "No transaction obligation clause",
            7: "Return/destruction of materials clause",
            9: "Standard contract boilerplate"
        }
        
        missing_desc = [req_descriptions.get(req, f"Requirement {req}") for req in missing_requirements]
        
        prompt = f"""Make targeted changes to address these missing institutional requirements:

MISSING REQUIREMENTS:
{chr(10).join([f"- {desc}" for desc in missing_desc])}

NDA TEXT:
{redacted_text}

INSTRUCTIONS:
- Make ONLY the changes needed to address the missing requirements
- Add missing clauses where appropriate
- Modify existing language only if necessary
- Keep changes minimal but ensure institutional compliance

Return ONLY JSON array:
[
  {{
    "type": "insert_after_paragraph",
    "paragraph_number": 5,
    "text": "exact text to insert",
    "addresses_requirement": 3,
    "reason": "adds advisor disclosure rights"
  }},
  {{
    "type": "replace",
    "find": "exact text to find",
    "replace": "exact replacement",
    "addresses_requirement": 4,
    "reason": "adds liability limitation"
  }}
]"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Targeted changes error: {e}")
            return []
    
    def create_redlined_version(self, original_doc_path, changes, output_path):
        """Create redlined version with institutional changes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("INSTITUTIONAL STANDARDS REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run(f"🟢 Green underline = Institutional additions | Total: {len(changes)} targeted changes")
        redlined_doc.add_paragraph("=" * 80)
        
        # Process paragraphs
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for change in changes:
                if change.get('type') == 'replace' and change['find'] in para_text:
                    parts = para_text.split(change['find'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Show deletion
                    del_run = new_para.add_run(change['find'])
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Show replacement
                    ins_run = new_para.add_run(change['replace'])
                    ins_run.underline = True
                    ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    break
            
            if para_text:
                new_para.add_run(para_text)
            
            # Add insertions after this paragraph
            for change in changes:
                if (change.get('type') == 'insert_after_paragraph' and 
                    change.get('paragraph_number') == paragraph_index):
                    
                    insert_para = redlined_doc.add_paragraph()
                    
                    label_run = insert_para.add_run(f"[INSTITUTIONAL ADDITION] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(0, 100, 0)
                    
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    reason_para = redlined_doc.add_paragraph()
                    reason_run = reason_para.add_run(f"REASON: {change['reason']}")
                    reason_run.italic = True
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)
            
            paragraph_index += 1
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, changes, output_path):
        """Create clean version with institutional changes applied"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for change in changes:
                if change.get('type') == 'replace' and change['find'] in para_text:
                    para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
            
            # Add insertions
            for change in changes:
                if (change.get('type') == 'insert_after_paragraph' and 
                    change.get('paragraph_number') == paragraph_index):
                    insert_para = clean_doc.add_paragraph()
                    insert_para.add_run(change['text'])
            
            paragraph_index += 1
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Process NDA with institutional standards approach"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: Analyze gaps
        print("[██░░░] Step 2: Analyzing institutional gaps...")
        self.gap_analysis = self.analyze_institutional_gaps(redacted_text)
        missing_count = len(self.gap_analysis.get('missing_requirements', []))
        print(f"   ✓ Found {missing_count} missing institutional requirements")
        
        # Step 3: Get targeted changes
        print("[███░░] Step 3: Getting targeted changes...")
        changes = self.get_targeted_changes(redacted_text, self.gap_analysis.get('missing_requirements', []))
        print(f"   ✓ Generated {len(changes)} targeted changes")
        
        # Step 4: Create versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[████░] Step 4: Creating redlined version...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_institutional_redlined.docx"
        self.create_redlined_version(input_path, changes, redlined_path)
        
        print("[█████] Step 5: Creating clean version...")
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_institutional_clean-version.docx"
        self.create_clean_version(input_path, changes, clean_path)
        
        # Step 6: Restore personal info
        print("[█████] Step 6: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ INSTITUTIONAL REDLINING COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Addressed {missing_count} institutional gaps with {len(changes)} changes")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 institutional_redliner.py input.docx")
        sys.exit(1)
    
    processor = InstitutionalRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
