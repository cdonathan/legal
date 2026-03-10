#!/usr/bin/env python3
"""
Dual Output NDA Redlining System
Creates both redlined version (with markup) and clean version (changes applied)
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class DualOutputRedliner:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_ai_changes(self, redacted_text):
        """Get AI changes using Golden NDA and Clause Library standards"""
        if not self.openai_client:
            return {"changes": []}
        
        # Load reference materials
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        with open('/home/cliff/redact/redline_project/nda_clause_library.md', 'r') as f:
            clause_library = f.read()
        
        prompt = f"""You are an experienced attorney redlining an NDA. Make MINIMAL inline edits to align with institutional standards.

GOLDEN NDA STANDARD:
{golden_nda[:1500]}...

CLAUSE LIBRARY:
{clause_library[:1500]}...

NDA TO REDLINE:
{redacted_text[:2000]}...

REDLINING APPROACH:
1. Make inline word/phrase replacements within existing sentences
2. Capitalize defined terms consistently
3. Standardize legal language to match Golden NDA
4. Fix awkward phrasing
5. Only add new clauses if critical protections are missing

Return ONLY JSON:
{{
  "changes": [
    {{
      "type": "replace",
      "paragraph": 1,
      "find": "exact text to find",
      "replace": "exact replacement text",
      "reason": "why this change aligns with standards"
    }}
  ]
}}

Focus on inline edits that match the Golden NDA standards."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"AI error: {e}")
            return {"changes": []}

    
    def create_redlined_version(self, original_doc_path, changes, output_path):
        """Create redlined version with track changes markup"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("REDLINED NDA - TRACK CHANGES")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Deletions | 🟢 Green underline = Insertions")
        redlined_doc.add_paragraph("=" * 80)
        
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            
            # Check for replacements
            para_text = para.text
            for change in changes.get('changes', []):
                if change['type'] == 'replace' and change.get('paragraph') == paragraph_index:
                    if change['find'] in para_text:
                        parts = para_text.split(change['find'])
                        
                        if parts[0]:
                            new_para.add_run(parts[0])
                        
                        # Deleted text
                        del_run = new_para.add_run(change['find'])
                        del_run.font.strike = True
                        del_run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        # Inserted text
                        ins_run = new_para.add_run(change['replace'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
                        
                        if len(parts) > 1 and parts[1]:
                            new_para.add_run(parts[1])
                        
                        para_text = ""
            
            if para_text:
                new_para.add_run(para_text)
            
            # Add insertions
            for change in changes.get('changes', []):
                if change['type'] == 'insert' and change.get('after_paragraph') == paragraph_index:
                    insert_para = redlined_doc.add_paragraph()
                    
                    label_run = insert_para.add_run("[INSERTION] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
            
            paragraph_index += 1
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, changes, output_path):
        """Create clean version with all changes applied, no markup"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            
            # Apply replacements
            para_text = para.text
            for change in changes.get('changes', []):
                if change['type'] == 'replace' and change.get('paragraph') == paragraph_index:
                    if change['find'] in para_text:
                        para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
            
            # Add insertions (clean, no markup)
            for change in changes.get('changes', []):
                if change['type'] == 'insert' and change.get('after_paragraph') == paragraph_index:
                    insert_para = clean_doc.add_paragraph()
                    insert_para.add_run(change['text'])
            
            paragraph_index += 1
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Process NDA and create both versions"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: AI redline
        print("[██░░░] Step 2: AI redlining...")
        changes = self.get_ai_changes(redacted_text)
        print(f"   ✓ {len(changes.get('changes', []))} changes suggested")
        
        # Step 3: Create both versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[███░░] Step 3: Creating redlined version...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_redlined.docx"
        self.create_redlined_version(input_path, changes, redlined_path)
        
        print("[████░] Step 4: Creating clean version...")
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_clean-version.docx"
        self.create_clean_version(input_path, changes, clean_path)
        
        # Step 5: Restore personal info in both
        print("[█████] Step 5: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 dual_output_redliner.py input.docx")
        sys.exit(1)
    
    processor = DualOutputRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
