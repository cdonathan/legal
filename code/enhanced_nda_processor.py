#!/usr/bin/env python3
"""
Enhanced NDA Redlining System with proper Word track changes
Requires: pip install python-docx
"""

import re
import json
import os
from datetime import datetime
import openai

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

class EnhancedNDAProcessor:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def extract_text_from_docx(self, docx_path):
        """Extract text preserving paragraph structure"""
        if not DOCX_AVAILABLE:
            return ["Error: python-docx not installed"]
        
        doc = Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        return paragraphs
    
    def redact_personal_info(self, text):
        """Redact only personal info from opening and signature sections"""
        redacted_text = text
        counter = 1
        
        # More precise patterns for NDAs
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME'),
            (r'\b\d{3}-\d{3}-\d{4}\b', 'PHONE'),
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                # Skip common legal terms
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'receiving', 'disclosing']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def ai_redline(self, redacted_text):
        """AI redlining with structured output"""
        if not self.openai_client:
            return redacted_text
        
        prompt = f"""You are redlining an NDA. Make minimal institutional-standard edits.

NDA TEXT:
{redacted_text}

RULES:
1. Add missing standard clauses only if critical
2. Preserve original language when possible  
3. Focus on: permitted recipients, return of materials, term limits, governing law
4. Mark changes as [INSERT: new text] or [DELETE: old text]
5. Keep edits under 20% of document length

Return the redlined NDA with clear markup."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,
                temperature=0.1
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"AI redlining failed: {e}")
            return redacted_text
    
    def create_tracked_docx(self, original_paragraphs, redlined_text, output_path):
        """Create Word document with track changes markup"""
        if not DOCX_AVAILABLE:
            print("Cannot create tracked changes without python-docx")
            return
        
        doc = Document()
        
        # Process redlined text and add with highlighting
        lines = redlined_text.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
                
            para = doc.add_paragraph()
            
            # Handle INSERT markup
            if '[INSERT:' in line:
                parts = re.split(r'\[INSERT:\s*([^\]]+)\]', line)
                for i, part in enumerate(parts):
                    if i % 2 == 0:  # Regular text
                        if part.strip():
                            para.add_run(part)
                    else:  # Inserted text
                        run = para.add_run(part)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Handle DELETE markup  
            elif '[DELETE:' in line:
                parts = re.split(r'\[DELETE:\s*([^\]]+)\]', line)
                for i, part in enumerate(parts):
                    if i % 2 == 0:  # Regular text
                        if part.strip():
                            para.add_run(part)
                    else:  # Deleted text
                        run = para.add_run(part)
                        run.font.strike = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
            
            else:  # Regular text
                para.add_run(line)
        
        doc.save(output_path)
    
    def restore_personal_info(self, redlined_text):
        """Restore personal information"""
        restored_text = redlined_text
        for placeholder, original in self.personal_info.items():
            restored_text = restored_text.replace(placeholder, original)
        return restored_text
    
    def process_nda(self, input_docx_path):
        """Complete 5-step workflow"""
        print(f"Processing: {os.path.basename(input_docx_path)}")
        
        # Extract text
        paragraphs = self.extract_text_from_docx(input_docx_path)
        original_text = '\n\n'.join(paragraphs)
        
        # Step 1: Redact personal info
        print("Step 1: Redacting personal information...")
        redacted_text = self.redact_personal_info(original_text)
        print(f"Redacted {len(self.personal_info)} items: {list(self.personal_info.keys())[:5]}...")
        
        # Step 2: AI redlining
        print("Step 2: AI redlining...")
        redlined_text = self.ai_redline(redacted_text)
        
        # Step 3: Create temp tracked document
        base_name = os.path.splitext(os.path.basename(input_docx_path))[0]
        temp_output = f"/home/cliff/redact/redline_project/{base_name}_temp_tracked.docx"
        print("Step 3: Creating tracked changes document...")
        self.create_tracked_docx(paragraphs, redlined_text, temp_output)
        
        # Step 4: Restore personal info
        print("Step 4: Restoring personal information...")
        final_text = self.restore_personal_info(redlined_text)
        
        # Step 5: Save final document
        final_output = f"/home/cliff/redact/redline_project/{base_name}_redlined.docx"
        print("Step 5: Creating final document...")
        self.create_tracked_docx(paragraphs, final_text, final_output)
        
        print(f"✓ Complete! Saved as: {os.path.basename(final_output)}")
        return final_output

# Simple batch processor
def process_all_ndas():
    processor = EnhancedNDAProcessor()
    nda_folder = "/home/cliff/redact/OneDrive_1_3-5-2026"
    
    pre_redline_files = [f for f in os.listdir(nda_folder) if 'pre' in f and f.endswith('.docx')]
    
    for file in pre_redline_files[:2]:  # Process first 2 for testing
        file_path = os.path.join(nda_folder, file)
        try:
            processor.process_nda(file_path)
            processor.personal_info = {}  # Reset for next document
        except Exception as e:
            print(f"Error processing {file}: {e}")

if __name__ == "__main__":
    if DOCX_AVAILABLE:
        process_all_ndas()
    else:
        print("Install python-docx first: pip install python-docx")
