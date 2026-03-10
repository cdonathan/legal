#!/usr/bin/env python3
"""
Problem-Solving Redliner - AI addresses specific problems identified in Sample 2
"""

import openai
import re
import os
import json
from docx import Document
from docx.shared import RGBColor

class ProblemSolvingRedliner:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_problem_solving_changes(self, redacted_text):
        """Get specific changes to address identified problems"""
        if not self.openai_client:
            return []
        
        # Load the problems analysis
        with open('/home/cliff/redact/redline_project/Sample_2_Problems_Analysis.md', 'r') as f:
            problems_analysis = f.read()
        
        # Load Golden NDA for reference
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        prompt = f"""You are an attorney fixing the specific problems identified in Sample 2 NDA.

GOLDEN NDA STANDARD:
{golden_nda}

PROBLEMS IDENTIFIED:
{problems_analysis}

SAMPLE 2 NDA TO FIX:
{redacted_text}

TASK: Generate specific changes to address each identified problem. Focus on FIXING THE PROBLEMS, not adding generic clauses.

For each change, specify:
- What problem it solves (reference the analysis)
- Exact text to find and replace OR where to insert new text
- The fix that aligns with Golden NDA standard

Return ONLY JSON array:
[
  {{
    "problem_addressed": "Definition of Confidential Information",
    "type": "replace",
    "find": "Informational Materials on the property such as financial information",
    "replace": "Confidential Information means any non-public financial, operational, legal, strategic, or business information relating to the Property that is disclosed to the Potential Purchaser, including without limitation: financial statements, rent rolls, leases, operating statements, tenant information, legal documentation, property information, and transaction materials",
    "reason": "Fixes vague definition by providing comprehensive scope like Golden NDA"
  }},
  {{
    "problem_addressed": "Missing Purpose Clause", 
    "type": "insert_after_paragraph",
    "paragraph_number": 1,
    "text": "The Broker is willing to disclose certain confidential information to the Potential Purchaser solely for the purpose of evaluating a potential purchase of the Property (the 'Purpose').",
    "reason": "Adds missing purpose limitation like Golden NDA"
  }}
]

Address ALL the high and medium risk problems identified in the analysis."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Problem-solving changes error: {e}")
            return []
    
    def create_redlined_version(self, original_doc_path, changes, output_path):
        """Create redlined version showing problem fixes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("PROBLEM-SOLVING REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = Problems removed | 🟢 Green underline = Problem fixes | Total: {len(changes)} targeted fixes")
        redlined_doc.add_paragraph("=" * 80)
        
        # Process paragraphs
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for change in changes:
                if change.get('type') == 'replace' and change['find'] in para_text:
                    parts = para_text.split(change['find'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Show problem being removed
                    del_run = new_para.add_run(change['find'])
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Show fix
                    ins_run = new_para.add_run(change['replace'])
                    ins_run.underline = True
                    ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    break
            
            if para_text:
                new_para.add_run(para_text)
            
            # Add problem fixes after this paragraph
            for change in changes:
                if (change.get('type') == 'insert_after_paragraph' and 
                    change.get('paragraph_number') == paragraph_index):
                    
                    insert_para = redlined_doc.add_paragraph()
                    
                    label_run = insert_para.add_run(f"[PROBLEM FIX: {change['problem_addressed']}] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(0, 100, 0)
                    
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    reason_para = redlined_doc.add_paragraph()
                    reason_run = reason_para.add_run(f"REASON: {change['reason']}")
                    reason_run.italic = True
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)
            
            paragraph_index += 1
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, changes, output_path):
        """Create clean version with problems fixed"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply replacements
            for change in changes:
                if change.get('type') == 'replace' and change['find'] in para_text:
                    para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
            
            # Add problem fixes
            for change in changes:
                if (change.get('type') == 'insert_after_paragraph' and 
                    change.get('paragraph_number') == paragraph_index):
                    insert_para = clean_doc.add_paragraph()
                    insert_para.add_run(change['text'])
            
            paragraph_index += 1
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Process NDA by fixing identified problems"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: Get problem-solving changes
        print("[██░░░] Step 2: Getting problem-solving changes...")
        changes = self.get_problem_solving_changes(redacted_text)
        print(f"   ✓ Generated {len(changes)} targeted fixes")
        
        # Show what problems are being addressed
        problems_addressed = set(change.get('problem_addressed', 'Unknown') for change in changes)
        print(f"   • Addressing {len(problems_addressed)} identified problems")
        
        # Step 3: Create versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[███░░] Step 3: Creating problem-fix redlined version...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_problem_fix_redlined.docx"
        self.create_redlined_version(input_path, changes, redlined_path)
        
        print("[████░] Step 4: Creating clean fixed version...")
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_problem_fix_clean-version.docx"
        self.create_clean_version(input_path, changes, clean_path)
        
        # Step 5: Restore personal info
        print("[█████] Step 5: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ PROBLEM-SOLVING REDLINING COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Fixed {len(problems_addressed)} identified problems with {len(changes)} changes")
        
        return redlined_path, clean_path

def main():
    import sys
    import os
    
    if len(sys.argv) != 2:
        print("Usage: python3 problem_solving_redliner.py input.docx")
        sys.exit(1)
    
    processor = ProblemSolvingRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
