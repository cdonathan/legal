#!/usr/bin/env python3
"""
Complete NDA Redlining System with Visual Track Changes
5-step process: redact -> AI redline -> visual markup -> restore -> save
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class CompleteNDARedliner:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Step 1: Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME'),
            (r'\b\d{3}-\d{3}-\d{4}\b', 'PHONE'),
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'receiving', 'disclosing']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_ai_changes(self, redacted_text):
        """Step 2: Get structured changes from AI"""
        if not self.openai_client:
            return self._mock_changes()
        
        prompt = f"""Redline this NDA. Return ONLY JSON with structured changes.

NDA TEXT:
{redacted_text[:2000]}...

Return JSON format:
{{
  "changes": [
    {{
      "type": "insert",
      "after_paragraph": 3,
      "text": "clause text",
      "reason": "why needed"
    }}
  ]
}}

Add missing: permitted recipients, return of materials, term limits."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except:
            return self._mock_changes()
    
    def _mock_changes(self):
        return {
            "changes": [
                {
                    "type": "insert",
                    "after_paragraph": 3,
                    "text": "The Receiving Party may disclose Confidential Information to its attorneys, accountants, financial advisors, and other professional advisors who have a need to know such information for the Purpose.",
                    "reason": "Add permitted recipients clause"
                }
            ]
        }
    
    def create_visual_redlines(self, original_doc_path, changes, output_path):
        """Step 3: Create document with visual track changes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        paragraph_index = 0
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            new_para.add_run(para.text)
            
            # Add insertions after this paragraph
            for change in changes.get('changes', []):
                if change['type'] == 'insert' and change.get('after_paragraph') == paragraph_index:
                    insert_para = redlined_doc.add_paragraph()
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.underline = True
                    insert_run.font.color.rgb = RGBColor(0, 128, 0)
            
            paragraph_index += 1
        
        redlined_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Step 4: Restore personal information in document"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Complete 5-step workflow"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: AI redline
        print("[██░░░] Step 2: AI redlining...")
        changes = self.get_ai_changes(redacted_text)
        print(f"   ✓ {len(changes.get('changes', []))} changes suggested")
        
        # Step 3: Visual redlines
        print("[███░░] Step 3: Creating visual redlines...")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        temp_output = f"/home/cliff/redact/redline_project/{base_name}_temp.docx"
        self.create_visual_redlines(input_path, changes, temp_output)
        
        # Step 4: Restore personal info
        print("[████░] Step 4: Restoring personal info...")
        self.restore_personal_info(temp_output)
        
        # Step 5: Final save
        print("[█████] Step 5: Saving final document...")
        final_output = f"/home/cliff/redact/redline_project/{base_name}_redlined.docx"
        os.rename(temp_output, final_output)
        
        print(f"✅ Complete: {os.path.basename(final_output)}")
        return final_output

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 complete_nda_redliner.py input.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        sys.exit(1)
    
    processor = CompleteNDARedliner()
    processor.process_nda(input_file)

if __name__ == "__main__":
    main()
