#!/usr/bin/env python3
"""
Complete 4-Call AI System with Document Creation
Call 1: Problems | Call 2: Prioritization | Call 3: Recommendations | Call 4: Implementation | Step 5: Documents
"""

import os
import sys
import subprocess
import json
import re
import openai
from docx import Document
from docx.shared import RGBColor

class Complete4CallSystem:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = len(self.personal_info) + 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'broker', 'seller']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                if placeholder not in self.personal_info:
                    self.personal_info[placeholder] = match.group()
                    redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                    counter += 1
        
        return redacted_text
    
    def convert_with_libreoffice(self, input_path):
        """Convert document using LibreOffice"""
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            input_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
        
        return None
    
    def ai_call_4_implementation_instructions(self, recommendations_content, redacted_text, base_name):
        """Call 4: Implementation instructions with line numbers"""
        if not self.openai_client:
            return None, []
        
        # Add line numbers to text
        lines = redacted_text.split('\n')
        numbered_lines = []
        line_number = 1
        
        for line in lines:
            if line.strip():
                numbered_lines.append(f"LINE {line_number:03d}: {line.strip()}")
                line_number += 1
        
        line_numbered_text = '\n'.join(numbered_lines)
        
        prompt = f"""You are an attorney providing precise implementation instructions.

RECOMMENDATIONS TO IMPLEMENT:
{recommendations_content}

LINE-NUMBERED NDA:
{line_numbered_text}

TASK: Convert recommendations into precise line-numbered implementation instructions.

Return ONLY JSON array:
[
  {{
    "recommendation": "Add definition of Confidential Information",
    "change_type": "insert_definition_section",
    "line_number": 1,
    "insert_text": "For purposes of this Agreement, 'Confidential Information' means...",
    "reason": "Addresses strategic goal #2 - clearly define confidential information"
  }},
  {{
    "recommendation": "Replace vague term",
    "change_type": "replace_term",
    "line_number": 5,
    "find_text": "Informational Materials",
    "replace_with": "Confidential Information",
    "reason": "Addresses strategic goal #2 - use consistent terminology"
  }}
]

Provide precise line numbers and exact text for each recommendation."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000,
                temperature=0.1
            )
            
            implementation_content = response.choices[0].message.content
            
            # Save implementation instructions
            implementation_file = f"/home/cliff/redact/redline_project/{base_name}_Call4_Implementation.md"
            with open(implementation_file, 'w') as f:
                f.write(f"# Implementation Instructions\n\n")
                f.write("**AI Implementation Output:**\n\n")
                f.write("```json\n")
                f.write(implementation_content)
                f.write("\n```\n\n")
            
            # Parse JSON
            json_start = implementation_content.find('[')
            json_end = implementation_content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = implementation_content[json_start:json_end]
                instructions = json.loads(json_str)
                
                print(f"   ✓ Call 4 - Implementation instructions: {os.path.basename(implementation_file)}")
                return implementation_file, instructions
            else:
                print("   ❌ Could not parse implementation JSON")
                return implementation_file, []
                
        except Exception as e:
            print(f"   ❌ Call 4 error: {e}")
            return None, []
    
    def implement_and_create_documents(self, instructions, original_path, base_name):
        """Step 5: Implement instructions and create redlined + clean documents"""
        # Convert to paragraphs
        text_content = self.convert_with_libreoffice(original_path)
        if not text_content:
            return None, None
        
        # Split into paragraphs
        paragraphs = []
        lines = text_content.split('\n')
        current_para = []
        para_number = 1
        
        for line in lines:
            if line.strip():
                current_para.append(line.strip())
            else:
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append({
                        'paragraph_number': para_number,
                        'original_text': para_text,
                        'current_text': para_text,
                        'changes': []
                    })
                    para_number += 1
                    current_para = []
        
        if current_para:
            para_text = ' '.join(current_para)
            paragraphs.append({
                'paragraph_number': para_number,
                'original_text': para_text,
                'current_text': para_text,
                'changes': []
            })
        
        # Implement instructions
        implemented_count = 0
        
        for inst in instructions:
            try:
                if inst.get('change_type') == 'replace_term':
                    find_text = inst.get('find_text', '')
                    replace_text = inst.get('replace_with', '')
                    
                    for para in paragraphs:
                        if find_text in para['current_text']:
                            new_text = para['current_text'].replace(find_text, replace_text)
                            para['current_text'] = new_text
                            para['changes'].append({
                                'type': 'replace_term',
                                'find': find_text,
                                'replace': replace_text,
                                'reason': inst.get('reason', '')
                            })
                            implemented_count += 1
                            print(f"   ✓ Replaced term: {find_text} → {replace_text}")
                            break
                
                elif inst.get('change_type') == 'insert_definition_section':
                    # Insert definition at beginning
                    insert_text = inst.get('insert_text', '')
                    new_para = {
                        'paragraph_number': None,
                        'original_text': '',
                        'current_text': insert_text,
                        'changes': [{
                            'type': 'insert_definition',
                            'text': insert_text,
                            'reason': inst.get('reason', '')
                        }]
                    }
                    paragraphs.insert(1, new_para)  # Insert after title
                    implemented_count += 1
                    print(f"   ✓ Inserted definition section")
                
                elif inst.get('change_type') == 'insert_clause':
                    line_num = inst.get('line_number', 1)
                    insert_text = inst.get('insert_text', '')
                    
                    if line_num <= len(paragraphs):
                        new_para = {
                            'paragraph_number': None,
                            'original_text': '',
                            'current_text': insert_text,
                            'changes': [{
                                'type': 'insert_clause',
                                'text': insert_text,
                                'reason': inst.get('reason', '')
                            }]
                        }
                        paragraphs.insert(line_num, new_para)
                        implemented_count += 1
                        print(f"   ✓ Inserted clause after line {line_num}")
                    
            except Exception as e:
                print(f"   ❌ Failed to implement {inst.get('recommendation', 'unknown')}: {e}")
        
        # Create redlined document
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_4Call_Redlined.docx"
        doc = Document()
        
        # Header
        header = doc.add_paragraph()
        header_run = header.add_run("4-CALL AI SYSTEM REDLINES")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = Original | 🟢 Green underline = AI changes | Total: {implemented_count} changes")
        doc.add_paragraph("=" * 80)
        
        # Process paragraphs
        for para_data in paragraphs:
            doc_para = doc.add_paragraph()
            
            if para_data['changes']:
                for change in para_data['changes']:
                    if change['type'] == 'replace_term':
                        original_text = para_data['original_text']
                        find_text = change['find']
                        replace_text = change['replace']
                        
                        if find_text in original_text:
                            parts = original_text.split(find_text)
                            
                            if parts[0]:
                                doc_para.add_run(parts[0])
                            
                            # Show deletion
                            del_run = doc_para.add_run(find_text)
                            del_run.font.strike = True
                            del_run.font.color.rgb = RGBColor(255, 0, 0)
                            
                            # Show addition
                            ins_run = doc_para.add_run(replace_text)
                            ins_run.underline = True
                            ins_run.font.color.rgb = RGBColor(0, 128, 0)
                            
                            if len(parts) > 1 and parts[1]:
                                doc_para.add_run(parts[1])
                        else:
                            doc_para.add_run(para_data['current_text'])
                    
                    elif change['type'] in ['insert_definition', 'insert_clause']:
                        # Show insertion
                        label_run = doc_para.add_run("[AI ADDITION] ")
                        label_run.bold = True
                        label_run.font.color.rgb = RGBColor(0, 100, 0)
                        
                        ins_run = doc_para.add_run(change['text'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                if para_data['current_text']:
                    doc_para.add_run(para_data['current_text'])
        
        doc.save(redlined_path)
        
        # Create clean document
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_4Call_Clean.docx"
        clean_doc = Document()
        
        for para_data in paragraphs:
            if para_data['current_text']:
                para = clean_doc.add_paragraph()
                para.add_run(para_data['current_text'])
        
        clean_doc.save(clean_path)
        
        print(f"   ✓ Created redlined document: {os.path.basename(redlined_path)}")
        print(f"   ✓ Created clean document: {os.path.basename(clean_path)}")
        print(f"   ✓ Implemented {implemented_count}/{len(instructions)} instructions")
        
        return redlined_path, clean_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information in document"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
    
    def process_complete_system(self, input_path):
        """Complete 4-call system with document creation"""
        print(f"🔄 Complete 4-Call AI System: {os.path.basename(input_path)}")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        # Convert document
        print("[█░░░░] Step 1: Converting document with LibreOffice...")
        text_content = self.convert_with_libreoffice(input_path)
        if not text_content:
            print("❌ LibreOffice conversion failed")
            return
        
        # Redact personal info
        redacted_text = self.redact_personal_info(text_content)
        print(f"   ✓ Redacted {len(self.personal_info)} personal items")
        
        # For now, use existing recommendations (Call 3 output from previous run)
        recommendations_content = """
## P1 Recommendations (Must Implement)

### Recommendation 3: Definition of Confidential Information
**Strategic Goal:** Clearly define what is and is not confidential  
**Recommendation:** Insert a comprehensive definition of what constitutes "Confidential Information" within the agreement.  
**Justification:** A clear definition helps to avoid misunderstandings about what information is protected, ensuring that both parties are aligned on the scope of confidentiality.
"""
        
        # Call 4: Implementation Instructions
        print("[██░░░] Call 4: Implementation instructions...")
        implementation_file, instructions = self.ai_call_4_implementation_instructions(recommendations_content, redacted_text, base_name)
        
        # Step 5: Implement and create documents
        if instructions:
            print("[███░░] Step 5: Creating redlined and clean documents...")
            redlined_path, clean_path = self.implement_and_create_documents(instructions, input_path, base_name)
            
            # Restore personal info
            if redlined_path:
                self.restore_personal_info(redlined_path)
            if clean_path:
                self.restore_personal_info(clean_path)
        else:
            redlined_path, clean_path = None, None
        
        print("✅ COMPLETE 4-CALL SYSTEM FINISHED!")
        print(f"📋 Call 4 - Implementation: {os.path.basename(implementation_file) if implementation_file else 'FAILED'}")
        print(f"📝 Redlined Document: {os.path.basename(redlined_path) if redlined_path else 'FAILED'}")
        print(f"📄 Clean Document: {os.path.basename(clean_path) if clean_path else 'FAILED'}")
        print(f"🎯 Total implementation instructions: {len(instructions)}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 complete_4call_system.py input.docx")
        sys.exit(1)
    
    system = Complete4CallSystem()
    system.process_complete_system(sys.argv[1])

if __name__ == "__main__":
    main()
