#!/usr/bin/env python3
"""
Analyze the redlined document output
"""

from docx import Document
import json

def analyze_redlined_doc(doc_path):
    """Check what's actually in the redlined document"""
    print(f"Analyzing: {doc_path}")
    
    doc = Document(doc_path)
    
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    green_text_found = False
    underlined_text_found = False
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"\nParagraph {i}: {para.text[:100]}...")
            
            # Check runs for formatting
            for j, run in enumerate(para.runs):
                if run.underline:
                    print(f"  Run {j}: UNDERLINED - {run.text[:50]}...")
                    underlined_text_found = True
                
                if run.font.color.rgb:
                    rgb = run.font.color.rgb
                    if rgb.red == 0 and rgb.green == 128 and rgb.blue == 0:
                        print(f"  Run {j}: GREEN TEXT - {run.text[:50]}...")
                        green_text_found = True
    
    print(f"\nSummary:")
    print(f"Green text found: {green_text_found}")
    print(f"Underlined text found: {underlined_text_found}")
    
    if not green_text_found and not underlined_text_found:
        print("❌ No redline formatting detected!")
        print("The document appears to be unchanged.")

if __name__ == "__main__":
    doc_path = "/home/cliff/redact/redline_project/REDLINE_NDA_Sample5_pre_redline_redlined.docx"
    analyze_redlined_doc(doc_path)
