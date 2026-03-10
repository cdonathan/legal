#!/usr/bin/env python3
"""
Analyze attorney redlines to understand their exact changes
"""

from docx import Document
import difflib

def analyze_attorney_redlines():
    """Compare pre-redline vs attorney redlined version to document every change"""
    
    # Load documents
    pre_redline = Document('/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE - NDA_Sample_4_pre_redline.docx')
    attorney_redlined = Document('/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE - NDA_Sample_4.docx')
    
    # Extract text by paragraphs
    pre_paras = [p.text.strip() for p in pre_redline.paragraphs if p.text.strip()]
    attorney_paras = [p.text.strip() for p in attorney_redlined.paragraphs if p.text.strip()]
    
    changes_found = []
    
    # Compare paragraph by paragraph
    for i, (pre, attorney) in enumerate(zip(pre_paras, attorney_paras)):
        if pre != attorney:
            changes_found.append({
                'paragraph': i + 1,
                'type': 'MODIFIED',
                'original': pre,
                'attorney_version': attorney,
                'change_type': analyze_change_type(pre, attorney)
            })
    
    # Check for added/removed paragraphs
    if len(attorney_paras) > len(pre_paras):
        for i in range(len(pre_paras), len(attorney_paras)):
            changes_found.append({
                'paragraph': i + 1,
                'type': 'ADDED',
                'original': '',
                'attorney_version': attorney_paras[i],
                'change_type': 'PARAGRAPH_ADDITION'
            })
    elif len(pre_paras) > len(attorney_paras):
        for i in range(len(attorney_paras), len(pre_paras)):
            changes_found.append({
                'paragraph': i + 1,
                'type': 'DELETED',
                'original': pre_paras[i],
                'attorney_version': '',
                'change_type': 'PARAGRAPH_DELETION'
            })
    
    return changes_found

def analyze_change_type(original, modified):
    """Determine what type of change was made"""
    if 'XXXX' in modified and 'XXXX' not in original:
        return 'REDACTION'
    elif len(modified) < len(original) * 0.8:
        return 'MAJOR_DELETION'
    elif len(modified) > len(original) * 1.2:
        return 'MAJOR_ADDITION'
    elif original.lower() != modified.lower():
        return 'TEXT_MODIFICATION'
    else:
        return 'FORMATTING_CHANGE'

def create_changes_markdown():
    """Create detailed markdown of all attorney changes"""
    
    changes = analyze_attorney_redlines()
    
    md_content = f"""# Attorney Redline Analysis - NDA Sample 4

**Total Changes Found:** {len(changes)}

## Change Summary by Type

"""
    
    # Count change types
    change_types = {}
    for change in changes:
        change_type = change['change_type']
        change_types[change_type] = change_types.get(change_type, 0) + 1
    
    for change_type, count in change_types.items():
        md_content += f"- **{change_type}:** {count} changes\n"
    
    md_content += "\n---\n\n## Detailed Changes\n\n"
    
    # Document each change
    for i, change in enumerate(changes, 1):
        md_content += f"### Change #{i} - Paragraph {change['paragraph']}\n\n"
        md_content += f"**Type:** {change['change_type']}\n\n"
        
        if change['type'] == 'MODIFIED':
            md_content += f"**Original:**\n```\n{change['original']}\n```\n\n"
            md_content += f"**Attorney Version:**\n```\n{change['attorney_version']}\n```\n\n"
            
            # Show word-level differences
            original_words = change['original'].split()
            attorney_words = change['attorney_version'].split()
            
            differ = difflib.unified_diff(original_words, attorney_words, lineterm='')
            diff_lines = list(differ)
            
            if diff_lines:
                md_content += "**Word-Level Changes:**\n```diff\n"
                for line in diff_lines:
                    md_content += line + "\n"
                md_content += "```\n\n"
        
        elif change['type'] == 'ADDED':
            md_content += f"**Added Text:**\n```\n{change['attorney_version']}\n```\n\n"
        
        elif change['type'] == 'DELETED':
            md_content += f"**Deleted Text:**\n```\n{change['original']}\n```\n\n"
        
        md_content += "---\n\n"
    
    # Add patterns analysis
    md_content += "## Redaction Patterns Observed\n\n"
    
    redaction_examples = []
    for change in changes:
        if change['change_type'] == 'REDACTION':
            # Find what was redacted
            original = change['original']
            modified = change['attorney_version']
            
            # Simple pattern detection
            if 'XXXX' in modified:
                redaction_examples.append({
                    'original_snippet': original[:100] + "..." if len(original) > 100 else original,
                    'redacted_snippet': modified[:100] + "..." if len(modified) > 100 else modified
                })
    
    for i, example in enumerate(redaction_examples[:5], 1):  # Show first 5
        md_content += f"**Redaction Example {i}:**\n"
        md_content += f"- Original: `{example['original_snippet']}`\n"
        md_content += f"- Redacted: `{example['redacted_snippet']}`\n\n"
    
    md_content += "## Key Insights for AI Prompting\n\n"
    md_content += "Based on this analysis, attorneys primarily:\n\n"
    
    if change_types.get('REDACTION', 0) > 0:
        md_content += f"1. **Redacted personal information** ({change_types.get('REDACTION', 0)} instances) - Replace specific details with XXXX\n"
    
    if change_types.get('TEXT_MODIFICATION', 0) > 0:
        md_content += f"2. **Made text modifications** ({change_types.get('TEXT_MODIFICATION', 0)} instances) - Improved language/terminology\n"
    
    if change_types.get('MAJOR_DELETION', 0) > 0:
        md_content += f"3. **Deleted content** ({change_types.get('MAJOR_DELETION', 0)} instances) - Removed unnecessary text\n"
    
    if change_types.get('MAJOR_ADDITION', 0) > 0:
        md_content += f"4. **Added content** ({change_types.get('MAJOR_ADDITION', 0)} instances) - Inserted missing clauses\n"
    
    md_content += "\n**Recommendation:** AI should focus on redaction and minimal text improvements, not major structural changes.\n"
    
    return md_content

def main():
    """Generate the markdown analysis"""
    print("🔍 Analyzing attorney redlines...")
    
    md_content = create_changes_markdown()
    
    output_file = "/home/cliff/redact/redline_project/Attorney_Redline_Analysis_NDA_Sample_4.md"
    
    with open(output_file, 'w') as f:
        f.write(md_content)
    
    print(f"✅ Analysis complete: {output_file}")
    print("📄 Review this file to understand exactly what attorneys changed")

if __name__ == "__main__":
    main()
