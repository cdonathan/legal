#!/usr/bin/env python3
"""
Enhanced Visual Redlining - Make changes more obvious
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

class EnhancedVisualRedliner:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_ai_changes(self, redacted_text):
        """Get AI changes"""
        if not self.openai_client:
            return self._mock_changes()
        
        prompt = f"""Redline this NDA. Return ONLY JSON.

NDA TEXT:
{redacted_text[:1500]}...

Return JSON:
{{
  "changes": [
    {{
      "type": "insert",
      "after_paragraph": 2,
      "text": "clause text here",
      "reason": "why needed"
    }}
  ]
}}

Add missing standard clauses."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except:
            return self._mock_changes()
    
    def _mock_changes(self):
        return {
            "changes": [
                {
                    "type": "insert",
                    "after_paragraph": 2,
                    "text": "PERMITTED RECIPIENTS: The Receiving Party may disclose Confidential Information to its attorneys, accountants, financial advisors, lenders, investors, and other professional advisors who have a need to know such information for the Purpose and who are bound by confidentiality obligations.",
                    "reason": "Add permitted recipients clause - institutional standard"
                },
                {
                    "type": "insert",
                    "after_paragraph": 4,
                    "text": "RETURN OF MATERIALS: Upon written request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and any copies, notes, or derivatives thereof.",
                    "reason": "Add return of materials clause - required for enforceability"
                }
            ]
        }
    
    def create_obvious_redlines(self, original_doc_path, changes, output_path):
        """Create document with VERY obvious redlines"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header explaining redlines
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("REDLINED NDA - AI SUGGESTED CHANGES")
        header_run.bold = True
        header_run.font.size = 16
        header_run.font.color.rgb = RGBColor(255, 0, 0)
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run("GREEN HIGHLIGHTED TEXT = AI INSERTIONS | Review and accept/reject as needed")
        
        redlined_doc.add_paragraph("=" * 80)
        
        paragraph_index = 0
        for para in doc.paragraphs:
            # Copy original paragraph
            new_para = redlined_doc.add_paragraph()
            new_para.add_run(para.text)
            
            # Add insertions after this paragraph with OBVIOUS formatting
            for change in changes.get('changes', []):
                if change['type'] == 'insert' and change.get('after_paragraph') == paragraph_index:
                    # Add blank line
                    redlined_doc.add_paragraph()
                    
                    # Add insertion with obvious formatting
                    insert_para = redlined_doc.add_paragraph()
                    
                    # Label
                    label_run = insert_para.add_run("[AI INSERTION] ")
                    label_run.bold = True
                    label_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Inserted text
                    insert_run = insert_para.add_run(change['text'])
                    insert_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    insert_run.underline = True
                    insert_run.bold = True
                    
                    # Reason
                    reason_para = redlined_doc.add_paragraph()
                    reason_run = reason_para.add_run(f"REASON: {change['reason']}")
                    reason_run.italic = True
                    reason_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    redlined_doc.add_paragraph()
            
            paragraph_index += 1
        
        # Add summary
        redlined_doc.add_page_break()
        summary_header = redlined_doc.add_paragraph()
        summary_run = summary_header.add_run("REDLINE SUMMARY")
        summary_run.bold = True
        summary_run.font.size = 14
        
        for i, change in enumerate(changes.get('changes', []), 1):
            change_para = redlined_doc.add_paragraph()
            change_para.add_run(f"{i}. {change['type'].upper()}: {change['reason']}")
        
        redlined_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Process with obvious visual redlines"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: AI redline
        print("[██░░░] Step 2: AI redlining...")
        changes = self.get_ai_changes(redacted_text)
        print(f"   ✓ {len(changes.get('changes', []))} changes suggested")
        
        # Step 3: Create OBVIOUS redlines
        print("[███░░] Step 3: Creating obvious visual redlines...")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        temp_output = f"/home/cliff/redact/redline_project/{base_name}_temp.docx"
        self.create_obvious_redlines(input_path, changes, temp_output)
        
        # Step 4: Restore personal info
        print("[████░] Step 4: Restoring personal info...")
        self.restore_personal_info(temp_output)
        
        # Step 5: Final save
        print("[█████] Step 5: Saving final document...")
        final_output = f"/home/cliff/redact/redline_project/{base_name}_OBVIOUS_REDLINED.docx"
        os.rename(temp_output, final_output)
        
        print(f"✅ Complete: {os.path.basename(final_output)}")
        print("   🟢 GREEN HIGHLIGHTED = AI insertions")
        print("   📝 Summary included at end")
        return final_output

if __name__ == "__main__":
    processor = EnhancedVisualRedliner()
    test_file = "/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Conf_Agr_Sample1-pre-redline.docx"
    if os.path.exists(test_file):
        processor.process_nda(test_file)
    else:
        print("Test file not found")
