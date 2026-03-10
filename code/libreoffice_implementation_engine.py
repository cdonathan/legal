#!/usr/bin/env python3
"""
LibreOffice Implementation Engine - Follows AI's precise instructions
"""

import os
import sys
import subprocess
import json
import re
from docx import Document
from docx.shared import RGBColor

class LibreOfficeImplementationEngine:
    def __init__(self):
        self.personal_info = {}
        self.original_lines = []
    
    def load_ai_recommendations(self, recommendations_file):
        """Load AI recommendations from markdown file"""
        try:
            with open(recommendations_file, 'r') as f:
                content = f.read()
            
            # Find the JSON array more carefully
            lines = content.split('\n')
            json_lines = []
            in_json = False
            
            for line in lines:
                if line.strip() == '```json' or line.strip().startswith('['):
                    in_json = True
                    if line.strip().startswith('['):
                        json_lines.append(line)
                    continue
                elif line.strip() == '```' and in_json:
                    break
                elif in_json:
                    json_lines.append(line)
            
            if json_lines:
                json_str = '\n'.join(json_lines)
                # Clean up any extra content after the JSON
                if ']' in json_str:
                    json_str = json_str[:json_str.rfind(']') + 1]
                
                recommendations = json.loads(json_str)
                print(f"   ✓ Loaded {len(recommendations)} AI recommendations")
                return recommendations
            else:
                print("   ❌ No JSON array found in recommendations file")
                return []
        except Exception as e:
            print(f"   ❌ Error loading recommendations: {e}")
            # Try to show what we found
            try:
                print(f"   Debug: JSON string was: {json_str[:200]}...")
            except:
                pass
            return []
    
    def load_line_numbered_document(self, line_numbered_file):
        """Load the line-numbered document for reference"""
        try:
            with open(line_numbered_file, 'r') as f:
                content = f.read()
            
            # Extract lines
            lines = content.split('\n')
            numbered_lines = {}
            
            for line in lines:
                if line.startswith('LINE '):
                    try:
                        # Extract line number and content
                        parts = line.split(': ', 1)
                        if len(parts) == 2:
                            line_num_str = parts[0].replace('LINE ', '').strip()
                            line_num = int(line_num_str)
                            line_content = parts[1]
                            numbered_lines[line_num] = line_content
                    except:
                        continue
            
            print(f"   ✓ Loaded {len(numbered_lines)} numbered lines")
            return numbered_lines
        except Exception as e:
            print(f"   ❌ Error loading line-numbered document: {e}")
            return {}
    
    def convert_original_to_lines(self, original_path):
        """Convert original document to line-based format for editing"""
        # Use LibreOffice to convert to text
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            original_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(original_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Split into lines and number them
                lines = content.split('\n')
                numbered_lines = []
                line_number = 1
                
                for line in lines:
                    if line.strip():
                        numbered_lines.append({
                            'line_number': line_number,
                            'original_text': line.strip(),
                            'current_text': line.strip(),
                            'modified': False
                        })
                        line_number += 1
                    else:
                        numbered_lines.append({
                            'line_number': None,
                            'original_text': '',
                            'current_text': '',
                            'modified': False
                        })
                
                print(f"   ✓ Converted to {line_number-1} numbered lines")
                return numbered_lines
        
        return []
    
    def implement_recommendations(self, original_lines, recommendations):
        """Implement AI recommendations on the line-based document"""
        modified_lines = original_lines.copy()
        implemented_count = 0
        
        print(f"   🔄 Implementing {len(recommendations)} recommendations...")
        
        for rec in recommendations:
            try:
                if rec.get('change_type') == 'replace_text_on_line':
                    line_num = rec.get('line_number')
                    find_text = rec.get('find_text', '')
                    replace_text = rec.get('replace_with', '')
                    
                    # Find the line
                    for i, line_data in enumerate(modified_lines):
                        if line_data.get('line_number') == line_num:
                            current_text = line_data['current_text']
                            
                            # Try exact match first
                            if find_text in current_text:
                                new_text = current_text.replace(find_text, replace_text)
                                modified_lines[i]['current_text'] = new_text
                                modified_lines[i]['modified'] = True
                                implemented_count += 1
                                print(f"   ✓ Replaced text on line {line_num}")
                                break
                            else:
                                # Try partial match (first 50 characters)
                                find_start = find_text[:50] if len(find_text) > 50 else find_text
                                if find_start in current_text:
                                    # Replace the whole line
                                    modified_lines[i]['current_text'] = replace_text
                                    modified_lines[i]['modified'] = True
                                    implemented_count += 1
                                    print(f"   ✓ Replaced line {line_num} (partial match)")
                                    break
                                else:
                                    print(f"   ⚠️ Could not find text on line {line_num}")
                
                elif rec.get('change_type') == 'insert_after_line':
                    line_num = rec.get('line_number')
                    insert_text = rec.get('insert_text', '')
                    
                    # Find insertion point
                    for i, line_data in enumerate(modified_lines):
                        if line_data.get('line_number') == line_num:
                            # Insert after this line
                            new_line = {
                                'line_number': None,
                                'original_text': '',
                                'current_text': insert_text,
                                'modified': True,
                                'inserted': True
                            }
                            modified_lines.insert(i + 1, new_line)
                            implemented_count += 1
                            print(f"   ✓ Inserted text after line {line_num}")
                            break
                
                elif rec.get('change_type') == 'insert_at_end':
                    insert_text = rec.get('insert_text', '')
                    new_line = {
                        'line_number': None,
                        'original_text': '',
                        'current_text': insert_text,
                        'modified': True,
                        'inserted': True
                    }
                    modified_lines.append(new_line)
                    implemented_count += 1
                    print(f"   ✓ Inserted text at end")
                    
            except Exception as e:
                print(f"   ❌ Failed to implement {rec.get('problem_addressed', 'unknown')}: {e}")
        
        print(f"   ✓ Successfully implemented {implemented_count}/{len(recommendations)} recommendations")
        return modified_lines, implemented_count
    
    def create_redlined_docx(self, original_lines, modified_lines, output_path):
        """Create redlined Word document showing changes"""
        doc = Document()
        
        # Add header
        header = doc.add_paragraph()
        header_run = header.add_run("LIBREOFFICE AI IMPLEMENTATION - REDLINED")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Original text | 🟢 Green underline = AI changes")
        doc.add_paragraph("=" * 80)
        
        # Process each line
        for i, modified_line in enumerate(modified_lines):
            para = doc.add_paragraph()
            
            if modified_line.get('inserted'):
                # This is an insertion
                label_run = para.add_run("[AI INSERTION] ")
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(0, 100, 0)
                
                insert_run = para.add_run(modified_line['current_text'])
                insert_run.underline = True
                insert_run.font.color.rgb = RGBColor(0, 128, 0)
                
            elif modified_line.get('modified'):
                # This line was modified
                original_text = modified_line['original_text']
                new_text = modified_line['current_text']
                
                if original_text:
                    # Show original (strikethrough)
                    del_run = para.add_run(original_text)
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    para.add_run(" → ")
                
                # Show new text (underlined)
                ins_run = para.add_run(new_text)
                ins_run.underline = True
                ins_run.font.color.rgb = RGBColor(0, 128, 0)
                
            else:
                # Unchanged line
                if modified_line['current_text']:
                    para.add_run(modified_line['current_text'])
        
        doc.save(output_path)
        print(f"   ✓ Created redlined document: {os.path.basename(output_path)}")
    
    def create_clean_docx(self, modified_lines, output_path):
        """Create clean Word document with all changes applied"""
        doc = Document()
        
        # Process each line
        for modified_line in modified_lines:
            if modified_line['current_text']:
                para = doc.add_paragraph()
                para.add_run(modified_line['current_text'])
        
        doc.save(output_path)
        print(f"   ✓ Created clean document: {os.path.basename(output_path)}")
    
    def process_implementation(self, original_path, recommendations_file, line_numbered_file):
        """Complete implementation workflow"""
        print(f"🔄 Implementing AI recommendations for: {os.path.basename(original_path)}")
        
        # Step 1: Load AI recommendations
        print("[█░░░░] Step 1: Loading AI recommendations...")
        recommendations = self.load_ai_recommendations(recommendations_file)
        if not recommendations:
            return None, None
        
        # Step 2: Convert original to line format
        print("[██░░░] Step 2: Converting original document...")
        original_lines = self.convert_original_to_lines(original_path)
        if not original_lines:
            return None, None
        
        # Step 3: Implement recommendations
        print("[███░░] Step 3: Implementing AI recommendations...")
        modified_lines, implemented_count = self.implement_recommendations(original_lines, recommendations)
        
        # Step 4: Create output documents
        print("[████░] Step 4: Creating output documents...")
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_implemented_redlined.docx"
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_implemented_clean.docx"
        
        self.create_redlined_docx(original_lines, modified_lines, redlined_path)
        self.create_clean_docx(modified_lines, clean_path)
        
        print("[█████] Step 5: Complete!")
        print("✅ LIBREOFFICE IMPLEMENTATION COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Implemented {implemented_count}/{len(recommendations)} AI recommendations")
        
        return redlined_path, clean_path

def main():
    if len(sys.argv) != 4:
        print("Usage: python3 libreoffice_implementation_engine.py original.docx recommendations.md line_numbered.txt")
        sys.exit(1)
    
    original_path = sys.argv[1]
    recommendations_file = sys.argv[2]
    line_numbered_file = sys.argv[3]
    
    engine = LibreOfficeImplementationEngine()
    engine.process_implementation(original_path, recommendations_file, line_numbered_file)

if __name__ == "__main__":
    main()
