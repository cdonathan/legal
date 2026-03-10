#!/usr/bin/env python3
"""
Multi-Pass AI Redlining System
Creates virtual working document and makes multiple focused passes like attorneys do
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class MultiPassRedliner:
    def __init__(self):
        self.personal_info = {}
        self.document_context = {}
        self.working_document = ""
        self.all_changes = []
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Step 1: Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def analyze_document_context(self, redacted_text):
        """Step 2: Analyze document for strategic context"""
        if not self.openai_client:
            return self._mock_context()
        
        prompt = f"""Analyze this NDA to understand the strategic context for institutional redlining.

NDA TEXT:
{redacted_text}

CORE REDLINING GOALS:
1. Limit liability for seller/broker
2. Clearly define confidential information  
3. Allow sharing within buyer's organization
4. Control buyer interactions with property/tenants
5. Create enforceable remedies
6. Add modern contract language

Return ONLY JSON:
{{
  "transaction_type": "Real estate acquisition/lease/etc",
  "current_weaknesses": ["What makes this NDA legally weak"],
  "missing_protections": ["What institutional protections are missing"],
  "broker_involvement": "Yes/No - is there broker protection needed",
  "liability_risks": ["What liability risks exist for seller"],
  "enforceability_issues": ["What makes this hard to enforce"]
}}"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=800,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Context analysis error: {e}")
            return self._mock_context()
    
    def _mock_context(self):
        return {
            "transaction_type": "Real estate confidentiality agreement",
            "current_weaknesses": ["Vague confidentiality definition", "No advisor disclosure rights"],
            "missing_protections": ["Liability limitation", "Injunctive relief"],
            "broker_involvement": "Yes",
            "liability_risks": ["Information accuracy claims"],
            "enforceability_issues": ["No governing law", "Weak remedies"]
        }
    
    def make_pass(self, pass_number, focus_area, working_doc):
        """Make a focused pass on the working document"""
        if not self.openai_client:
            return working_doc, []
        
        pass_prompts = {
            1: {
                "focus": "DEFINITIONS AND SCOPE",
                "instruction": """Focus ONLY on defining confidential information clearly and adding standard exceptions.

GOALS:
- Add formal definition of Confidential Information
- Include financial data, strategy, customer lists, transaction discussions
- Add standard exceptions (already known, public, third party, independently developed, legally required)

Make minimal inline edits. Do NOT add new sections."""
            },
            2: {
                "focus": "ADVISOR DISCLOSURE RIGHTS", 
                "instruction": """Focus ONLY on allowing disclosure to advisors and internal teams.

GOALS:
- Allow sharing with investors, employees, attorneys, accountants, lenders, advisors
- Add language: "Recipient may disclose to its attorneys, accountants, financial advisors who need to know"

Make minimal inline edits within existing paragraphs."""
            },
            3: {
                "focus": "LIABILITY PROTECTION",
                "instruction": """Focus ONLY on limiting seller/broker liability.

GOALS:
- Add "no representation or warranty as to accuracy"
- Limit seller liability for information quality
- Protect broker from circumvention

Make minimal inline edits."""
            },
            4: {
                "focus": "ENFORCEABILITY",
                "instruction": """Focus ONLY on making the NDA legally enforceable.

GOALS:
- Add injunctive relief language
- Add governing law and jurisdiction
- Add "no obligation to transact" clause
- Add document return/destruction requirements

Make minimal inline edits."""
            }
        }
        
        if pass_number not in pass_prompts:
            return working_doc, []
        
        pass_info = pass_prompts[pass_number]
        
        prompt = f"""PASS {pass_number}: {pass_info['focus']}

DOCUMENT CONTEXT:
Transaction: {self.document_context.get('transaction_type', 'Real estate')}
Weaknesses: {self.document_context.get('current_weaknesses', [])}
Missing: {self.document_context.get('missing_protections', [])}

CURRENT WORKING DOCUMENT:
{working_doc}

PASS INSTRUCTIONS:
{pass_info['instruction']}

CRITICAL RULES:
- Make ONLY inline word/phrase replacements within existing sentences
- Do NOT add new paragraphs or sections
- Focus ONLY on this pass's specific goal
- Preserve document structure

Return ONLY JSON:
{{
  "changes": [
    {{
      "type": "replace",
      "find": "exact text to replace",
      "replace": "exact replacement text", 
      "reason": "why this supports the pass goal"
    }}
  ]
}}"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            changes = json.loads(json_str)
            
            # Apply changes to working document
            updated_doc = working_doc
            for change in changes.get('changes', []):
                if change['find'] in updated_doc:
                    updated_doc = updated_doc.replace(change['find'], change['replace'])
            
            return updated_doc, changes.get('changes', [])
            
        except Exception as e:
            print(f"Pass {pass_number} error: {e}")
            return working_doc, []
    
    def create_redlined_version(self, original_doc_path, all_changes, output_path):
        """Create redlined version showing all changes from all passes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("MULTI-PASS AI REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Deletions | 🟢 Green underline = Insertions")
        
        # Show pass summary
        summary = redlined_doc.add_paragraph()
        summary_run = summary.add_run(f"Applied {len(all_changes)} changes across 4 focused passes")
        summary_run.italic = True
        
        redlined_doc.add_paragraph("=" * 80)
        
        # Process each paragraph with all changes
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply all changes
            for change in all_changes:
                if change['find'] in para_text:
                    parts = para_text.split(change['find'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Deleted text
                    del_run = new_para.add_run(change['find'])
                    del_run.font.strike = True
                    del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Inserted text
                    ins_run = new_para.add_run(change['replace'])
                    ins_run.underline = True
                    ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    break
            
            if para_text:
                new_para.add_run(para_text)
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, all_changes, output_path):
        """Create clean version with all changes applied"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply all changes
            for change in all_changes:
                if change['find'] in para_text:
                    para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Complete multi-pass workflow"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: Analyze context
        print("[██░░░░░] Step 2: Analyzing strategic context...")
        self.document_context = self.analyze_document_context(redacted_text)
        print(f"   ✓ Transaction: {self.document_context['transaction_type']}")
        
        # Step 3: Initialize working document
        print("[███░░░░] Step 3: Creating virtual working document...")
        self.working_document = redacted_text
        
        # Step 4-7: Multi-pass redlining
        pass_names = ["Definitions", "Advisor Rights", "Liability Protection", "Enforceability"]
        
        for pass_num in range(1, 5):
            print(f"[{'█' * (3 + pass_num)}{'░' * (7 - (3 + pass_num))}] Pass {pass_num}: {pass_names[pass_num-1]}...")
            
            self.working_document, pass_changes = self.make_pass(pass_num, pass_names[pass_num-1], self.working_document)
            self.all_changes.extend(pass_changes)
            
            print(f"   ✓ Made {len(pass_changes)} changes")
        
        # Step 8: Create both versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[███████] Step 8: Creating final documents...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_multipass_redlined.docx"
        self.create_redlined_version(input_path, self.all_changes, redlined_path)
        
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_multipass_clean-version.docx"
        self.create_clean_version(input_path, self.all_changes, clean_path)
        
        # Step 9: Restore personal info
        print("[███████] Step 9: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ MULTI-PASS REDLINING COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Total Changes: {len(self.all_changes)} across 4 passes")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 multipass_redliner.py input.docx")
        sys.exit(1)
    
    processor = MultiPassRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
