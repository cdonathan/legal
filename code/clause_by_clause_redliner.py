#!/usr/bin/env python3
"""
Clause-by-Clause NDA Redlining System
Proper 5+ step workflow: redact -> analyze -> chunk -> clause redline -> reassemble
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class ClauseByClauseRedliner:
    def __init__(self):
        self.personal_info = {}
        self.document_context = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Step 1: Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def analyze_full_document(self, redacted_text):
        """Step 2: Full document analysis to create context"""
        if not self.openai_client:
            return self._mock_context()
        
        prompt = f"""Analyze this NDA and create a context document for clause-by-clause redlining.

FULL NDA TEXT:
{redacted_text}

Return ONLY JSON with document context:
{{
  "transaction_type": "Real estate acquisition/Commercial lease/etc",
  "purpose": "Brief description of business purpose",
  "tone": "Balanced/Seller-favorable/Buyer-favorable",
  "governing_law": "State if mentioned",
  "special_clauses": ["Any unique provisions"],
  "missing_standard_clauses": ["What institutional clauses are missing"],
  "document_structure": "Well-organized/Needs improvement/etc"
}}

Focus on understanding the overall transaction and legal context."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=800,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Document analysis error: {e}")
            return self._mock_context()
    
    def _mock_context(self):
        return {
            "transaction_type": "Real estate confidentiality agreement",
            "purpose": "Property evaluation and due diligence",
            "tone": "Balanced institutional NDA",
            "governing_law": "Not specified",
            "special_clauses": [],
            "missing_standard_clauses": ["permitted_recipients", "return_materials"],
            "document_structure": "Standard NDA format"
        }
    
    def chunk_by_clauses(self, doc_paragraphs):
        """Step 3: Chunk document by logical clauses"""
        clauses = []
        current_clause = []
        clause_number = 1
        
        for para in doc_paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect new clause (numbered sections, headers, etc.)
            if (re.match(r'^\d+\.', text) or 
                re.match(r'^[A-Z][A-Z\s]{10,}', text) or
                len(current_clause) > 3):  # Max 3 paragraphs per clause
                
                if current_clause:
                    clauses.append({
                        "clause_number": clause_number,
                        "title": self._extract_clause_title(current_clause[0]),
                        "text": '\n'.join(current_clause)
                    })
                    clause_number += 1
                
                current_clause = [text]
            else:
                current_clause.append(text)
        
        # Add final clause
        if current_clause:
            clauses.append({
                "clause_number": clause_number,
                "title": self._extract_clause_title(current_clause[0]),
                "text": '\n'.join(current_clause)
            })
        
        return clauses
    
    def _extract_clause_title(self, first_line):
        """Extract clause title from first line"""
        # Remove numbering and get first few words
        clean_line = re.sub(r'^\d+\.?\s*', '', first_line)
        words = clean_line.split()[:4]
        return ' '.join(words) if words else "Untitled Clause"
    
    def redline_clause(self, clause, document_context):
        """Step 4+: Redline individual clause with full context"""
        if not self.openai_client:
            return {"changes": []}
        
        # Load reference materials
        with open('/home/cliff/redact/redline_project/golden_nda.md', 'r') as f:
            golden_nda = f.read()
        
        with open('/home/cliff/redact/redline_project/nda_clause_library.md', 'r') as f:
            clause_library = f.read()
        
        prompt = f"""You are redlining Clause {clause['clause_number']}: "{clause['title']}" in a {document_context['transaction_type']}.

DOCUMENT CONTEXT:
Transaction: {document_context['transaction_type']}
Purpose: {document_context['purpose']}
Tone: {document_context['tone']}
Missing Standard Clauses: {document_context['missing_standard_clauses']}

GOLDEN NDA REFERENCE:
{golden_nda[:1000]}...

CLAUSE LIBRARY:
{clause_library[:1000]}...

CLAUSE TO REDLINE:
{clause['text']}

REDLINING INSTRUCTIONS:
1. Make minimal inline edits within existing sentences
2. Capitalize defined terms consistently
3. Align language with Golden NDA standards
4. Fix awkward phrasing
5. Only suggest new clauses if critical protections are missing

Return ONLY JSON:
{{
  "clause_number": {clause['clause_number']},
  "changes": [
    {{
      "type": "replace",
      "find": "exact text to replace",
      "replace": "exact replacement text",
      "reason": "alignment with institutional standards"
    }}
  ],
  "status": "MODIFIED" or "NO_CHANGE"
}}

Focus on precise inline edits that improve this specific clause."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_str = content[json_start:json_end]
            
            return json.loads(json_str)
        except Exception as e:
            print(f"Clause redline error: {e}")
            return {"clause_number": clause['clause_number'], "changes": [], "status": "NO_CHANGE"}
    
    def create_redlined_version(self, original_doc_path, all_clause_changes, output_path):
        """Create redlined version with all clause changes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("CLAUSE-BY-CLAUSE REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Deletions | 🟢 Green underline = Insertions")
        redlined_doc.add_paragraph("=" * 80)
        
        # Process each paragraph
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply all clause changes to this paragraph
            for clause_changes in all_clause_changes:
                for change in clause_changes.get('changes', []):
                    if change['type'] == 'replace' and change['find'] in para_text:
                        parts = para_text.split(change['find'])
                        
                        if parts[0]:
                            new_para.add_run(parts[0])
                        
                        # Deleted text
                        del_run = new_para.add_run(change['find'])
                        del_run.font.strike = True
                        del_run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        # Inserted text
                        ins_run = new_para.add_run(change['replace'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
                        
                        if len(parts) > 1 and parts[1]:
                            new_para.add_run(parts[1])
                        
                        para_text = ""
                        break
            
            if para_text:
                new_para.add_run(para_text)
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, all_clause_changes, output_path):
        """Create clean version with all changes applied"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply all changes
            for clause_changes in all_clause_changes:
                for change in clause_changes.get('changes', []):
                    if change['type'] == 'replace' and change['find'] in para_text:
                        para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Complete clause-by-clause workflow"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text and paragraphs
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: Analyze full document
        print("[██░░░░] Step 2: Analyzing full document for context...")
        self.document_context = self.analyze_full_document(redacted_text)
        print(f"   ✓ Transaction: {self.document_context['transaction_type']}")
        
        # Step 3: Chunk by clauses
        print("[███░░░] Step 3: Chunking by clauses...")
        clauses = self.chunk_by_clauses(doc.paragraphs)
        print(f"   ✓ Found {len(clauses)} clauses")
        
        # Step 4+: Redline each clause
        print("[████░░] Step 4+: Redlining clauses with context...")
        all_clause_changes = []
        modified_clauses = 0
        
        for clause in clauses:
            clause_result = self.redline_clause(clause, self.document_context)
            all_clause_changes.append(clause_result)
            if clause_result['status'] == 'MODIFIED':
                modified_clauses += 1
        
        print(f"   ✓ Modified {modified_clauses}/{len(clauses)} clauses")
        
        # Step 5: Create both versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[█████░] Step 5: Creating redlined version...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_clause_redlined.docx"
        self.create_redlined_version(input_path, all_clause_changes, redlined_path)
        
        print("[██████] Step 6: Creating clean version...")
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_clause_clean-version.docx"
        self.create_clean_version(input_path, all_clause_changes, clean_path)
        
        # Step 7: Restore personal info
        print("[██████] Step 7: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ CLAUSE-BY-CLAUSE REDLINING COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🔍 Context: {self.document_context['transaction_type']}")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 clause_by_clause_redliner.py input.docx")
        sys.exit(1)
    
    processor = ClauseByClauseRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
