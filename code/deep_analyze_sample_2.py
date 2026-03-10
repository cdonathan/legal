#!/usr/bin/env python3
"""
Deep Analysis of Attorney Changes in Sample 2
"""

from docx import Document
import difflib

def deep_analyze_sample_2():
    """Deep analysis of what attorneys actually accomplished in Sample 2"""
    
    # Load documents
    pre_redline = Document('/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Confidentiality Agreement_Sample_2_pre_redline.docx')
    attorney_redlined = Document('/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Confidentiality Agreement_Sample_2.docx')
    
    # Extract text by paragraphs
    pre_paras = [p.text.strip() for p in pre_redline.paragraphs if p.text.strip()]
    attorney_paras = [p.text.strip() for p in attorney_redlined.paragraphs if p.text.strip()]
    
    print("DEEP ANALYSIS: ATTORNEY CHANGES IN SAMPLE 2")
    print("=" * 80)
    print(f"Pre-redline paragraphs: {len(pre_paras)}")
    print(f"Attorney redlined paragraphs: {len(attorney_paras)}")
    print(f"Net change: {len(attorney_paras) - len(pre_paras)} paragraphs")
    
    print("\nDETAILED CHANGE ANALYSIS:")
    print("=" * 80)
    
    # Track different types of changes
    redactions = []
    replacements = []
    additions = []
    deletions = []
    
    # Compare paragraph by paragraph
    max_len = max(len(pre_paras), len(attorney_paras))
    
    for i in range(max_len):
        pre_text = pre_paras[i] if i < len(pre_paras) else ""
        attorney_text = attorney_paras[i] if i < len(attorney_paras) else ""
        
        if pre_text and not attorney_text:
            deletions.append(f"DELETED Para {i+1}: {pre_text[:100]}...")
        elif not pre_text and attorney_text:
            additions.append(f"ADDED Para {i+1}: {attorney_text[:100]}...")
        elif pre_text != attorney_text and pre_text and attorney_text:
            # Analyze the type of change
            if "XXXX" in attorney_text and "XXXX" not in pre_text:
                redactions.append({
                    'para': i+1,
                    'original': pre_text,
                    'redacted': attorney_text,
                    'type': 'REDACTION'
                })
            else:
                replacements.append({
                    'para': i+1,
                    'original': pre_text,
                    'modified': attorney_text,
                    'type': 'REPLACEMENT'
                })
    
    # Print findings
    print(f"\nREDACTIONS (Personal Info → XXXX): {len(redactions)}")
    for redaction in redactions[:5]:  # Show first 5
        print(f"  Para {redaction['para']}: {redaction['original'][:80]}...")
        print(f"  →: {redaction['redacted'][:80]}...")
        print()
    
    print(f"\nREPLACEMENTS (Content Changes): {len(replacements)}")
    for replacement in replacements[:5]:  # Show first 5
        print(f"  Para {replacement['para']}:")
        print(f"    ORIGINAL: {replacement['original'][:80]}...")
        print(f"    MODIFIED: {replacement['modified'][:80]}...")
        
        # Show word-level differences
        orig_words = replacement['original'].split()
        mod_words = replacement['modified'].split()
        
        if len(orig_words) > len(mod_words):
            print(f"    SHORTENED: {len(orig_words)} → {len(mod_words)} words")
        elif len(mod_words) > len(orig_words):
            print(f"    LENGTHENED: {len(orig_words)} → {len(mod_words)} words")
        else:
            print(f"    SAME LENGTH: {len(orig_words)} words")
        print()
    
    print(f"\nADDITIONS (New Content): {len(additions)}")
    for addition in additions[:3]:
        print(f"  {addition}")
    
    print(f"\nDELETIONS (Removed Content): {len(deletions)}")
    for deletion in deletions[:3]:
        print(f"  {deletion}")
    
    # Analyze attorney strategy
    print("\n" + "=" * 80)
    print("ATTORNEY STRATEGY ANALYSIS:")
    print("=" * 80)
    
    total_changes = len(redactions) + len(replacements) + len(additions) + len(deletions)
    
    if total_changes > 0:
        redaction_pct = (len(redactions) / total_changes) * 100
        replacement_pct = (len(replacements) / total_changes) * 100
        addition_pct = (len(additions) / total_changes) * 100
        deletion_pct = (len(deletions) / total_changes) * 100
        
        print(f"Change Distribution:")
        print(f"  Redactions (XXXX): {redaction_pct:.1f}%")
        print(f"  Replacements: {replacement_pct:.1f}%")
        print(f"  Additions: {addition_pct:.1f}%")
        print(f"  Deletions: {deletion_pct:.1f}%")
    
    # Look for patterns in redactions
    print(f"\nREDACTION PATTERNS:")
    redacted_items = []
    for redaction in redactions:
        original = redaction['original']
        redacted = redaction['redacted']
        
        # Find what was replaced with XXXX
        orig_words = original.split()
        red_words = redacted.split()
        
        for i, (o_word, r_word) in enumerate(zip(orig_words, red_words)):
            if r_word == "XXXX" and o_word != "XXXX":
                redacted_items.append(o_word)
    
    # Count redacted item types
    from collections import Counter
    redacted_counts = Counter(redacted_items)
    
    print("Most commonly redacted items:")
    for item, count in redacted_counts.most_common(10):
        print(f"  '{item}': {count} times")
    
    # Analyze replacement patterns
    print(f"\nREPLACEMENT PATTERNS:")
    for replacement in replacements[:3]:
        orig_len = len(replacement['original'])
        mod_len = len(replacement['modified'])
        
        if mod_len < orig_len * 0.8:
            print(f"  Para {replacement['para']}: SIGNIFICANT SHORTENING ({orig_len} → {mod_len} chars)")
        elif mod_len > orig_len * 1.2:
            print(f"  Para {replacement['para']}: SIGNIFICANT LENGTHENING ({orig_len} → {mod_len} chars)")
        else:
            print(f"  Para {replacement['para']}: MINOR MODIFICATION ({orig_len} → {mod_len} chars)")
    
    print("\n" + "=" * 80)
    print("CONCLUSION: WHAT ATTORNEYS ACTUALLY DID")
    print("=" * 80)
    
    if len(redactions) > len(replacements) + len(additions) + len(deletions):
        print("PRIMARY FOCUS: REDACTION - Replacing personal info with XXXX")
    elif len(replacements) > len(redactions):
        print("PRIMARY FOCUS: CONTENT MODIFICATION - Changing existing language")
    elif len(additions) > len(redactions):
        print("PRIMARY FOCUS: CONTENT ADDITION - Adding new clauses")
    else:
        print("MIXED APPROACH: Multiple types of changes")
    
    print(f"\nThis was NOT about removing verbose language.")
    print(f"This was about: {redaction_pct:.0f}% redaction, {replacement_pct:.0f}% content changes")

if __name__ == "__main__":
    deep_analyze_sample_2()
