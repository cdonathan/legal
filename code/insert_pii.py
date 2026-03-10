import random
import re
from docx import Document

class PIIInserter:
    def __init__(self):
        self.names = [
            "John Smith", "Sarah Johnson", "Michael Brown", "Emily Davis", 
            "David Wilson", "Jennifer Miller", "Robert Taylor", "Lisa Anderson",
            "Christopher Martinez", "Amanda Thompson", "Matthew Garcia", "Jessica Rodriguez"
        ]
        
        self.addresses = [
            "123 Main Street, Anytown, CA 90210",
            "456 Oak Avenue, Springfield, IL 62701", 
            "789 Pine Road, Austin, TX 78701",
            "321 Elm Drive, Seattle, WA 98101",
            "654 Maple Lane, Denver, CO 80202",
            "987 Cedar Court, Miami, FL 33101"
        ]
        
        self.emails = [
            "john.smith@email.com", "sarah.j@company.com", "mbrown@business.net",
            "emily.davis@firm.org", "dwilson@legal.com", "jmiller@realty.biz"
        ]
        
        self.phones = [
            "(555) 123-4567", "(555) 987-6543", "(555) 456-7890",
            "(555) 321-0987", "(555) 654-3210", "(555) 789-0123"
        ]
        
        self.companies = [
            "ABC Corporation", "XYZ Industries", "Global Enterprises LLC",
            "Premier Holdings Inc", "Strategic Partners Group", "Innovative Solutions Co"
        ]

    def insert_pii_into_docx(self, input_path, output_path):
        """Insert realistic PII data into DOCX document"""
        try:
            doc = Document(input_path)
            
            # Get random PII data
            buyer_name = random.choice(self.names)
            seller_name = random.choice(self.names)
            buyer_address = random.choice(self.addresses)
            seller_address = random.choice(self.addresses)
            buyer_email = random.choice(self.emails)
            seller_email = random.choice(self.emails)
            buyer_phone = random.choice(self.phones)
            seller_phone = random.choice(self.phones)
            buyer_company = random.choice(self.companies)
            seller_company = random.choice(self.companies)
            
            # Replace placeholders and add PII
            replacements = {
                # Common placeholders
                '[BUYER NAME]': buyer_name,
                '[SELLER NAME]': seller_name,
                '[PURCHASER NAME]': buyer_name,
                '[DISCLOSING PARTY]': seller_name,
                '[RECIPIENT]': buyer_name,
                '[BUYER]': buyer_name,
                '[SELLER]': seller_name,
                
                # Addresses
                '[BUYER ADDRESS]': buyer_address,
                '[SELLER ADDRESS]': seller_address,
                '[ADDRESS]': buyer_address,
                
                # Contact info
                '[EMAIL]': buyer_email,
                '[PHONE]': buyer_phone,
                
                # Companies
                '[COMPANY]': buyer_company,
                '[BUYER COMPANY]': buyer_company,
                '[SELLER COMPANY]': seller_company,
                
                # Generic patterns
                '____': buyer_name,
                '______': seller_name,
                'XXXXX': buyer_company,
                
                # Date placeholders
                '____ day of': f"{random.randint(1,28)} day of",
                '201[ ]': f"202{random.randint(3,6)}",
                
                # Signature lines
                'Signature _______': f"Signature: {buyer_name}",
                'Print Name _______': f"Print Name: {buyer_name}",
            }
            
            # Process all paragraphs
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)
            
            # Add additional PII in strategic locations
            if doc.paragraphs:
                # Add contact info to first paragraph if it mentions parties
                first_para = doc.paragraphs[0].text.lower()
                if any(word in first_para for word in ['agreement', 'party', 'between']):
                    # Insert contact details
                    new_para = doc.paragraphs[0]._element
                    new_para.getparent().insert(1, doc.add_paragraph(
                        f"Buyer Contact: {buyer_name}, {buyer_address}, {buyer_email}, {buyer_phone}"
                    )._element)
                    new_para.getparent().insert(2, doc.add_paragraph(
                        f"Seller Contact: {seller_name}, {seller_address}, {seller_email}, {seller_phone}"
                    )._element)
            
            doc.save(output_path)
            print(f"   ✓ Inserted PII data: {buyer_name}, {seller_name}, addresses, emails, phones")
            return True
            
        except Exception as e:
            print(f"   ❌ PII insertion failed: {e}")
            return False

def main():
    inserter = PIIInserter()
    
    # Process all converted DOCX files from testExamples
    import os
    test_files = [
        "/tmp/example1_converted.docx",
        "/tmp/example2_converted.docx", 
        "/tmp/example3_converted.docx",
        "/tmp/example4_converted.docx",
        "/home/cliff/redact/redline_project/testExamples/example5.docx"
    ]
    
    output_dir = "/home/cliff/redact/redline_project/testExamples_with_pii"
    os.makedirs(output_dir, exist_ok=True)
    
    for i, file_path in enumerate(test_files, 1):
        if os.path.exists(file_path):
            output_path = os.path.join(output_dir, f"example{i}_with_pii.docx")
            print(f"Processing example{i}...")
            inserter.insert_pii_into_docx(file_path, output_path)
        else:
            print(f"File not found: {file_path}")

if __name__ == "__main__":
    main()
