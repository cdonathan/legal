#!/usr/bin/env python3
"""
Fix track changes implementation for NDA redlining
"""

import re
import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor

def create_proper_track_changes(original_text, redlined_text, output_path):
    """Create Word document with proper track changes"""
    doc = Document()
    
    # Split into lines for processing
    original_lines = original_text.split('\n')
    redlined_lines = redlined_text.split('\n')
    
    for line in redlined_lines:
        if not line.strip():
            doc.add_paragraph()
            continue
            
        para = doc.add_paragraph()
        
        # Handle INSERT markup - show as inserted text
        if '[INSERT:' in line:
            parts = re.split(r'\[INSERT:\s*([^\]]+)\]', line)
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Regular text
                    if part.strip():
                        para.add_run(part)
                else:  # Inserted text - make it green and underlined
                    run = para.add_run(part)
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    run.underline = True
                    
                    # Add revision markup
                    run_xml = run._element
                    ins_xml = parse_xml(f'<w:ins {nsdecls("w")} w:id="1" w:author="AI Redline" w:date="2024-03-05T18:00:00Z">{run_xml.xml}</w:ins>')
                    run_xml.getparent().replace(run_xml, ins_xml)
        
        # Handle DELETE markup - show as struck through
        elif '[DELETE:' in line:
            parts = re.split(r'\[DELETE:\s*([^\]]+)\]', line)
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Regular text
                    if part.strip():
                        para.add_run(part)
                else:  # Deleted text - make it red and struck through
                    run = para.add_run(part)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    run.font.strike = True
        
        else:  # Regular text
            para.add_run(line)
    
    doc.save(output_path)
    print(f"Created track changes document: {os.path.basename(output_path)}")

def test_track_changes():
    """Test with sample redlined content"""
    original = "This is the original confidentiality agreement."
    redlined = """This is the original confidentiality agreement. [INSERT: The Receiving Party may disclose Confidential Information to its attorneys, accountants, and financial advisors who have a need to know such information for the Purpose.] [DELETE: old clause text]"""
    
    create_proper_track_changes(original, redlined, "/home/cliff/redact/redline_project/test_track_changes.docx")

if __name__ == "__main__":
    test_track_changes()
