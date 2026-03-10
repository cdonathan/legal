#!/usr/bin/env python3
"""
Minimal Changes NDA System - Uses Prioritized Golden NDA
Only makes P1 (Required) and essential P2 (Conditional) changes
"""

import os
import sys
import subprocess
import json
import re
import openai
from docx import Document
from docx.shared import RGBColor

class MinimalChangesNDASystem:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = len(self.personal_info) + 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential', 'broker', 'seller']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                if placeholder not in self.personal_info:
                    self.personal_info[placeholder] = match.group()
                    redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                    counter += 1
        
        return redacted_text
    
    def convert_with_libreoffice(self, input_path):
        """Convert document using LibreOffice"""
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            input_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
        
        return None
    
    def get_minimal_ai_recommendations(self, redacted_text, base_name):
        """Get MINIMAL AI recommendations using prioritized Golden NDA"""
        if not self.openai_client:
            return None, []
        
        # Load prioritized Golden NDA
        with open('/mnt/c/Users/cliff/Downloads/golden_nda_prioritized.md', 'r') as f:
            prioritized_golden = f.read()
        
        # Add line numbers
        lines = redacted_text.split('\n')
        numbered_lines = []
        line_number = 1
        
        for line in lines:
            if line.strip():
                numbered_lines.append(f"LINE {line_number:03d}: {line.strip()}")
                line_number += 1
        
        line_numbered_text = '\n'.join(numbered_lines)
        
        prompt = f"""You are an attorney making MINIMAL changes to align with institutional standards.

PRIORITIZED GOLDEN NDA RULES:
{prioritized_golden}

LINE-NUMBERED NDA:
{line_numbered_text}

CRITICAL RULES:
- Only make P1 (Required) changes if completely missing
- Only make P2 (Conditional) changes if clearly necessary
- NEVER add P3 or P4 sections
- Make MINIMAL word-level changes, not paragraph additions
- Focus on fixing existing language, not adding new sections

MINIMAL CHANGES ONLY:
1. If "Informational Materials" is vague → replace with "Confidential Information" (simple term swap)
2. If purpose is missing → add ONE sentence about evaluation purpose
3. If confidentiality obligation is weak → strengthen existing language
4. If no exceptions exist → add ONE sentence about standard exceptions

Return ONLY JSON array (maximum 3-4 changes):
[
  {{
    "problem_addressed": "Vague term definition",
    "change_type": "replace_term",
    "line_number": 3,
    "find_text": "Informational Materials",
    "replace_with": "Confidential Information",
    "reason": "P1 Required - clarifies protected information",
    "priority": "P1"
  }},
  {{
    "problem_addressed": "Missing purpose",
    "change_type": "insert_sentence",
    "line_number": 2,
    "insert_text": "This information is provided solely for evaluating a potential transaction.",
    "reason": "P1 Required - states disclosure purpose",
    "priority": "P1"
  }}
]

MAXIMUM 4 CHANGES. Focus on P1 requirements only."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            
            # Save recommendations
            recommendations_file = f"/home/cliff/redact/redline_project/{base_name}_Minimal_AI_Recommendations.md"
            with open(recommendations_file, 'w') as f:
                f.write(f"# Minimal AI Recommendations for {base_name}\n\n")
                f.write("**Focus: P1 (Required) changes only - MINIMAL modifications**\n\n")
                f.write("**AI Analysis Output:**\n\n")
                f.write("```json\n")
                f.write(content)
                f.write("\n```\n\n")
            
            # Parse JSON
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = content[json_start:json_end]
                recommendations = json.loads(json_str)
                
                # Add detailed recommendations
                with open(recommendations_file, 'a') as f:
                    f.write(f"**Parsed Recommendations:** {len(recommendations)} (MINIMAL)\n\n")
                    
                    for i, rec in enumerate(recommendations, 1):
                        f.write(f"## Minimal Change {i}: {rec.get('problem_addressed', 'Unknown')}\n\n")
                        f.write(f"**Priority:** {rec.get('priority', 'UNKNOWN')}\n\n")
                        f.write(f"**Change Type:** {rec.get('change_type', 'unknown')}\n\n")
                        f.write(f"**Line Number:** {rec.get('line_number', 'N/A')}\n\n")
                        f.write(f"**Find Text:**\n```\n{rec.get('find_text', 'N/A')}\n```\n\n")
                        f.write(f"**Replace/Insert:**\n```\n{rec.get('replace_with', rec.get('insert_text', 'N/A'))}\n```\n\n")
                        f.write(f"**Reason:** {rec.get('reason', 'N/A')}\n\n")
                        f.write("---\n\n")
                
                print(f"   ✓ Created minimal AI recommendations: {os.path.basename(recommendations_file)}")
                return recommendations_file, recommendations
            else:
                print("   ❌ Could not parse AI recommendations JSON")
                return recommendations_file, []
                
        except Exception as e:
            print(f"   ❌ AI recommendations error: {e}")
            return None, []
    
    def implement_minimal_changes(self, original_path, recommendations, base_name):
        """Implement MINIMAL changes only"""
        # Convert to paragraphs
        text_content = self.convert_with_libreoffice(original_path)
        if not text_content:
            return None, None
        
        # Split into paragraphs
        paragraphs = []
        lines = text_content.split('\n')
        current_para = []
        para_number = 1
        
        for line in lines:
            if line.strip():
                current_para.append(line.strip())
            else:
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append({
                        'paragraph_number': para_number,
                        'original_text': para_text,
                        'current_text': para_text,
                        'changes': []
                    })
                    para_number += 1
                    current_para = []
        
        if current_para:
            para_text = ' '.join(current_para)
            paragraphs.append({
                'paragraph_number': para_number,
                'original_text': para_text,
                'current_text': para_text,
                'changes': []
            })
        
        # Implement MINIMAL changes
        implemented_count = 0
        
        for rec in recommendations:
            try:
                if rec.get('change_type') == 'replace_term':
                    # Simple term replacement
                    find_text = rec.get('find_text', '')
                    replace_text = rec.get('replace_with', '')
                    
                    for para in paragraphs:
                        if find_text in para['current_text']:
                            new_text = para['current_text'].replace(find_text, replace_text)
                            para['current_text'] = new_text
                            para['changes'].append({
                                'type': 'replace_term',
                                'find': find_text,
                                'replace': replace_text,
                                'reason': rec.get('reason', '')
                            })
                            implemented_count += 1
                            print(f"   ✓ Replaced term: {find_text} → {replace_text}")
                            break
                
                elif rec.get('change_type') == 'insert_sentence':
                    # Insert single sentence
                    line_num = rec.get('line_number')
                    insert_text = rec.get('insert_text', '')
                    
                    if line_num <= len(paragraphs):
                        # Add to existing paragraph, not new paragraph
                        target_para = paragraphs[line_num - 1]
                        target_para['current_text'] += f" {insert_text}"
                        target_para['changes'].append({
                            'type': 'insert_sentence',
                            'text': insert_text,
                            'reason': rec.get('reason', '')
                        })
                        implemented_count += 1
                        print(f"   ✓ Added sentence to paragraph {line_num}")
                
                elif rec.get('change_type') == 'strengthen_language':
                    # Strengthen existing language
                    line_num = rec.get('line_number')
                    find_text = rec.get('find_text', '')
                    replace_text = rec.get('replace_with', '')
                    
                    if line_num <= len(paragraphs):
                        para = paragraphs[line_num - 1]
                        if find_text in para['current_text']:
                            new_text = para['current_text'].replace(find_text, replace_text)
                            para['current_text'] = new_text
                            para['changes'].append({
                                'type': 'strengthen',
                                'find': find_text,
                                'replace': replace_text,
                                'reason': rec.get('reason', '')
                            })
                            implemented_count += 1
                            print(f"   ✓ Strengthened language in paragraph {line_num}")
                    
            except Exception as e:
                print(f"   ❌ Failed to implement {rec.get('problem_addressed', 'unknown')}: {e}")
        
        # Create redlined document
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_Minimal_Redlined.docx"
        doc = Document()
        
        # Header
        header = doc.add_paragraph()
        header_run = header.add_run("MINIMAL AI CHANGES - INSTITUTIONAL ALIGNMENT")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = Original | 🟢 Green underline = Minimal changes | Total: {implemented_count} changes")
        doc.add_paragraph("=" * 80)
        
        # Process paragraphs with minimal changes
        for para_data in paragraphs:
            doc_para = doc.add_paragraph()
            
            if para_data['changes']:
                for change in para_data['changes']:
                    if change['type'] in ['replace_term', 'strengthen']:
                        original_text = para_data['original_text']
                        find_text = change['find']
                        replace_text = change['replace']
                        
                        if find_text in original_text:
                            parts = original_text.split(find_text)
                            
                            if parts[0]:
                                doc_para.add_run(parts[0])
                            
                            # Show minimal deletion
                            del_run = doc_para.add_run(find_text)
                            del_run.font.strike = True
                            del_run.font.color.rgb = RGBColor(255, 0, 0)
                            
                            # Show minimal addition
                            ins_run = doc_para.add_run(replace_text)
                            ins_run.underline = True
                            ins_run.font.color.rgb = RGBColor(0, 128, 0)
                            
                            if len(parts) > 1 and parts[1]:
                                doc_para.add_run(parts[1])
                        else:
                            doc_para.add_run(para_data['current_text'])
                    
                    elif change['type'] == 'insert_sentence':
                        # Show sentence addition within paragraph
                        original_text = para_data['original_text']
                        added_text = change['text']
                        
                        doc_para.add_run(original_text)
                        
                        # Show added sentence
                        ins_run = doc_para.add_run(f" {added_text}")
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                if para_data['current_text']:
                    doc_para.add_run(para_data['current_text'])
        
        doc.save(redlined_path)
        
        # Create clean document
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_Minimal_Clean.docx"
        clean_doc = Document()
        
        for para_data in paragraphs:
            if para_data['current_text']:
                para = clean_doc.add_paragraph()
                para.add_run(para_data['current_text'])
        
        clean_doc.save(clean_path)
        
        print(f"   ✓ Created minimal redlined document: {os.path.basename(redlined_path)}")
        print(f"   ✓ Created minimal clean document: {os.path.basename(clean_path)}")
        print(f"   ✓ Implemented {implemented_count}/{len(recommendations)} minimal changes")
        
        return redlined_path, clean_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information in document"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
    
    def process_minimal_nda_changes(self, input_path):
        """Process NDA with MINIMAL changes only"""
        print(f"🔄 Minimal NDA Changes: {os.path.basename(input_path)}")
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        # Convert document
        print("[█░░░] Step 1: Converting document with LibreOffice...")
        text_content = self.convert_with_libreoffice(input_path)
        if not text_content:
            print("❌ LibreOffice conversion failed")
            return
        
        # Redact personal info
        redacted_text = self.redact_personal_info(text_content)
        print(f"   ✓ Redacted {len(self.personal_info)} personal items")
        
        # Get MINIMAL AI recommendations
        print("[██░░] Step 2: Getting MINIMAL AI recommendations...")
        recommendations_file, recommendations = self.get_minimal_ai_recommendations(redacted_text, base_name)
        
        # Implement MINIMAL changes
        print("[███░] Step 3: Implementing MINIMAL changes...")
        redlined_path, clean_path = self.implement_minimal_changes(input_path, recommendations, base_name)
        
        # Restore personal info
        print("[████] Step 4: Restoring personal information...")
        if redlined_path:
            self.restore_personal_info(redlined_path)
        if clean_path:
            self.restore_personal_info(clean_path)
        
        print("✅ MINIMAL NDA CHANGES COMPLETE!")
        print(f"📋 Minimal AI Recommendations: {os.path.basename(recommendations_file) if recommendations_file else 'FAILED'}")
        print(f"📝 Minimal Redlined: {os.path.basename(redlined_path) if redlined_path else 'FAILED'}")
        print(f"📄 Minimal Clean: {os.path.basename(clean_path) if clean_path else 'FAILED'}")
        print(f"🎯 Total changes: {len(recommendations)} (MINIMAL P1 only)")

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 minimal_changes_nda_system.py input.docx")
        sys.exit(1)
    
    system = MinimalChangesNDASystem()
    system.process_minimal_nda_changes(sys.argv[1])

if __name__ == "__main__":
    main()
