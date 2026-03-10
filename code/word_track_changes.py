#!/usr/bin/env python3
"""
AI-driven Word Track Changes System
AI returns structured changes, code applies them as native Word revisions
"""

import json
import re
import os
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from datetime import datetime

class WordTrackChanges:
    def __init__(self):
        self.change_id = 1
        self.author = "AI Redline System"
        self.date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    
    def get_ai_changes(self, redacted_text):
        """Get structured changes from AI"""
        # Simplified for testing - in real version, call OpenAI
        sample_changes = {
            "document_analysis": {
                "transaction_type": "Real estate NDA",
                "missing_clauses": ["permitted_recipients", "return_materials"],
                "recommendations": "Add standard institutional clauses"
            },
            "changes": [
                {
                    "type": "insert",
                    "location": "after_paragraph",
                    "paragraph_number": 3,
                    "text": "The Receiving Party may disclose Confidential Information to its attorneys, accountants, financial advisors, and other professional advisors who have a need to know such information for the Purpose.",
                    "reason": "Add permitted recipients clause - standard institutional protection"
                },
                {
                    "type": "insert", 
                    "location": "after_paragraph",
                    "paragraph_number": 5,
                    "text": "Upon written request, the Receiving Party shall promptly return or destroy all Confidential Information and any copies thereof.",
                    "reason": "Add return of materials clause - required for enforceability"
                },
                {
                    "type": "replace",
                    "paragraph_number": 2,
                    "original_text": "confidential information",
                    "new_text": "Confidential Information",
                    "reason": "Capitalize defined term for consistency"
                }
            ]
        }
        return sample_changes
    
    def create_insertion_element(self, text, run_props=None):
        """Create Word insertion revision element"""
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), str(self.change_id))
        ins.set(qn('w:author'), self.author)
        ins.set(qn('w:date'), self.date)
        
        # Create run inside insertion
        run = OxmlElement('w:r')
        if run_props:
            run.append(run_props)
        
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)
        ins.append(run)
        
        self.change_id += 1
        return ins
    
    def create_deletion_element(self, text):
        """Create Word deletion revision element"""
        delete = OxmlElement('w:del')
        delete.set(qn('w:id'), str(self.change_id))
        delete.set(qn('w:author'), self.author)
        delete.set(qn('w:date'), self.date)
        
        # Create run inside deletion
        run = OxmlElement('w:r')
        t = OxmlElement('w:delText')
        t.text = text
        run.append(t)
        delete.append(run)
        
        self.change_id += 1
        return delete
    
    def apply_changes_to_document(self, doc_path, changes, output_path):
        """Apply structured changes to Word document with track changes"""
        doc = Document(doc_path)
        
        # Sort changes by paragraph number (reverse order for insertions)
        sorted_changes = sorted(changes['changes'], key=lambda x: x.get('paragraph_number', 0), reverse=True)
        
        for change in sorted_changes:
            if change['type'] == 'insert' and change['location'] == 'after_paragraph':
                para_num = change['paragraph_number']
                if para_num < len(doc.paragraphs):
                    # Insert new paragraph after specified paragraph
                    new_para = doc.paragraphs[para_num]._element.getparent().insert(
                        doc.paragraphs[para_num]._element.getparent().index(doc.paragraphs[para_num]._element) + 1,
                        OxmlElement('w:p')
                    )
                    
                    # Add insertion revision to new paragraph
                    ins_element = self.create_insertion_element(change['text'])
                    new_para.append(ins_element)
            
            elif change['type'] == 'replace':
                para_num = change['paragraph_number']
                if para_num < len(doc.paragraphs):
                    para = doc.paragraphs[para_num]
                    
                    # Find and replace text with track changes
                    for run in para.runs:
                        if change['original_text'] in run.text:
                            # Split the run at the replacement point
                            before_text = run.text.split(change['original_text'])[0]
                            after_text = run.text.split(change['original_text'], 1)[1]
                            
                            # Clear original run
                            run.clear()
                            
                            # Add before text
                            if before_text:
                                run.text = before_text
                            
                            # Add deletion
                            del_element = self.create_deletion_element(change['original_text'])
                            run._element.addnext(del_element)
                            
                            # Add insertion
                            ins_element = self.create_insertion_element(change['new_text'])
                            del_element.addnext(ins_element)
                            
                            # Add after text
                            if after_text:
                                after_run = OxmlElement('w:r')
                                after_t = OxmlElement('w:t')
                                after_t.text = after_text
                                after_run.append(after_t)
                                ins_element.addnext(after_run)
                            
                            break
        
        doc.save(output_path)
        return output_path
    
    def process_nda_with_track_changes(self, input_path):
        """Complete workflow with proper track changes"""
        print(f"Processing: {os.path.basename(input_path)}")
        
        # Step 1: Extract and redact (simplified)
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Step 2: Get AI changes
        print("Getting AI redline suggestions...")
        changes = self.get_ai_changes(text)
        
        print(f"AI suggested {len(changes['changes'])} changes:")
        for i, change in enumerate(changes['changes'], 1):
            print(f"  {i}. {change['type'].title()}: {change['reason']}")
        
        # Step 3: Apply changes with track changes
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = f"/home/cliff/redact/redline_project/{base_name}_tracked_redlined.docx"
        
        print("Applying changes with Word track changes...")
        self.apply_changes_to_document(input_path, changes, output_path)
        
        print(f"✓ Complete: {os.path.basename(output_path)}")
        print("Track changes are now visible in Word - users can Accept/Reject")
        
        return output_path

def test_track_changes():
    """Test the track changes system"""
    tracker = WordTrackChanges()
    
    # Test with a sample document
    test_file = "/home/cliff/redact/OneDrive_1_3-5-2026/REDLINE_Conf_Agr_Sample1-pre-redline.docx"
    if os.path.exists(test_file):
        tracker.process_nda_with_track_changes(test_file)
    else:
        print(f"Test file not found: {test_file}")

if __name__ == "__main__":
    test_track_changes()
