#!/usr/bin/env python3
"""
Minimal Changes AI Redliner - Focus on deletion and simplification
"""

import json
import os
import re
import openai
from docx import Document
from docx.shared import RGBColor

class MinimalChangesRedliner:
    def __init__(self):
        self.personal_info = {}
        self.openai_client = self._setup_openai()
    
    def _setup_openai(self):
        try:
            with open('/home/cliff/redact/openai_api_key.txt', 'r') as f:
                api_key = f.read().strip()
            return openai.OpenAI(api_key=api_key)
        except:
            return None
    
    def redact_personal_info(self, text):
        """Redact personal information"""
        redacted_text = text
        counter = 1
        
        patterns = [
            (r'\b[A-Z][A-Z\s&,\.]{3,}(?:LLC|INC|CORP|LP|LLP|COMPANY|CO\.)\b', 'COMPANY'),
            (r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Boulevard|Blvd|Lane|Ln)[^,\n]*', 'ADDRESS'),
            (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', 'DATE'),
            (r'\b[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b', 'NAME')
        ]
        
        for pattern, label in patterns:
            for match in re.finditer(pattern, text):
                if any(term in match.group().lower() for term in 
                      ['party', 'agreement', 'information', 'confidential']):
                    continue
                
                placeholder = f"[{label}_{counter}]"
                self.personal_info[placeholder] = match.group()
                redacted_text = redacted_text.replace(match.group(), placeholder, 1)
                counter += 1
        
        return redacted_text
    
    def get_minimal_changes(self, redacted_text):
        """Get minimal changes focused on deletion and simplification"""
        if not self.openai_client:
            return []
        
        prompt = f"""You are an attorney making MINIMAL changes to achieve institutional NDA standards.

CRITICAL RULES:
1. Make the FEWEST possible changes
2. PREFER DELETION over addition
3. SIMPLIFY verbose language
4. REMOVE unnecessary words/phrases
5. Only change what's essential for institutional compliance

NDA TEXT:
{redacted_text}

INSTITUTIONAL GOALS (achieve with minimal changes):
- Limit seller liability
- Define confidential information clearly
- Allow advisor disclosure
- Create enforceable remedies

APPROACH:
- DELETE verbose language first
- SIMPLIFY complex sentences second
- ENHANCE only if absolutely necessary

Return ONLY JSON array (maximum 8 changes):
[
  {{
    "type": "delete",
    "find": "unnecessary phrase to remove",
    "replace": "",
    "reason": "removes verbose language"
  }},
  {{
    "type": "simplify", 
    "find": "complex sentence",
    "replace": "simplified version",
    "reason": "clarifies meaning"
  }}
]

Focus on DELETION and SIMPLIFICATION, not enhancement."""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.1
            )
            
            content = response.choices[0].message.content
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = content[json_start:json_end]
                return json.loads(json_str)
            else:
                return []
                
        except Exception as e:
            print(f"AI error: {e}")
            return []
    
    def create_redlined_version(self, original_doc_path, changes, output_path):
        """Create redlined version with minimal changes"""
        doc = Document(original_doc_path)
        redlined_doc = Document()
        
        # Add header
        header = redlined_doc.add_paragraph()
        header_run = header.add_run("MINIMAL CHANGES REDLINED NDA")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = redlined_doc.add_paragraph()
        explanation.add_run(f"🔴 Red strikethrough = Deletions | 🟢 Green underline = Simplifications | Total: {len(changes)} minimal changes")
        redlined_doc.add_paragraph("=" * 80)
        
        # Process paragraphs
        for para in doc.paragraphs:
            new_para = redlined_doc.add_paragraph()
            para_text = para.text
            
            # Apply changes
            for change in changes:
                if change['find'] in para_text:
                    parts = para_text.split(change['find'])
                    
                    if parts[0]:
                        new_para.add_run(parts[0])
                    
                    # Show deletion
                    if change['find']:
                        del_run = new_para.add_run(change['find'])
                        del_run.font.strike = True
                        del_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Show replacement (if any)
                    if change['replace']:
                        ins_run = new_para.add_run(change['replace'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
                    
                    if len(parts) > 1 and parts[1]:
                        new_para.add_run(parts[1])
                    
                    para_text = ""
                    break
            
            if para_text:
                new_para.add_run(para_text)
        
        redlined_doc.save(output_path)
        return output_path
    
    def create_clean_version(self, original_doc_path, changes, output_path):
        """Create clean version with changes applied"""
        doc = Document(original_doc_path)
        clean_doc = Document()
        
        for para in doc.paragraphs:
            new_para = clean_doc.add_paragraph()
            para_text = para.text
            
            # Apply changes
            for change in changes:
                if change['find'] in para_text:
                    para_text = para_text.replace(change['find'], change['replace'])
            
            new_para.add_run(para_text)
        
        clean_doc.save(output_path)
        return output_path
    
    def restore_personal_info(self, doc_path):
        """Restore personal information"""
        doc = Document(doc_path)
        
        for para in doc.paragraphs:
            for placeholder, original in self.personal_info.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, original)
        
        doc.save(doc_path)
        return doc_path
    
    def process_nda(self, input_path):
        """Process with minimal changes approach"""
        print(f"🔄 Processing: {os.path.basename(input_path)}")
        
        # Extract text
        doc = Document(input_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Step 1: Redact
        print("[█░░░] Step 1: Redacting personal info...")
        redacted_text = self.redact_personal_info(text)
        print(f"   ✓ Redacted {len(self.personal_info)} items")
        
        # Step 2: Get minimal changes
        print("[██░░] Step 2: Getting minimal AI changes...")
        changes = self.get_minimal_changes(redacted_text)
        print(f"   ✓ AI proposed {len(changes)} minimal changes")
        
        # Show change summary
        deletions = len([c for c in changes if c.get('type') == 'delete'])
        simplifications = len([c for c in changes if c.get('type') == 'simplify'])
        print(f"   • {deletions} deletions, {simplifications} simplifications")
        
        # Step 3: Create versions
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        print("[███░] Step 3: Creating redlined version...")
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_minimal_redlined.docx"
        self.create_redlined_version(input_path, changes, redlined_path)
        
        print("[████] Step 4: Creating clean version...")
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_minimal_clean-version.docx"
        self.create_clean_version(input_path, changes, clean_path)
        
        # Step 5: Restore personal info
        print("[████] Step 5: Restoring personal info...")
        self.restore_personal_info(redlined_path)
        self.restore_personal_info(clean_path)
        
        print("✅ MINIMAL CHANGES COMPLETE!")
        print(f"📝 Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Total Changes: {len(changes)} (focused on deletion/simplification)")
        
        return redlined_path, clean_path

def main():
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 minimal_changes_redliner.py input.docx")
        sys.exit(1)
    
    processor = MinimalChangesRedliner()
    processor.process_nda(sys.argv[1])

if __name__ == "__main__":
    main()
