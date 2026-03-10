#!/usr/bin/env python3
"""
Surgical Word-Level Implementation Engine
Makes precise word/phrase changes like attorneys do, not paragraph replacements
"""

import os
import sys
import subprocess
import json
import re
from docx import Document
from docx.shared import RGBColor

class SurgicalImplementationEngine:
    def __init__(self):
        self.personal_info = {}
    
    def load_ai_recommendations(self, recommendations_file):
        """Load AI recommendations from markdown file"""
        try:
            with open(recommendations_file, 'r') as f:
                content = f.read()
            
            # Find the JSON array more carefully
            lines = content.split('\n')
            json_lines = []
            in_json = False
            
            for line in lines:
                if line.strip() == '```json' or line.strip().startswith('['):
                    in_json = True
                    if line.strip().startswith('['):
                        json_lines.append(line)
                    continue
                elif line.strip() == '```' and in_json:
                    break
                elif in_json:
                    json_lines.append(line)
            
            if json_lines:
                json_str = '\n'.join(json_lines)
                if ']' in json_str:
                    json_str = json_str[:json_str.rfind(']') + 1]
                
                recommendations = json.loads(json_str)
                print(f"   ✓ Loaded {len(recommendations)} AI recommendations")
                return recommendations
            else:
                print("   ❌ No JSON array found in recommendations file")
                return []
        except Exception as e:
            print(f"   ❌ Error loading recommendations: {e}")
            return []
    
    def convert_to_paragraphs(self, original_path):
        """Convert document to paragraph-based format for surgical editing"""
        # Use LibreOffice to convert to text
        output_dir = "/tmp"
        cmd = [
            'libreoffice', 
            '--headless', 
            '--convert-to', 'txt',
            '--outdir', output_dir,
            original_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(original_path))[0]
            txt_file = os.path.join(output_dir, f"{base_name}.txt")
            
            if os.path.exists(txt_file):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Split into paragraphs
                paragraphs = []
                lines = content.split('\n')
                current_para = []
                para_number = 1
                
                for line in lines:
                    if line.strip():
                        current_para.append(line.strip())
                    else:
                        if current_para:
                            para_text = ' '.join(current_para)
                            paragraphs.append({
                                'paragraph_number': para_number,
                                'original_text': para_text,
                                'current_text': para_text,
                                'changes': []
                            })
                            para_number += 1
                            current_para = []
                
                # Add final paragraph if exists
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append({
                        'paragraph_number': para_number,
                        'original_text': para_text,
                        'current_text': para_text,
                        'changes': []
                    })
                
                print(f"   ✓ Converted to {len(paragraphs)} paragraphs")
                return paragraphs
        
        return []
    
    def find_best_paragraph_match(self, paragraphs, target_line_num, find_text):
        """Find the best paragraph that contains the text to change"""
        # First try to match by approximate line number
        if target_line_num <= len(paragraphs):
            target_para = paragraphs[target_line_num - 1]
            if find_text[:50] in target_para['current_text']:
                return target_line_num - 1
        
        # Search nearby paragraphs
        search_range = 3
        start = max(0, target_line_num - search_range - 1)
        end = min(len(paragraphs), target_line_num + search_range)
        
        for i in range(start, end):
            if find_text[:50] in paragraphs[i]['current_text']:
                return i
        
        # Last resort: search all paragraphs
        for i, para in enumerate(paragraphs):
            if find_text[:30] in para['current_text']:
                return i
        
        return None
    
    def make_surgical_change(self, paragraph, find_text, replace_text, change_reason):
        """Make surgical word-level change within paragraph"""
        current_text = paragraph['current_text']
        
        # Try exact match first
        if find_text in current_text:
            new_text = current_text.replace(find_text, replace_text, 1)  # Only replace first occurrence
            paragraph['current_text'] = new_text
            paragraph['changes'].append({
                'type': 'replace',
                'find': find_text,
                'replace': replace_text,
                'reason': change_reason
            })
            return True
        
        # Try to find key phrases for partial replacement
        find_words = find_text.split()
        if len(find_words) > 5:
            # Try first 5 words
            partial_find = ' '.join(find_words[:5])
            if partial_find in current_text:
                # Find the sentence containing this phrase
                sentences = current_text.split('.')
                for i, sentence in enumerate(sentences):
                    if partial_find in sentence:
                        # Replace just the key phrase, not the whole sentence
                        key_phrase = self.extract_key_phrase(sentence, partial_find)
                        if key_phrase:
                            replacement_phrase = self.extract_key_phrase(replace_text, partial_find)
                            if replacement_phrase:
                                new_text = current_text.replace(key_phrase, replacement_phrase, 1)
                                paragraph['current_text'] = new_text
                                paragraph['changes'].append({
                                    'type': 'replace',
                                    'find': key_phrase,
                                    'replace': replacement_phrase,
                                    'reason': change_reason
                                })
                                return True
        
        return False
    
    def extract_key_phrase(self, text, anchor):
        """Extract a meaningful phrase around the anchor text"""
        words = text.split()
        anchor_words = anchor.split()
        
        # Find where anchor starts in the text
        for i in range(len(words) - len(anchor_words) + 1):
            if ' '.join(words[i:i+len(anchor_words)]) == anchor:
                # Extract phrase around this location (±3 words)
                start = max(0, i - 3)
                end = min(len(words), i + len(anchor_words) + 3)
                return ' '.join(words[start:end])
        
        return None
    
    def implement_surgical_changes(self, paragraphs, recommendations):
        """Implement AI recommendations with surgical precision"""
        implemented_count = 0
        
        print(f"   🔄 Making surgical changes for {len(recommendations)} recommendations...")
        
        for rec in recommendations:
            try:
                if rec.get('change_type') == 'replace_text_on_line':
                    line_num = rec.get('line_number')
                    find_text = rec.get('find_text', '')
                    replace_text = rec.get('replace_with', '')
                    reason = rec.get('reason', '')
                    
                    # Find the best matching paragraph
                    para_index = self.find_best_paragraph_match(paragraphs, line_num, find_text)
                    
                    if para_index is not None:
                        if self.make_surgical_change(paragraphs[para_index], find_text, replace_text, reason):
                            implemented_count += 1
                            print(f"   ✓ Made surgical change in paragraph {para_index + 1}")
                        else:
                            print(f"   ⚠️ Could not make surgical change in paragraph {para_index + 1}")
                    else:
                        print(f"   ⚠️ Could not find target paragraph for line {line_num}")
                
                elif rec.get('change_type') == 'insert_after_line':
                    line_num = rec.get('line_number')
                    insert_text = rec.get('insert_text', '')
                    
                    if line_num <= len(paragraphs):
                        # Insert new paragraph
                        new_para = {
                            'paragraph_number': None,
                            'original_text': '',
                            'current_text': insert_text,
                            'changes': [{
                                'type': 'insert',
                                'text': insert_text,
                                'reason': rec.get('reason', '')
                            }]
                        }
                        paragraphs.insert(line_num, new_para)
                        implemented_count += 1
                        print(f"   ✓ Inserted new paragraph after line {line_num}")
                
                elif rec.get('change_type') == 'insert_at_end':
                    insert_text = rec.get('insert_text', '')
                    new_para = {
                        'paragraph_number': None,
                        'original_text': '',
                        'current_text': insert_text,
                        'changes': [{
                            'type': 'insert',
                            'text': insert_text,
                            'reason': rec.get('reason', '')
                        }]
                    }
                    paragraphs.append(new_para)
                    implemented_count += 1
                    print(f"   ✓ Inserted paragraph at end")
                    
            except Exception as e:
                print(f"   ❌ Failed to implement {rec.get('problem_addressed', 'unknown')}: {e}")
        
        print(f"   ✓ Made {implemented_count}/{len(recommendations)} surgical changes")
        return paragraphs, implemented_count
    
    def create_surgical_redlined_docx(self, paragraphs, output_path):
        """Create redlined document showing surgical word-level changes"""
        doc = Document()
        
        # Add header
        header = doc.add_paragraph()
        header_run = header.add_run("SURGICAL WORD-LEVEL REDLINES")
        header_run.bold = True
        header_run.font.size = 16
        
        explanation = doc.add_paragraph()
        explanation.add_run("🔴 Red strikethrough = Deleted words/phrases | 🟢 Green underline = Added words/phrases")
        doc.add_paragraph("=" * 80)
        
        # Process each paragraph
        for para_data in paragraphs:
            doc_para = doc.add_paragraph()
            
            if para_data['changes']:
                # This paragraph has changes - show surgical edits
                for change in para_data['changes']:
                    if change['type'] == 'replace':
                        # Show the surgical replacement
                        original_text = para_data['original_text']
                        find_text = change['find']
                        replace_text = change['replace']
                        
                        # Split around the change
                        if find_text in original_text:
                            parts = original_text.split(find_text)
                            
                            # Before the change
                            if parts[0]:
                                doc_para.add_run(parts[0])
                            
                            # Show deletion (strikethrough)
                            del_run = doc_para.add_run(find_text)
                            del_run.font.strike = True
                            del_run.font.color.rgb = RGBColor(255, 0, 0)
                            
                            # Show addition (underlined)
                            ins_run = doc_para.add_run(replace_text)
                            ins_run.underline = True
                            ins_run.font.color.rgb = RGBColor(0, 128, 0)
                            
                            # After the change
                            if len(parts) > 1 and parts[1]:
                                doc_para.add_run(parts[1])
                        else:
                            # Fallback: show current text
                            doc_para.add_run(para_data['current_text'])
                    
                    elif change['type'] == 'insert':
                        # Show insertion
                        label_run = doc_para.add_run("[INSERTED] ")
                        label_run.bold = True
                        label_run.font.color.rgb = RGBColor(0, 100, 0)
                        
                        ins_run = doc_para.add_run(change['text'])
                        ins_run.underline = True
                        ins_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                # Unchanged paragraph
                if para_data['current_text']:
                    doc_para.add_run(para_data['current_text'])
        
        doc.save(output_path)
        print(f"   ✓ Created surgical redlined document: {os.path.basename(output_path)}")
    
    def create_clean_docx(self, paragraphs, output_path):
        """Create clean document with surgical changes applied"""
        doc = Document()
        
        for para_data in paragraphs:
            if para_data['current_text']:
                para = doc.add_paragraph()
                para.add_run(para_data['current_text'])
        
        doc.save(output_path)
        print(f"   ✓ Created clean document: {os.path.basename(output_path)}")
    
    def process_surgical_implementation(self, original_path, recommendations_file):
        """Complete surgical implementation workflow"""
        print(f"🔄 Making surgical changes for: {os.path.basename(original_path)}")
        
        # Step 1: Load AI recommendations
        print("[█░░░░] Step 1: Loading AI recommendations...")
        recommendations = self.load_ai_recommendations(recommendations_file)
        if not recommendations:
            return None, None
        
        # Step 2: Convert to paragraph format
        print("[██░░░] Step 2: Converting to paragraph format...")
        paragraphs = self.convert_to_paragraphs(original_path)
        if not paragraphs:
            return None, None
        
        # Step 3: Make surgical changes
        print("[███░░] Step 3: Making surgical word-level changes...")
        modified_paragraphs, implemented_count = self.implement_surgical_changes(paragraphs, recommendations)
        
        # Step 4: Create output documents
        print("[████░] Step 4: Creating surgical redlined documents...")
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        redlined_path = f"/home/cliff/redact/redline_project/{base_name}_surgical_redlined.docx"
        clean_path = f"/home/cliff/redact/redline_project/{base_name}_surgical_clean.docx"
        
        self.create_surgical_redlined_docx(modified_paragraphs, redlined_path)
        self.create_clean_docx(modified_paragraphs, clean_path)
        
        print("[█████] Step 5: Complete!")
        print("✅ SURGICAL IMPLEMENTATION COMPLETE!")
        print(f"📝 Surgical Redlined: {os.path.basename(redlined_path)}")
        print(f"📄 Clean: {os.path.basename(clean_path)}")
        print(f"🎯 Made {implemented_count}/{len(recommendations)} surgical changes")
        
        return redlined_path, clean_path

def main():
    if len(sys.argv) != 3:
        print("Usage: python3 surgical_implementation_engine.py original.docx recommendations.md")
        sys.exit(1)
    
    original_path = sys.argv[1]
    recommendations_file = sys.argv[2]
    
    engine = SurgicalImplementationEngine()
    engine.process_surgical_implementation(original_path, recommendations_file)

if __name__ == "__main__":
    main()
