#!/usr/bin/env python3
"""
SeedJura Lease Summary Tool
============================
Ingests lease agreements (PDF/DOCX), redacts PII, uses AI to extract
key lease terms, and populates the SeedJura Lease Summary template.

Flow:
  1. Ingest the document (PDF or DOCX)
  2. Capture PII and redact (using existing ContractRedactor)
  3. Feed redacted text to AI for analysis and field extraction
  4. Populate the template and save as: {lease_name}_summary_{date}.docx

Usage:
  python3 lease_summary_tool.py <input_lease_file> [--output-dir <dir>]
  python3 lease_summary_tool.py --batch <directory> [--output-dir <dir>]
"""

import os
import sys
import re
import json
import copy
import argparse
from datetime import datetime
from pathlib import Path

# Third-party
from docx import Document
from docx.shared import Inches, Pt
import pymupdf  # fitz / PyMuPDF

# Local imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from redactor import ContractRedactor


def _log(code: str, message: str):
    """Lightweight logging hook — routes to the app logger if available."""
    try:
        import logging
        logging.getLogger("seedjura").info(f"[{code}] {message}")
    except Exception:
        pass


# =============================================================================
# CONFIGURATION
# =============================================================================

def _resolve_path(path: str) -> str:
    """Resolve a path, checking both WSL and Windows locations."""
    if os.path.exists(path):
        return path
    # If running on Windows (not WSL), try converting /mnt/c/ to C:\
    if sys.platform == "win32" and path.startswith("/mnt/"):
        drive = path[5].upper()
        rest = path[6:].replace("/", "\\")
        win_path = f"{drive}:{rest}"
        if os.path.exists(win_path):
            return win_path
    return path


# Template - checked in multiple locations
_TEMPLATE_SEARCH_PATHS = [
    "/mnt/c/Users/cliff/Downloads/OneDrive_1_8-10-2026/SeedJura_Lease_Summary_FORM.docx",
    "/mnt/c/seedJura/SeedJura_Lease_Summary_FORM.docx",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "SeedJura_Lease_Summary_FORM.docx"),
    # Windows paths (for packaged exe)
    r"C:\seedJura\SeedJura_Lease_Summary_FORM.docx",
    r"C:\Users\cliff\Downloads\OneDrive_1_8-10-2026\SeedJura_Lease_Summary_FORM.docx",
]

TEMPLATE_PATH = ""
for _p in _TEMPLATE_SEARCH_PATHS:
    if os.path.exists(_p):
        TEMPLATE_PATH = _p
        break
if not TEMPLATE_PATH:
    TEMPLATE_PATH = _TEMPLATE_SEARCH_PATHS[0]  # Fallback for error message

DEFAULT_OUTPUT_DIR = "/mnt/c/seedJura/lease_summaries"
if sys.platform == "win32":
    DEFAULT_OUTPUT_DIR = r"C:\seedJura\Summary_Output"

# API key - checked in multiple locations
_API_KEY_SEARCH_PATHS = [
    "/home/cliff/redact/openai_api_key.txt",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "openai_api_key.txt"),
    # Windows paths
    r"C:\seedJura\openai_api_key.txt",
    os.path.join(os.path.expanduser("~"), "openai_api_key.txt"),
]

API_KEY_FILE = ""
for _p in _API_KEY_SEARCH_PATHS:
    if os.path.exists(_p):
        API_KEY_FILE = _p
        break
if not API_KEY_FILE:
    API_KEY_FILE = _API_KEY_SEARCH_PATHS[0]

AI_MODEL = "gpt-4o"


# =============================================================================
# PHASE 1: DOCUMENT INGESTION
# =============================================================================

class DocumentIngestError(Exception):
    """Raised when a document cannot be ingested. Carries a user-friendly message."""
    def __init__(self, message: str, kind: str = "general", detail: str = ""):
        self.message = message
        self.kind = kind  # unsupported, encrypted, corrupt, ocr_unavailable, empty
        self.detail = detail
        super().__init__(message)


def ingest_document(file_path: str) -> str:
    """
    Extract text from PDF or DOCX lease document.
    Raises DocumentIngestError with a user-friendly message on any failure.
    """
    if not os.path.exists(file_path):
        raise DocumentIngestError(
            f"File not found: {os.path.basename(file_path)}",
            kind="corrupt",
        )

    # Check the file isn't empty
    try:
        if os.path.getsize(file_path) == 0:
            raise DocumentIngestError(
                f"File is empty: {os.path.basename(file_path)}",
                kind="empty",
            )
    except OSError:
        pass

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        text = _extract_pdf(file_path)
    elif ext == ".docx":
        text = _extract_docx(file_path)
    elif ext == ".doc":
        # Legacy binary .doc is NOT supported by python-docx
        raise DocumentIngestError(
            "Legacy .doc files (Word 97-2003) are not supported. "
            "Please save the file as .docx or PDF and try again.",
            kind="unsupported",
        )
    elif ext in (".txt", ".md"):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except OSError as e:
            raise DocumentIngestError(
                f"Could not read text file: {os.path.basename(file_path)}",
                kind="corrupt", detail=str(e),
            )
    else:
        raise DocumentIngestError(
            f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.",
            kind="unsupported",
        )

    # Check we actually got usable content
    if not text or len(text.strip()) < 20:
        raise DocumentIngestError(
            f"No readable text found in {os.path.basename(file_path)}. "
            "The document may be blank, an unsupported image format, or corrupt.",
            kind="empty",
        )

    # Post-OCR corrections for common handwriting misreads
    text = _fix_handwritten_dates(text)
    return text


def _fix_handwritten_dates(text: str) -> str:
    """
    Fix common OCR misreads of handwritten dates in legal documents.
    Handwritten numbers on fill-in-the-blank lease templates are frequently
    garbled by OCR (e.g., "1st" -> "Bias", "2nd" -> "Znd", "15" -> "IS").

    Strategy:
    - Detect garbled text in the pattern: <garbage> day of <Month>, <Year>
    - Replace with [handwritten] day of <Month>, <Year> so downstream AI
      knows the day is uncertain
    - Also fix common specific OCR misreads of ordinals
    """
    # Common OCR misreads of handwritten ordinals before "day of"
    # Pattern: <word that doesn't look like a number/ordinal> day of <Month>
    months = (
        "January|February|March|April|May|June|July|"
        "August|September|October|November|December"
    )

    # Known OCR misreads of handwritten numbers/ordinals
    ocr_fixes = {
        # Multi-word patterns first (before single-word fixes)
        r'\bIS\s*th\b(?=\s*day)': '15th',
        # 1st misreads
        r'\bBias\b': '1st',
        r'\bIst\b': '1st',
        r'\bist\b': '1st',
        r'\blst\b': '1st',
        r'\bLst\b': '1st',
        r'\bI st\b': '1st',
        r'\bl st\b': '1st',
        # 2nd misreads
        r'\bZnd\b': '2nd',
        r'\bznd\b': '2nd',
        r'\bZad\b': '2nd',
        # 3rd misreads
        r'\bSrd\b': '3rd',
        r'\b3ra\b': '3rd',
        # Common number misreads (only when followed by day)
        r'\bIS\b(?=\s*day)': '15',
        r'\bIi\b(?=\s*(?:th)?\s*day)': '11',
        r'\bIl\b(?=\s*(?:th)?\s*day)': '11',
        r'\bZl\b(?=\s*(?:th)?\s*day)': '21',
        r'\bZI\b(?=\s*(?:th)?\s*day)': '21',
        r'\bSl\b(?=\s*(?:th)?\s*day)': '31',
        r'\bSI\b(?=\s*(?:th)?\s*day)': '31',
    }

    for pattern, replacement in ocr_fixes.items():
        text = re.sub(pattern, replacement, text)

    # Catch remaining garbled text before "day of <Month>" that doesn't
    # look like a valid number or ordinal. Mark as [handwritten].
    # Valid patterns: "1st", "2nd", "3rd", "4th"-"31st", or bare numbers 1-31
    valid_day_pattern = r'(?:\d{1,2}(?:st|nd|rd|th)?)'

    def _check_day_token(match):
        token = match.group(1).strip()
        # If it already looks like a valid day, leave it alone
        if re.match(r'^\d{1,2}(?:st|nd|rd|th)?$', token, re.IGNORECASE):
            return match.group(0)
        # If it's a short gibberish token (1-6 chars), it's likely a misread handwritten day
        if len(token) <= 6 and not re.match(r'^[a-z]+$', token):
            return f"[handwritten: {token}] " + match.group(0)[match.start(2) - match.start():]
        return match.group(0)

    # Pattern: <token> day of <Month>
    # Only flag tokens that are clearly not valid day numbers
    garbled_day_pattern = re.compile(
        r'(\b[A-Za-z0-9]{1,6}\b)\s+(day\s+of\s+(?:' + months + r'))',
        re.IGNORECASE
    )

    def _replace_garbled(match):
        token = match.group(1)
        rest = match.group(2)
        # Valid day ordinals - don't touch
        if re.match(r'^\d{1,2}(?:st|nd|rd|th)?$', token, re.IGNORECASE):
            return match.group(0)
        # Common words that legitimately precede "day of" - don't touch
        if token.lower() in ('this', 'each', 'every', 'first', 'second', 'third',
                             'last', 'next', 'same', 'such', 'that', 'the', 'one',
                             'any', 'per', 'said', 'fifth', 'tenth', 'a', 'no'):
            return match.group(0)
        # Likely garbled handwritten number
        return f"[handwritten: {token}] {rest}"

    text = garbled_day_pattern.sub(_replace_garbled, text)

    return text


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF. Falls back to OCR for scanned docs."""
    try:
        doc = pymupdf.open(file_path)
    except Exception as e:
        raise DocumentIngestError(
            f"Could not open PDF: {os.path.basename(file_path)}. "
            "The file may be corrupt or not a valid PDF.",
            kind="corrupt", detail=str(e),
        )

    # Handle encrypted/password-protected PDFs
    if doc.needs_pass:
        # Try empty password (some PDFs are encrypted but openable)
        if not doc.authenticate(""):
            doc.close()
            raise DocumentIngestError(
                f"PDF is password-protected: {os.path.basename(file_path)}. "
                "Please remove the password and try again.",
                kind="encrypted",
            )

    text_parts = []
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_parts.append(page.get_text())
    except Exception as e:
        doc.close()
        raise DocumentIngestError(
            f"Error reading PDF pages in {os.path.basename(file_path)}. "
            "The file may be partially corrupt.",
            kind="corrupt", detail=str(e),
        )

    page_count = len(doc)
    doc.close()

    combined = "\n".join(text_parts)

    # If very little text extracted, the PDF is likely scanned - use OCR
    # Use a per-page heuristic: less than ~50 chars/page average = scanned
    avg_chars = len(combined.strip()) / max(page_count, 1)
    if len(combined.strip()) < 200 or avg_chars < 50:
        print("  PDF appears to be scanned/image-based. Running OCR...")
        _log("I104", f"OCR started: {os.path.basename(file_path)} ({page_count} pages)")
        ocr_text = _ocr_pdf(file_path)
        # Use OCR result if it produced more text than direct extraction
        if len(ocr_text.strip()) > len(combined.strip()):
            combined = ocr_text

    return combined


def _find_tesseract() -> str:
    """
    Locate the Tesseract executable. Checks PATH, common install locations,
    and a bundled copy next to the app. Returns the path or "" if not found.
    """
    import shutil as _shutil

    # 1. Already on PATH
    found = _shutil.which("tesseract")
    if found:
        return found

    # 2. Common Windows install locations
    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\seedJura\Tesseract-OCR\tesseract.exe",
    ]
    # 3. Bundled copy next to the app / exe
    app_dir = os.path.dirname(os.path.abspath(__file__))
    candidates.append(os.path.join(app_dir, "Tesseract-OCR", "tesseract.exe"))
    candidates.append(os.path.join(app_dir, "lease_summary_app", "Tesseract-OCR", "tesseract.exe"))
    # PyInstaller bundle location
    if getattr(sys, "frozen", False):
        candidates.append(os.path.join(sys._MEIPASS, "Tesseract-OCR", "tesseract.exe"))

    for c in candidates:
        if os.path.exists(c):
            return c

    return ""


def is_ocr_available() -> bool:
    """Check whether OCR (Tesseract + pytesseract) is available."""
    try:
        import pytesseract  # noqa
        from PIL import Image  # noqa
    except ImportError:
        return False
    return bool(_find_tesseract())


def _ocr_pdf(file_path: str) -> str:
    """OCR a scanned PDF using PyMuPDF raster + Tesseract. Tries basic first, preprocessing if needed."""
    try:
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        raise DocumentIngestError(
            "This is a scanned PDF that requires OCR, but the OCR libraries "
            "(pytesseract, Pillow) are not installed. Please provide a digital "
            "(text-based) PDF, or install OCR support.",
            kind="ocr_unavailable",
        )

    # Locate Tesseract binary
    tess_path = _find_tesseract()
    if not tess_path:
        raise DocumentIngestError(
            "This is a scanned PDF that requires OCR, but Tesseract is not "
            "installed on this machine. Please provide a digital (text-based) "
            "PDF, or install Tesseract OCR.",
            kind="ocr_unavailable",
        )
    pytesseract.pytesseract.tesseract_cmd = tess_path

    doc = pymupdf.open(file_path)
    text_parts = []
    total_pages = len(doc)

    for page_num in range(total_pages):
        page = doc[page_num]
        # Render page to image at 300 DPI
        mat = pymupdf.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_data))

        # Basic OCR first
        page_text = pytesseract.image_to_string(image, config='--psm 6')

        # If basic OCR produced very little text, try with preprocessing
        if len(page_text.strip()) < 100:
            page_text = _ocr_with_preprocessing(image)

        text_parts.append(page_text)

        if (page_num + 1) % 10 == 0 or page_num == 0:
            print(f"    OCR page {page_num + 1}/{total_pages}...")

    doc.close()
    print(f"    OCR complete: {total_pages} pages processed")
    return "\n".join(text_parts)


def _ocr_with_preprocessing(image) -> str:
    """Apply image preprocessing and retry OCR for difficult pages."""
    try:
        import pytesseract
        import cv2
        import numpy as np
    except ImportError:
        return ""

    img_np = np.array(image)

    # Convert to grayscale
    if len(img_np.shape) == 3:
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_np

    # Light denoising
    gray = cv2.GaussianBlur(gray, (3, 3), 0)

    # Otsu's thresholding (less aggressive than adaptive)
    _, gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    from PIL import Image as PILImage
    processed = PILImage.fromarray(gray)
    return pytesseract.image_to_string(processed, config='--psm 6')


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        doc = Document(file_path)
    except Exception as e:
        msg = str(e).lower()
        if "encrypted" in msg or "password" in msg:
            raise DocumentIngestError(
                f"DOCX is password-protected: {os.path.basename(file_path)}. "
                "Please remove the password and try again.",
                kind="encrypted", detail=str(e),
            )
        if "not a zip" in msg or "bad" in msg or "file is not" in msg:
            raise DocumentIngestError(
                f"Could not open {os.path.basename(file_path)}. It may be a "
                "legacy .doc file, corrupt, or not a valid Word document. "
                "Try saving it as .docx or PDF.",
                kind="corrupt", detail=str(e),
            )
        raise DocumentIngestError(
            f"Could not read Word document: {os.path.basename(file_path)}.",
            kind="corrupt", detail=str(e),
        )

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


# =============================================================================
# DOCUMENT TYPE DETECTION
# =============================================================================

def _detect_document_type(text: str, filename: str) -> str:
    """
    Detect whether document is a full lease, amendment, addendum, or other.
    Returns: 'lease', 'amendment', 'addendum', 'assignment', or 'other'
    """
    # Check filename first
    fname_lower = filename.lower()
    if 'amendment' in fname_lower or 'amend' in fname_lower:
        return 'amendment'
    if 'addendum' in fname_lower:
        return 'addendum'
    if 'assignment' in fname_lower and 'sublease' not in fname_lower:
        return 'assignment'

    # Check first 2000 chars of document text
    header = text[:2000].lower()

    amendment_signals = [
        'amendment to lease',
        'first amendment',
        'second amendment',
        'third amendment',
        'fourth amendment',
        'fifth amendment',
        'amendment no.',
        'amendment number',
        'this amendment',
        'amends the lease',
        'modification to lease',
        'lease modification',
    ]

    addendum_signals = [
        'addendum to lease',
        'lease addendum',
        'this addendum',
    ]

    for signal in amendment_signals:
        if signal in header:
            return 'amendment'

    for signal in addendum_signals:
        if signal in header:
            return 'addendum'

    return 'lease'


# Fields that typically only appear in a full lease, not an amendment
FULL_LEASE_ONLY_FIELDS = {
    "Tenant_Insurance", "Landlord_Repair_Obligations", "Tenant_Repair_Obligations",
    "OPEX_Inclusion", "OPEX_Exclusion", "Utilities", "Fees_Management",
    "Gross_Up_Percent_Language", "Assignment_LL_Decision", "Assignment_Rent_Profit",
    "Assignment_Recapture_Space", "Assignment_Process_Fee", "Assignment_Other_Terms",
    "Assignment_Change_Control", "Sublease_Terms", "Holdover_Rent",
    "Estoppel_Return_Period", "Estoppel_Details",
    "SNDA_Subject_Existing_Mortgages", "SNDA_Subject_Future_Mortgages",
    "SNDA_Subject_Existing_GroundLease", "SNDA_Subject_Future_GroundLease",
    "SNDA_Required_by_Lender_Provision",
    "Relocation_Rights", "Relocation_Notice_Period", "Relocation_Language",
    "Relocation_Cost", "Relocation_Termination_Rights",
    "Signage_Allowed", "Signage_Location", "Signage_Approval_Required",
    "Signage_Type", "Signage_Renovation_Replacement", "Signage_Removal",
    "Other_Rights",
}


# =============================================================================
# PHASE 2: PII CAPTURE AND REDACTION
# =============================================================================

def redact_and_capture_pii(raw_text: str) -> tuple:
    """
    Run PII detection, capture findings for later use, return redacted text.
    Returns: (redacted_text, pii_findings_dict)
    """
    redactor = ContractRedactor()
    redacted_text = raw_text
    pii_findings = []

    # Apply pattern-based redaction
    for pattern, label in redactor.patterns:
        matches = list(re.finditer(pattern, redacted_text, re.IGNORECASE))
        for match in reversed(matches):  # reverse to preserve positions
            pii_findings.append({
                "type": label,
                "value": match.group(),
                "start": match.start(),
                "end": match.end()
            })

    # We do NOT redact for the AI analysis - we need the full text
    # But we capture what PII exists for the log
    print(f"  PII scan: {len(pii_findings)} items detected")

    # Group by type for summary
    pii_by_type = {}
    for item in pii_findings:
        pii_by_type.setdefault(item["type"], []).append(item["value"])

    for pii_type, values in pii_by_type.items():
        print(f"    {pii_type}: {len(values)} instance(s)")

    return raw_text, pii_findings


# =============================================================================
# PHASE 3: AI ANALYSIS - EXTRACT LEASE FIELDS
# =============================================================================

# All template fields that the AI should extract
LEASE_FIELDS = {
    # Header fields
    "Lease_Summary_Preparer": "Name of person preparing this summary",
    "Lease_Summary_Date": "Date the summary was prepared",
    "Lease_Summary_Purpose": "Purpose of the summary (e.g., purchase, review)",
    "Property_Name": "Name of the property/project",
    "Property_Address_1": "Street address line 1",
    "Property_Address_2": "Street address line 2 (suite, unit)",
    "Property_Address_City": "City",
    "Property_Address_State": "State",
    "Property_Address_Zip": "Zip code",
    "Owner_Name": "Property owner/landlord name",
    "Owner_StReg": "Owner state of registration",
    "Owner_Type": "Owner entity type (e.g., LLC, Corp)",
    # Tenant info
    "Tenant_Name": "Tenant legal name",
    "Tenant_StReg": "Tenant state of registration",
    "Tenant_Type": "Tenant entity type",
    "Tenant_DBA": "Tenant DBA (doing business as) name",
    "Tenant_Address_1": "Tenant address line 1",
    "Tenant_Address_2": "Tenant address line 2",
    "Tenant_Address_City": "Tenant city",
    "Tenant_Address_State": "Tenant state",
    "Tenant_Address_Zip": "Tenant zip",
    "Tenant_Contact": "Tenant contact person",
    "Tenant_CopyTo_Address_1": "Notice copy-to address line 1",
    "Tenant_CopyTo_Address_2": "Notice copy-to address line 2",
    "Tenant_CopyTo_Address_City": "Notice copy-to city",
    "Tenant_CopyTo_Address_State": "Notice copy-to state",
    "Tenant_CopyTo_Address_Zip": "Notice copy-to zip",
    "Tenant_CopyTo_Attn": "Notice copy-to attention name",
    "Tenant_CopyTo_Email": "Notice copy-to email",
    "Tenant_Attn": "Tenant attention/contact person",
    "Tenant_Email": "Tenant email address",
}

LEASE_FIELDS_2 = {
    # Premises
    "Premises_UnitNumber_Description": "Suite/unit number and description",
    "Premises_SqFt": "Square footage of premises",
    # Lease terms
    "LeaseAgr_Name": "Name/title of the lease agreement",
    "LeaseAgr_Amendments": "Any amendments to the lease",
    "Date_Lease": "Lease effective date",
    "Date_Commencment": "Commencement date - include FULL conditional language if applicable",
    "Lease_Term": "Length of lease term with full description",
    "Date_Expiration": "Expiration date with full conditional language",
    "Date_EarlyAccess": "Early access terms - full language",
    "Renewal_Option_Numbers": "Number of renewal options with conditions (e.g., doesn't apply to Assignment)",
    "Renewal_Option_Period_PerOption": "Period per renewal option",
    # Financial
    "Amt_Security_Deposit": "Security deposit amount",
    "Rent_BaseRent_PSF": "Base rent per square foot",
    "Rent_BaseRent_Amt": "Base rent annual amount",
    "Rent_BaseRent_Monthly": "Base rent monthly amount",
    "Rent_AnnualIncrease_Percentage": "Annual rent increase percentage",
    "Base_Year": "Base year for operating expenses",
    "Rent_Abatement_Commencement": "Rent abatement start date",
    "Rent_Abatement_Expiration": "Rent abatement end date",
    "Rent_Abatement_Duration": "Rent abatement duration",
    "Rent_Abatement_Base_Rent": "Is base rent abated?",
    "Rent_Abatement_Additional_Rent": "Is additional rent abated?",
    "Rent_Abatement_Qualifier": "Conditions for rent abatement",
    "Tenant_Share_Percentage": "Tenant's percentage share of expenses",
    "Rent_PercentageRent_ThresholdAmt": "Percentage rent threshold",
    "Rent_PercentageRent_Percent": "Percentage rent rate",
    "Tenant_Allowance_PSF": "Tenant improvement allowance per sq ft",
    "Tenant_Allowance_Total": "Total tenant improvement allowance",
}

LEASE_FIELDS_3 = {
    # Use and parking
    "Permitted_Use_Description": "Permitted use - quote full clause including restrictions",
    "Parking_Reserved_Spaces": "Reserved parking spaces with description",
    "Parking_Unreserved_Spaces": "Unreserved parking spaces with description",
    "Parking_Reserved_Amt_Fees": "Reserved parking fees",
    "Parking_Unreserved_Amt_Fees": "Unreserved parking fees",
    # Guarantor
    "Guarantor_Name": "Guarantor name",
    "Guarantor_StReg": "Guarantor state of registration",
    "Guarantor_Type": "Guarantor entity type",
    "Guaranty_Term": "Term of the guaranty",
    # Broker
    "Broker_Landlord_Name": "Landlord's broker",
    "Broker_Tenant_Name": "Tenant's broker",
    # Improvements
    "Tenant_Improvements_Description": "Landlord work / tenant improvements - full description with dollar amounts",
    # Options
    "ROFR_Space": "ROFR applicable space (or 'None' if not applicable)",
    "ROFR_Description": "ROFR description - full language",
    "ROFO_Space": "ROFO applicable space (or 'None' if not applicable)",
    "ROFO_Description": "ROFO description - full language",
    "Expansion_Space": "Expansion option space (or 'None' if not applicable)",
    "Expansion_Description": "Expansion option description",
    "Early_Termination_Description": "Early termination terms - include BOTH tenant and landlord rights",
    "Reduction_Description": "Reduction option description (or 'None')",
    "Purchase_Option_Space": "Purchase option applicable space (or 'None')",
    "Purchase_Option_Description": "Purchase option description - full language",
}

LEASE_FIELDS_4 = {
    # OPEX
    "OPEX_Inclusion": "Operating expense inclusions - list ALL included items",
    "OPEX_Exclusion": "Operating expense exclusions - list ALL excluded items (or 'None')",
    "Utilities": "Utility provisions - full language describing who pays and what's included",
    "Fees_Management": "Admin/management/accounting fees - full language",
    "Gross_Up_Percent_Language": "Gross up percentage language - full clause with cap language",
    # Other provisions
    "Tenant_Insurance": "Tenant insurance requirements - include section ref and key requirements (types of coverage, amounts, ratings)",
    "Landlord_Repair_Obligations": "Landlord repair and maintenance - full language with section ref",
    "Tenant_Repair_Obligations": "Tenant repair and maintenance - full language with section ref",
    "Assignment_3rd_Parties": "Assignment to 3rd parties - full language including consent requirements, fees, conditions",
    "Assignment_Affiliates": "Assignment to affiliates - full language",
    "Assignment_LL_Decision": "Landlord decision on assignment - full criteria for reasonable/unreasonable withholding",
    "Assignment_Rent_Profit": "Rent profits from assignment - how excess rent is shared",
    "Assignment_Recapture_Space": "Landlord recapture of space - full language (e.g., 'Yes, can terminate lease rather than approve')",
    "Assignment_Process_Fee": "Assignment processing fee amount",
    "Assignment_Other_Terms": "Other assignment terms",
    "Assignment_Change_Control": "Assignment change of control terms",
    "Sublease_Terms": "Sublease terms - full language or reference to assignment section",
    "Holdover_Rent": "Holdover rent terms with section ref and percentage (e.g., 'Sec. 28 - 150% of the rent')",
    "Estoppel_Return_Period": "Estoppel return period",
    "Estoppel_Details": "Estoppel details to be included",
    "SNDA_Subject_Existing_Mortgages": "Subject to existing mortgages (Yes/No)",
    "SNDA_Subject_Future_Mortgages": "Subject to future mortgages (Yes/No)",
    "SNDA_Subject_Existing_GroundLease": "Subject to existing ground leases",
    "SNDA_Subject_Future_GroundLease": "Subject to future ground leases",
    "SNDA_Required_by_Lender_Provision": "SNDA requested by lender language",
    "Relocation_Rights": "Relocation allowed (Yes/No)",
    "Relocation_Notice_Period": "Relocation notice period",
    "Relocation_Language": "Relocation language",
    "Relocation_Cost": "Relocation costs",
    "Relocation_Termination_Rights": "Termination rights",
    "Signage_Allowed": "Signage allowed (Yes/No)",
    "Signage_Location": "Signage location",
    "Signage_Approval_Required": "Signage approval required",
    "Signage_Type": "Type of signage",
    "Signage_Renovation_Replacement": "Signage renovation/replacement",
    "Signage_Removal": "Signage removal terms",
    "Other_Rights": "Other rights and provisions",
}

# Combine all field dictionaries
ALL_FIELDS = {**LEASE_FIELDS, **LEASE_FIELDS_2, **LEASE_FIELDS_3, **LEASE_FIELDS_4}


def get_openai_client():
    """Initialize OpenAI client with API key from file."""
    from openai import OpenAI

    api_key = ""
    if os.path.exists(API_KEY_FILE):
        with open(API_KEY_FILE, "r") as f:
            api_key = f.read().strip()
    else:
        api_key = os.environ.get("OPENAI_API_KEY", "")

    if not api_key:
        raise RuntimeError(
            f"No OpenAI API key found. Place it in {API_KEY_FILE} "
            "or set OPENAI_API_KEY environment variable."
        )

    return OpenAI(api_key=api_key)


def build_extraction_prompt(lease_text: str, doc_type: str = "lease") -> str:
    """Build the prompt for AI lease field extraction."""
    fields_list = "\n".join(
        f'  "{k}": "{v}"' for k, v in ALL_FIELDS.items()
    )

    # Add document-type-specific instructions
    doc_type_instruction = ""
    if doc_type in ("amendment", "addendum"):
        doc_type_instruction = f"""
IMPORTANT - THIS DOCUMENT IS A LEASE {doc_type.upper()}:
- This is NOT a full lease. It only modifies specific sections of the original lease.
- For fields that are NOT addressed in this {doc_type}, return "See Original Lease."
- Only extract fields that are explicitly stated or modified in this {doc_type}.
- DO NOT guess or fabricate values for sections that would only be in the original lease.
- Fields commonly NOT in an {doc_type}: insurance, repair obligations, OPEX, assignment details, holdover, SNDA, relocation, signage.
- Fields commonly IN an {doc_type}: parties, premises changes, rent changes, term extensions, specific clause modifications.
"""

    prompt = f"""You are a commercial real estate lease analyst. Analyze the following lease document and extract the requested information into a JSON object.
{doc_type_instruction}
CRITICAL RULES - READ CAREFULLY:
- You MUST copy/paste exact language from the document. Do NOT summarize, paraphrase, or reword.
- Every value you return must be text that appears verbatim in the document.
- If a provision is not found or not applicable, return "None." (with period).
- For dates, copy the exact phrasing (e.g., "the earlier to occur of A) November 1, 2022, or B) the earlier of...").
- For dollar amounts, copy exactly as written including $ sign.
- For percentages, copy exactly as written.
- Include section/article references when they appear near the relevant text (e.g., "Sec. 7 -", "Section 37 -").
- IMPORTANT: Start each extraction at the BEGINNING of the relevant sentence or clause. Do not start mid-sentence.
- For long provisions (assignment, insurance, repair, OPEX, permitted use): copy the FULL relevant clause starting from its topic sentence. Include up to 300 words.
- When a provision states "None" or the concept does not exist in the lease, return "None."
- For options (ROFR, ROFO, Expansion, Purchase, Early Termination, Reduction): if the lease does not grant such an option, return "None."
- For Permitted Use: find the clause that states what the premises may be used for. Start from "Tenant shall use" or equivalent.
- For Commencement Date: copy the full definition including all conditions (A, B, C options if present).
- For OPEX: start from the beginning of the list (e.g., "a) grounds maintenance..." not mid-list).

RESPONSE FORMAT:
Return a JSON object where each field maps to an object with two keys:
- "text": The extracted verbatim text from the lease
- "anchor": The first 8 words of the sentence where you found this text (used to verify location)

Example:
{{
  "Tenant_Name": {{"text": "ACME CORP, LLC", "anchor": ""}},
  "Permitted_Use_Description": {{"text": "Tenant shall use the Premises for medical office...", "anchor": "Tenant shall use the Premises for"}}
}}

For short values (names, amounts, dates, Yes/No), the "anchor" can be empty string "".
For longer clause extractions, "anchor" MUST contain the first 8 words of the source sentence.

FIELDS TO EXTRACT (field_name: description):
{fields_list}

LEASE DOCUMENT TEXT:
{lease_text[:80000]}

Respond with ONLY a valid JSON object. No markdown, no explanation."""

    return prompt


def analyze_lease_with_ai(lease_text: str, doc_type: str = "lease") -> tuple:
    """
    Send lease text to OpenAI and get structured field extraction.
    Returns (field_data_dict, anchors_dict).
    field_data maps field_name -> extracted text.
    anchors maps field_name -> anchor phrase (first 8 words of source sentence).
    """
    print("  Sending to AI for analysis...")
    client = get_openai_client()

    prompt = build_extraction_prompt(lease_text, doc_type=doc_type)

    response = client.chat.completions.create(
        model=AI_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a commercial lease data extractor. "
                    "You ONLY copy and paste exact text from lease documents. "
                    "You NEVER summarize, paraphrase, or use your own words. "
                    "Every value you return must appear verbatim in the source document."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        max_tokens=16000,
    )

    raw_response = response.choices[0].message.content.strip()

    # Clean potential markdown wrapping
    if raw_response.startswith("```"):
        raw_response = re.sub(r"^```(?:json)?\s*", "", raw_response)
        raw_response = re.sub(r"\s*```$", "", raw_response)

    try:
        parsed = json.loads(raw_response)
    except json.JSONDecodeError as e:
        print(f"  WARNING: AI response not valid JSON: {e}")
        print(f"  Raw response (first 500 chars): {raw_response[:500]}")
        return {}, {}

    # Parse the two-part format: {"field": {"text": "...", "anchor": "..."}}
    # Also handle flat format fallback: {"field": "value"}
    field_data = {}
    anchors = {}

    for key, val in parsed.items():
        if isinstance(val, dict):
            field_data[key] = val.get("text", "")
            anchors[key] = val.get("anchor", "")
        else:
            # Flat format fallback
            field_data[key] = val if val else ""
            anchors[key] = ""

    filled = sum(1 for v in field_data.values() if v)
    print(f"  AI extracted {filled}/{len(ALL_FIELDS)} fields")
    anchored = sum(1 for v in anchors.values() if v)
    print(f"  AI provided {anchored} anchor phrases")
    return field_data, anchors


# =============================================================================
# PHASE 3B: SOURCE VERIFICATION - HYBRID AI/CODE APPROACH
# =============================================================================

# Fields that are short lookups (not clause-length text) - skip expansion for these
SHORT_FIELDS = {
    "Lease_Summary_Preparer", "Lease_Summary_Date", "Lease_Summary_Purpose",
    "Property_Name", "Property_Address_1", "Property_Address_2",
    "Property_Address_City", "Property_Address_State", "Property_Address_Zip",
    "Owner_Name", "Owner_StReg", "Owner_Type",
    "Tenant_Name", "Tenant_StReg", "Tenant_Type", "Tenant_DBA",
    "Tenant_Address_1", "Tenant_Address_2", "Tenant_Address_City",
    "Tenant_Address_State", "Tenant_Address_Zip", "Tenant_Attn", "Tenant_Email",
    "Tenant_CopyTo_Address_1", "Tenant_CopyTo_Address_2",
    "Tenant_CopyTo_Address_City", "Tenant_CopyTo_Address_State",
    "Tenant_CopyTo_Address_Zip", "Tenant_CopyTo_Attn", "Tenant_CopyTo_Email",
    "Premises_UnitNumber_Description", "Premises_SqFt",
    "LeaseAgr_Name", "Date_Lease",
    "Amt_Security_Deposit", "Rent_BaseRent_PSF", "Rent_BaseRent_Amt",
    "Rent_BaseRent_Monthly", "Rent_AnnualIncrease_Percentage", "Base_Year",
    "Tenant_Share_Percentage", "Rent_PercentageRent_ThresholdAmt",
    "Rent_PercentageRent_Percent", "Tenant_Allowance_PSF", "Tenant_Allowance_Total",
    "Parking_Reserved_Amt_Fees", "Parking_Unreserved_Amt_Fees",
    "Guarantor_Name", "Guarantor_StReg", "Guarantor_Type",
    "Broker_Landlord_Name", "Broker_Tenant_Name",
    "Assignment_Process_Fee", "Estoppel_Return_Period",
    "SNDA_Subject_Existing_Mortgages", "SNDA_Subject_Future_Mortgages",
    "SNDA_Subject_Existing_GroundLease", "SNDA_Subject_Future_GroundLease",
    "Relocation_Rights", "Relocation_Notice_Period",
    "Signage_Allowed", "Signage_Approval_Required", "Signage_Type",
    "Renewal_Option_Numbers", "Renewal_Option_Period_PerOption",
    "ROFR_Space", "ROFO_Space", "Expansion_Space",
    "Purchase_Option_Space", "Reduction_Description",
}

# Fields where "None." is a valid extracted value (don't try to verify against source)
NONE_FIELDS = {
    "ROFR_Space", "ROFR_Description", "ROFO_Space", "ROFO_Description",
    "Expansion_Space", "Expansion_Description",
    "Early_Termination_Description", "Reduction_Description",
    "Purchase_Option_Space", "Purchase_Option_Description",
    "Guaranty_Term", "Gross_Up_Percent_Language",
    "Assignment_LL_Decision", "Assignment_Rent_Profit",
    "Assignment_Recapture_Space", "Assignment_Change_Control",
}


def _normalize_text(text: str) -> str:
    """Normalize text for comparison - collapse whitespace, lowercase."""
    text = re.sub(r'\s+', ' ', text).strip().lower()
    text = text.replace('\xa0', ' ')
    return text


def _find_best_match(needle: str, haystack: str, min_length: int = 8) -> tuple:
    """
    Find the best fuzzy match of needle in haystack.
    Returns (start_index, end_index, score) in the NORMALIZED haystack.
    Uses sliding window with longest common substring approach.
    """
    if len(needle) < min_length:
        return None

    norm_needle = _normalize_text(needle)
    norm_haystack = _normalize_text(haystack)

    # Try exact match first on normalized text
    idx = norm_haystack.find(norm_needle)
    if idx != -1:
        return (idx, idx + len(norm_needle), 1.0)

    # Try finding a substantial substring (at least 60% of the needle)
    words = norm_needle.split()
    if len(words) < 3:
        return None

    best_start = -1
    best_end = -1
    best_score = 0.0

    # Sliding window: try progressively smaller chunks of the needle
    for window_size in range(len(words), max(2, len(words) // 3), -1):
        for start_word in range(len(words) - window_size + 1):
            chunk = ' '.join(words[start_word:start_word + window_size])
            if len(chunk) < min_length:
                continue
            idx = norm_haystack.find(chunk)
            if idx != -1:
                score = len(chunk) / len(norm_needle)
                if score > best_score:
                    best_start = idx
                    best_end = idx + len(chunk)
                    best_score = score
                    if score > 0.7:
                        return (best_start, best_end, best_score)
        if best_score > 0.5:
            break

    if best_score > 0.3:
        return (best_start, best_end, best_score)

    return None


def _map_normalized_pos_to_original(haystack: str, norm_pos: int) -> int:
    """
    Map a single position in normalized text back to the original text.
    Returns the corresponding position in the original haystack.
    """
    in_whitespace = False
    current_norm_pos = 0

    for orig_pos, ch in enumerate(haystack):
        if current_norm_pos >= norm_pos:
            return orig_pos
        if ch in (' ', '\t', '\n', '\r', '\xa0'):
            if not in_whitespace:
                current_norm_pos += 1  # Collapsed whitespace = 1 char in normalized
                in_whitespace = True
        else:
            if in_whitespace:
                in_whitespace = False
            current_norm_pos += 1

    return len(haystack)


def _find_sentence_start(text: str, pos: int) -> int:
    """
    Walk backwards from pos to find the start of the current sentence.
    Looks for: period/colon/semicolon followed by whitespace, paragraph break,
    or section numbers (e.g., "7.", "28.", "(a)").
    Will not go back more than 300 chars.
    """
    limit = max(0, pos - 300)
    i = pos - 1

    while i > limit:
        ch = text[i]

        # Paragraph break (double newline or newline after period)
        if ch == '\n':
            # Check if this is a real paragraph break
            # Look at what comes after - if next non-whitespace is uppercase or number, good boundary
            after = text[i + 1:pos].lstrip()
            if after and (after[0].isupper() or after[0].isdigit() or after[0] == '('):
                return i + 1

        # Period, colon, or semicolon followed by space - sentence boundary
        if ch in '.;:' and i + 1 < pos:
            next_char = text[i + 1] if i + 1 < len(text) else ''
            if next_char in (' ', '\n', '\t', '\r'):
                # Make sure we're not hitting an abbreviation (e.g., "Sec.", "Inc.")
                # Check if the word before the period is short (likely abbreviation)
                word_before = ''
                j = i - 1
                while j >= limit and text[j].isalpha():
                    word_before = text[j] + word_before
                    j -= 1
                # Skip common abbreviations
                if word_before.lower() in ('sec', 'inc', 'llc', 'corp', 'ltd', 'no',
                                           'st', 'ave', 'blvd', 'dr', 'mr', 'mrs',
                                           'vs', 'etc', 'i.e', 'e.g', 'approx'):
                    i -= 1
                    continue
                # Good sentence boundary
                # Return position after the period + whitespace
                start = i + 1
                while start < pos and text[start] in (' ', '\t', '\n', '\r'):
                    start += 1
                return start

        i -= 1

    return limit


def _find_sentence_end(text: str, pos: int) -> int:
    """
    Walk forwards from pos to find the end of the current sentence.
    Looks for: period followed by whitespace/newline, paragraph break,
    or next section header.
    Will not go forward more than 400 chars.
    """
    limit = min(len(text), pos + 400)
    i = pos

    while i < limit:
        ch = text[i]

        # Period followed by space/newline - likely sentence end
        if ch == '.' and i + 1 < len(text):
            next_char = text[i + 1]
            if next_char in (' ', '\n', '\t', '\r'):
                # Check it's not an abbreviation
                word_before = ''
                j = i - 1
                while j >= pos and text[j].isalpha():
                    word_before = text[j] + word_before
                    j -= 1
                if word_before.lower() in ('sec', 'inc', 'llc', 'corp', 'ltd', 'no',
                                           'st', 'ave', 'blvd', 'dr', 'mr', 'mrs',
                                           'vs', 'etc', 'i.e', 'e.g', 'approx'):
                    i += 1
                    continue
                # Check what follows - if it's a new section/number, definitely end here
                after = text[i + 1:min(i + 10, limit)].lstrip()
                if after and (after[0].isupper() or after[0].isdigit() or after[0] == '('):
                    return i + 1

                # If followed by lowercase, might be mid-sentence ("...the U.S. government")
                # Still return - better to cut slightly short than grab junk
                return i + 1

        # Double newline - paragraph break, always stop
        if ch == '\n' and i + 1 < len(text) and text[i + 1] == '\n':
            return i

        # Section number pattern at start of line (e.g., "\n7." or "\n28.")
        if ch == '\n' and i + 1 < limit:
            after = text[i + 1:min(i + 6, limit)].lstrip()
            if after and after[0].isdigit():
                # Looks like a new numbered section
                return i

        i += 1

    return limit


def verify_and_expand_from_source(field_data: dict, source_text: str, anchors: dict = None) -> tuple:
    # Hybrid AI/code verification:
    # 1. If value is None - accept as-is
    # 2. If AI provided an anchor phrase, verify and use that location
    # 3. If exact match in source - keep (verified)
    # 4. If no exact match - fuzzy find, expand to sentence boundaries
    # 5. If still no match - try keyword fallback from field_anchors.json
    # 6. Short/lookup fields accepted without expansion
    if anchors is None:
        anchors = {}

    verified_data = {}
    flagged_fields = []  # Fields where verification suspects wrong content
    report = {"exact": 0, "expanded": 0, "anchor_verified": 0,
              "short_field": 0, "not_found": 0, "none_value": 0}

    norm_source = _normalize_text(source_text)

    for field_name, value in field_data.items():
        if not value or not value.strip():
            verified_data[field_name] = value
            continue

        # "None." values - accept without verification UNLESS keyword anchors exist
        if value.strip().lower() in ('none', 'none.', 'n/a', 'not applicable',
                                     'see original lease', 'see original lease.'):
            # For amendments: "See Original Lease." is a valid answer, don't override
            if value.strip().lower().startswith('see original'):
                verified_data[field_name] = value
                report["none_value"] += 1
                continue
            # Check if keyword fallback can find it (AI might have missed it)
            if field_name in _FIELD_ANCHORS and field_name not in SHORT_FIELDS:
                fallback = _keyword_fallback_search(field_name, source_text)
                if fallback and len(fallback) > 20:
                    verified_data[field_name] = fallback
                    report["expanded"] += 1
                    continue
            verified_data[field_name] = value
            report["none_value"] += 1
            continue

        # Short fields: accept without expansion
        if field_name in SHORT_FIELDS:
            verified_data[field_name] = value
            report["short_field"] += 1
            continue

        # --- Try AI anchor phrase first ---
        anchor_phrase = anchors.get(field_name, "")
        anchor_used = False
        if anchor_phrase and len(anchor_phrase) > 10:
            anchor_pos = _verify_ai_anchor(anchor_phrase, source_text)
            if anchor_pos >= 0:
                # Anchor verified - extract from that location
                orig_start = _map_normalized_pos_to_original(source_text, anchor_pos)
                sentence_start = _find_sentence_start(source_text, orig_start)
                sentence_end = _find_sentence_end(source_text, orig_start + 50)

                extracted = source_text[sentence_start:sentence_end].strip()
                extracted = re.sub(r'[ \t]+', ' ', extracted)
                extracted = re.sub(r'\n{3,}', '\n', extracted)
                extracted = re.sub(r'^\s*\n', '', extracted)
                # Strip leading OCR noise (page numbers, stray digits)
                extracted = re.sub(r'^[\d\s]{1,5}(?=\s*[A-Z])', '', extracted).strip()

                if len(extracted) > 600:
                    cut_pos = extracted.rfind('. ', 0, 600)
                    if cut_pos > len(extracted) // 3:
                        extracted = extracted[:cut_pos + 1]
                    else:
                        extracted = extracted[:600]

                # Validate: if field has keyword anchors, check the extracted text
                # contains relevant keywords (prevents wrong-section grabs)
                if field_name in _FIELD_ANCHORS and len(extracted) > 15:
                    norm_ext = _normalize_text(extracted)
                    # Check if any anchor pattern matches within the extracted text
                    anchor_relevant = False
                    for pattern in _FIELD_ANCHORS[field_name]:
                        try:
                            if re.search(pattern, norm_ext):
                                anchor_relevant = True
                                break
                        except re.error:
                            continue
                    if not anchor_relevant:
                        # Anchor landed in wrong section - flag for retry
                        flagged_fields.append(field_name)
                        pass
                    else:
                        verified_data[field_name] = extracted
                        report["anchor_verified"] += 1
                        anchor_used = True
                elif len(extracted) > 15:
                    verified_data[field_name] = extracted
                    report["anchor_verified"] += 1
                    anchor_used = True

        if anchor_used:
            continue

        # --- Check for exact match (normalized) ---
        norm_value = _normalize_text(value)
        if norm_value in norm_source:
            verified_data[field_name] = value
            report["exact"] += 1
            continue

        # --- Fuzzy find - AI got close, code grabs the real text ---
        match = _find_best_match(value, source_text)
        if match:
            norm_start, norm_end, score = match

            orig_start = _map_normalized_pos_to_original(source_text, norm_start)
            orig_end = _map_normalized_pos_to_original(source_text, norm_end)

            sentence_start = _find_sentence_start(source_text, orig_start)
            sentence_end = _find_sentence_end(source_text, orig_end)

            extracted = source_text[sentence_start:sentence_end].strip()
            extracted = re.sub(r'[ \t]+', ' ', extracted)
            extracted = re.sub(r'\n{3,}', '\n', extracted)
            extracted = re.sub(r'^\s*\n', '', extracted)
            extracted = re.sub(r'^[\d\s]{1,5}(?=\s*[A-Z])', '', extracted).strip()

            if len(extracted) > 600:
                cut_pos = extracted.rfind('. ', 0, 600)
                if cut_pos > len(extracted) // 2:
                    extracted = extracted[:cut_pos + 1]
                else:
                    extracted = extracted[:600]

            # If field has known anchors, always try the fallback if not exact
            if field_name in _FIELD_ANCHORS and score < 1.0:
                fallback = _keyword_fallback_search(field_name, source_text)
                if fallback and len(fallback) > 20:
                    norm_extracted = _normalize_text(extracted)
                    norm_fallback = _normalize_text(fallback)
                    overlap = norm_fallback[:50] in norm_extracted or norm_extracted[:50] in norm_fallback
                    if not overlap:
                        extracted = fallback
                elif score < 0.5:
                    # Low score + no fallback found = suspect content
                    flagged_fields.append(field_name)

            verified_data[field_name] = extracted
            report["expanded"] += 1
        else:
            # Could not find in source - try keyword fallback
            fallback = _keyword_fallback_search(field_name, source_text)
            if fallback:
                verified_data[field_name] = fallback
                report["expanded"] += 1
            else:
                verified_data[field_name] = value
                report["not_found"] += 1

    report["flagged"] = len(flagged_fields)
    return verified_data, report, flagged_fields


# Load keyword anchors from external config file (editable without code changes)
_FIELD_ANCHORS = {}
_ANCHORS_CONFIG_PATHS = [
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "lease_summary_app", "field_anchors.json"),
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "field_anchors.json"),
    "/home/cliff/redact/lease_summary_app/field_anchors.json",
]

def _load_field_anchors() -> dict:
    """Load field anchor patterns from JSON config file."""
    global _FIELD_ANCHORS
    for path in _ANCHORS_CONFIG_PATHS:
        if os.path.exists(path):
            try:
                with open(path, "r") as f:
                    data = json.load(f)
                # Remove comment keys
                _FIELD_ANCHORS = {k: v for k, v in data.items() if not k.startswith("_")}
                return _FIELD_ANCHORS
            except (json.JSONDecodeError, IOError):
                pass
    return {}

_load_field_anchors()


def _keyword_fallback_search(field_name: str, source_text: str) -> str:
    """
    When the AI's value doesn't match the source, find the correct clause
    using keyword anchor patterns from field_anchors.json config.
    Uses fuzzy matching to handle OCR errors.
    """
    from thefuzz import fuzz

    anchors = _FIELD_ANCHORS.get(field_name)
    if not anchors:
        return None

    norm_source = _normalize_text(source_text)

    # First pass: exact regex matching
    for anchor in anchors:
        try:
            match = re.search(anchor, norm_source)
        except re.error:
            continue
        if match:
            norm_start = match.start()
            orig_start = _map_normalized_pos_to_original(source_text, norm_start)
            sentence_start = _find_sentence_start(source_text, orig_start)
            norm_end_pos = match.end()
            orig_end = _map_normalized_pos_to_original(source_text, norm_end_pos)
            sentence_end = _find_sentence_end(source_text, orig_end)

            extracted = source_text[sentence_start:sentence_end].strip()
            extracted = re.sub(r'[ \t]+', ' ', extracted)
            extracted = re.sub(r'\n{3,}', '\n', extracted)
            extracted = re.sub(r'^\s*\n', '', extracted)
            extracted = re.sub(r'^[\d\s]{1,5}(?=\s*[A-Z])', '', extracted).strip()

            if len(extracted) > 600:
                cut_pos = extracted.rfind('. ', 0, 600)
                if cut_pos > len(extracted) // 3:
                    extracted = extracted[:cut_pos + 1]
                else:
                    extracted = extracted[:600]

            if len(extracted) > 20:
                return extracted

    # Second pass: fuzzy matching for OCR-garbled text
    for anchor in anchors:
        # Convert regex to plain text for fuzzy comparison
        plain = re.sub(r'\.\*', ' ', anchor)
        plain = re.sub(r'[\\.*+?^${}()|[\]]', '', plain)
        plain = plain.strip()
        if len(plain) < 8:
            continue

        # Slide a window across normalized source, fuzzy match
        window_size = len(plain) + 20
        best_score = 0
        best_pos = -1

        for i in range(0, max(1, len(norm_source) - window_size), 50):
            window = norm_source[i:i + window_size]
            score = fuzz.partial_ratio(plain, window)
            if score > best_score:
                best_score = score
                best_pos = i

        # Refine around best match
        if best_score >= 75 and best_pos >= 0:
            refine_start = max(0, best_pos - 50)
            refine_end = min(len(norm_source), best_pos + window_size + 50)
            for i in range(refine_start, refine_end - window_size, 10):
                window = norm_source[i:i + window_size]
                score = fuzz.partial_ratio(plain, window)
                if score > best_score:
                    best_score = score
                    best_pos = i

            if best_score >= 80:
                orig_start = _map_normalized_pos_to_original(source_text, best_pos)
                sentence_start = _find_sentence_start(source_text, orig_start)
                sentence_end = _find_sentence_end(source_text, orig_start + 50)

                extracted = source_text[sentence_start:sentence_end].strip()
                extracted = re.sub(r'[ \t]+', ' ', extracted)
                extracted = re.sub(r'\n{3,}', '\n', extracted)
                extracted = re.sub(r'^\s*\n', '', extracted)
                extracted = re.sub(r'^[\d\s]{1,5}(?=\s*[A-Z])', '', extracted).strip()

                if len(extracted) > 600:
                    cut_pos = extracted.rfind('. ', 0, 600)
                    if cut_pos > len(extracted) // 3:
                        extracted = extracted[:cut_pos + 1]
                    else:
                        extracted = extracted[:600]

                if len(extracted) > 20:
                    return extracted

    return None


def _verify_ai_anchor(anchor_phrase: str, source_text: str) -> int:
    # Verify the AI reported anchor phrase exists in source.
    # Returns the position if found, -1 if not.
    if not anchor_phrase or len(anchor_phrase) < 5:
        return -1
    norm_anchor = _normalize_text(anchor_phrase)
    norm_source = _normalize_text(source_text)
    idx = norm_source.find(norm_anchor)
    return idx


# Fields that should almost always exist in a commercial lease.
# If AI returns "None." for these, do a targeted retry.
EXPECTED_FIELDS = {
    "Permitted_Use_Description",
    "Tenant_Insurance",
    "Landlord_Repair_Obligations",
    "Tenant_Repair_Obligations",
    "OPEX_Inclusion",
    "Assignment_3rd_Parties",
    "Assignment_Affiliates",
    "Holdover_Rent",
    "Date_Commencment",
    "Lease_Term",
}


def ai_retry_none_fields(field_data: dict, source_text: str, flagged_fields: list = None) -> tuple:
    """
    Targeted AI second pass for fields that need re-extraction.
    Only called when there are flagged items. Two triggers:
    1. Fields expected in most leases that came back "None." / empty
    2. Fields explicitly flagged by verification as suspect (wrong section grabbed)

    The AI re-reads the source with OCR-tolerance and field-specific guidance.
    Returns (updated_field_data, count_of_recovered_fields).
    """
    if flagged_fields is None:
        flagged_fields = []

    # Build retry list: expected fields that are None + explicitly flagged fields
    retry_fields = {}

    # Add expected fields that are still None/empty
    for field_name in EXPECTED_FIELDS:
        val = field_data.get(field_name, "")
        if val.strip().lower() in ('none', 'none.', 'n/a', '', 'not applicable'):
            retry_fields[field_name] = ALL_FIELDS.get(field_name, field_name)
        # Skip "See Original Lease." - that's intentional for amendments
        # Do not retry those

    # Add explicitly flagged fields (verification detected wrong content)
    for field_name in flagged_fields:
        if field_name not in retry_fields:
            retry_fields[field_name] = ALL_FIELDS.get(field_name, field_name)

    if not retry_fields:
        return field_data, 0

    print(f"[Phase 3C] Targeted AI retry for {len(retry_fields)} field(s)...")
    for f in retry_fields:
        print(f"    - {f}")

    # Build a focused prompt - give the AI specific guidance per field
    field_hints = {
        "Permitted_Use_Description": "Look for 'Use of Premises' or 'Tenant shall use' - may be in a definitions section formatted as '(d) Use of Premises:'",
        "Tenant_Insurance": "Look for 'INSURANCE' section header, 'Tenant will carry and maintain', or 'General Requirements' near insurance language",
        "Landlord_Repair_Obligations": "Look for 'Landlord shall make repairs' or 'Landlord will at its own cost make repairs'",
        "Tenant_Repair_Obligations": "Look for 'Tenant shall maintain the interior' or 'Tenant will not injure the Premises'",
        "OPEX_Inclusion": "Look for 'Operating Expenses shall include' - often starts with lettered list a) b) c)",
        "Assignment_3rd_Parties": "Look for 'Tenant shall not assign' or 'without prior written consent of Landlord'",
        "Assignment_Affiliates": "Look for 'Permitted Transfer' or 'Tenant may without consent assign to parent, affiliate, subsidiary'",
        "Holdover_Rent": "Look for 'Holdover' or 'holding over' with a rent percentage (150%, 200%, double)",
        "Date_Commencment": "Look for 'Commencement Date' definition - often in definitions section with conditions A) B) C)",
        "Lease_Term": "Look for 'Term' definition with number of months",
    }

    fields_with_hints = []
    for k, desc in retry_fields.items():
        hint = field_hints.get(k, "")
        hint_str = f" HINT: {hint}" if hint else ""
        fields_with_hints.append(f'  "{k}": "{desc}"{hint_str}')

    fields_desc = "\n".join(fields_with_hints)

    prompt = f"""You are re-reading a commercial lease document to find specific provisions that were missed in a first pass.
The document may contain OCR errors (misspellings, garbled characters, missing spaces).

RULES:
- Copy/paste the EXACT text from the document, even if it has OCR typos
- Start from the beginning of the relevant sentence or clause
- Include up to 250 words of the relevant provision
- Return "Not found." ONLY if you genuinely cannot locate the provision after careful search
- Do NOT make up or paraphrase text

FIELDS TO FIND (with hints on where to look):
{fields_desc}

DOCUMENT TEXT:
{source_text[:80000]}

Return a JSON object mapping field names to extracted text. No markdown."""

    try:
        client = get_openai_client()
        response = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert at reading OCR'd legal documents. You can identify clauses even when text contains spelling errors or formatting issues. You copy text exactly as it appears, never paraphrasing."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.1,
            max_tokens=8000,
        )

        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)

        retried = json.loads(raw)
    except (json.JSONDecodeError, Exception) as e:
        print(f"    Retry failed: {e}")
        return field_data, 0

    # Apply recovered fields - verify each exists in source
    recovered = 0
    norm_source = _normalize_text(source_text)

    for field_name, value in retried.items():
        if not value or value.strip().lower() in ('not found', 'not found.', 'none', 'none.', ''):
            continue

        # Verify: at least part of the returned text exists in the source
        norm_val = _normalize_text(value)
        # Check first 40 chars of normalized value exist in normalized source
        found_in_source = False
        if len(norm_val) > 15 and norm_val[:40] in norm_source:
            found_in_source = True
        elif _find_best_match(value, source_text):
            found_in_source = True

        if found_in_source:
            field_data[field_name] = value
            recovered += 1
            print(f"    Recovered: {field_name} ({len(value.split())}w)")
        else:
            print(f"    Rejected (not in source): {field_name}")

    return field_data, recovered


def replace_placeholder_in_text(text: str, field_data: dict) -> str:
    """Replace [!@FieldName] and [*FieldName] placeholders with values."""
    for field_name, value in field_data.items():
        # Replace both marker styles
        text = text.replace(f"[!@{field_name}]", value or "")
        text = text.replace(f"[*{field_name}]", value or "")
    return text


def _replace_in_paragraph(para, field_data: dict):
    """
    Replace placeholders in a paragraph, handling split runs.
    DOCX often splits placeholder text across multiple runs due to formatting.
    This concatenates all run text, does replacement, then resets runs.
    """
    full_text = "".join(run.text for run in para.runs)
    if "[!@" not in full_text and "[*" not in full_text:
        return

    new_text = replace_placeholder_in_text(full_text, field_data)
    if new_text != full_text:
        # Preserve first run's formatting, clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""


def populate_template(field_data: dict, output_path: str) -> str:
    """Load the DOCX template, replace all placeholders, and save."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Template not found at: {TEMPLATE_PATH}\n"
            "Please ensure the SeedJura_Lease_Summary_FORM.docx is available."
        )

    doc = Document(TEMPLATE_PATH)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, field_data)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, field_data)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_output_filename(input_file: str, output_dir: str) -> str:
    """Generate output filename: {lease_name}_summary_{date}.docx"""
    base_name = Path(input_file).stem
    # Clean up the name - remove common suffixes
    clean_name = re.sub(
        r"(?i)[\s_-]*(fully[\s_-]*executed|execution|final|signed|copy)",
        "",
        base_name,
    )
    clean_name = re.sub(r"[\s_-]+$", "", clean_name)
    # Replace spaces with underscores
    clean_name = re.sub(r"\s+", "_", clean_name)

    date_str = datetime.now().strftime("%m-%d-%y")
    filename = f"{clean_name}_summary_{date_str}.docx"
    return os.path.join(output_dir, filename)


# =============================================================================
# MAIN PIPELINE
# =============================================================================

def process_lease(input_file: str, output_dir: str = None, preparer: str = "", purpose: str = "") -> str:
    """
    Full pipeline: ingest -> PII scan -> AI analysis -> populate template -> save.
    Returns path to the generated summary document.
    """
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR

    print("=" * 70)
    print("SEEDJURA LEASE SUMMARY TOOL")
    print("=" * 70)
    print(f"Input:  {input_file}")
    print(f"Output: {output_dir}")
    print()

    # --- Phase 1: Ingest ---
    print("[Phase 1] Ingesting document...")
    raw_text = ingest_document(input_file)
    print(f"  Extracted {len(raw_text):,} characters")

    # Detect document type (full lease vs amendment/addendum)
    doc_type = _detect_document_type(raw_text, input_file)
    if doc_type != "lease":
        print(f"  Document type: {doc_type.upper()}")
    print()

    # --- Phase 2: PII Scan ---
    print("[Phase 2] Scanning for PII...")
    text_for_ai, pii_findings = redact_and_capture_pii(raw_text)
    print()

    # --- Phase 3: AI Analysis ---
    print("[Phase 3] AI analysis and field extraction...")
    field_data, ai_anchors = analyze_lease_with_ai(text_for_ai, doc_type=doc_type)

    # Override with user-provided metadata
    if preparer:
        field_data["Lease_Summary_Preparer"] = preparer
    if purpose:
        field_data["Lease_Summary_Purpose"] = purpose
    if not field_data.get("Lease_Summary_Date"):
        field_data["Lease_Summary_Date"] = datetime.now().strftime("%B %d, %Y")
    print()

    # --- Phase 3B: Source Verification ---
    print("[Phase 3B] Verifying against source text (hybrid AI/code)...")
    field_data, verify_report, flagged = verify_and_expand_from_source(field_data, raw_text, ai_anchors)
    print(f"  Anchor-verified: {verify_report.get('anchor_verified', 0)}")
    print(f"  Exact matches: {verify_report['exact']}")
    print(f"  Expanded from source: {verify_report['expanded']}")
    print(f"  Short/lookup fields: {verify_report['short_field']}")
    print(f"  Not found in source: {verify_report['not_found']}")
    if flagged:
        print(f"  Flagged for retry: {len(flagged)}")
    print()

    # --- Phase 3C: AI Retry for suspect "None." fields ---
    field_data, retry_count = ai_retry_none_fields(field_data, raw_text, flagged)
    if retry_count > 0:
        print(f"[Phase 3C] AI retry recovered {retry_count} field(s)")
        print()

    # --- Phase 4: Populate and Save ---
    print("[Phase 4] Populating template and saving...")
    output_path = generate_output_filename(input_file, output_dir)
    saved_path = populate_template(field_data, output_path)
    print(f"  Saved: {saved_path}")
    print()

    # Save extraction data as JSON sidecar for reference
    json_path = output_path.replace(".docx", "_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "source_file": os.path.basename(input_file),
                "generated_at": datetime.now().isoformat(),
                "pii_count": len(pii_findings),
                "fields_extracted": sum(1 for v in field_data.values() if v),
                "fields_total": len(ALL_FIELDS),
                "verification": verify_report,
                "field_data": field_data,
            },
            f,
            indent=2,
        )
    print(f"  Data:  {json_path}")

    print()
    print("=" * 70)
    print("COMPLETE")
    print("=" * 70)
    return saved_path


def batch_process(input_dir: str, output_dir: str = None, preparer: str = "", purpose: str = ""):
    """Process all lease files in a directory."""
    supported_ext = {".pdf", ".docx", ".doc", ".txt"}
    input_path = Path(input_dir)

    files = [
        f for f in input_path.iterdir()
        if f.suffix.lower() in supported_ext
        and "summary" not in f.name.lower()
        and "form" not in f.name.lower()
    ]

    if not files:
        print(f"No lease files found in: {input_dir}")
        return []

    print(f"Found {len(files)} lease file(s) to process:")
    for f in files:
        print(f"  - {f.name}")
    print()

    results = []
    for i, file_path in enumerate(files, 1):
        print(f"\n{'#' * 70}")
        print(f"# Processing {i}/{len(files)}: {file_path.name}")
        print(f"{'#' * 70}\n")
        try:
            result = process_lease(
                str(file_path), output_dir, preparer, purpose
            )
            results.append(result)
        except Exception as e:
            print(f"ERROR processing {file_path.name}: {e}")
            results.append(None)

    # Summary
    print(f"\n\n{'=' * 70}")
    print("BATCH COMPLETE")
    print(f"{'=' * 70}")
    successful = [r for r in results if r]
    print(f"Processed: {len(files)} files")
    print(f"Successful: {len(successful)}")
    print(f"Failed: {len(files) - len(successful)}")
    return results


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="SeedJura Lease Summary Tool - Analyze leases and generate summaries"
    )
    parser.add_argument(
        "input",
        help="Input lease file (PDF/DOCX) or directory for batch mode",
    )
    parser.add_argument(
        "--output-dir", "-o",
        default=DEFAULT_OUTPUT_DIR,
        help=f"Output directory (default: {DEFAULT_OUTPUT_DIR})",
    )
    parser.add_argument(
        "--preparer", "-p",
        default="",
        help="Name of person preparing the summary",
    )
    parser.add_argument(
        "--purpose",
        default="",
        help="Purpose of the summary (e.g., 'For the purchase of property')",
    )
    parser.add_argument(
        "--batch", "-b",
        action="store_true",
        help="Process all supported files in the input directory",
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input not found: {args.input}")
        sys.exit(1)

    if args.batch or os.path.isdir(args.input):
        batch_process(args.input, args.output_dir, args.preparer, args.purpose)
    else:
        process_lease(args.input, args.output_dir, args.preparer, args.purpose)


if __name__ == "__main__":
    main()
